"""Job posting and candidate matching cog for the 508.dev Discord bot."""

from __future__ import annotations

import asyncio
import html
import ipaddress
import io
import logging
import re
import socket
from collections import OrderedDict
from dataclasses import dataclass
from typing import Any, Awaitable, Callable, Literal
from urllib.parse import urljoin, urlsplit

import aiohttp
import discord
from discord import app_commands
from discord.ext import commands

from five08.audit import update_person_discord_roles, upsert_discord_member
from five08.candidate_search import search_candidates
from five08.discord_bot.config import settings
from five08.discord_bot.utils.audit import DiscordAuditCogMixin
from five08.discord_bot.utils.role_decorators import (
    check_user_roles_with_hierarchy,
    require_role,
)
from five08.job_channels import (
    list_registered_job_post_channels,
    register_job_post_channel,
    unregister_job_post_channel,
)
from five08.job_match import (
    DISCORD_ROLES_EXCLUDE_FROM_SYNC,
    JobRequirements,
    extract_job_requirements,
)

logger = logging.getLogger(__name__)

MATCH_CANDIDATES_MAX_ATTACHMENT_SCAN = 5
MATCH_CANDIDATES_MAX_LINK_SCAN = 3
MATCH_CANDIDATES_MAX_ATTACHMENT_BYTES = 8 * 1024 * 1024
MATCH_CANDIDATES_MAX_ATTACHMENT_TEXT_CHARS = 10000
MATCH_CANDIDATES_MAX_LINK_TEXT_CHARS = 12000
MATCH_CANDIDATES_MAX_LINK_BYTES = 2 * 1024 * 1024
MATCH_CANDIDATES_MAX_LINK_REDIRECTS = 2
MATCH_CANDIDATES_MAX_POSTING_CHARS = 36000
MATCH_CANDIDATES_SUPPORTED_ATTACHMENT_EXTENSIONS = frozenset(
    {".txt", ".md", ".pdf", ".doc", ".docx", ".html", ".htm", ".rtf"}
)
MATCH_CANDIDATES_URL_PATTERN = re.compile(r"(?i)\bhttps?://[^\s<>()\[\]\"']+")
MATCH_CANDIDATES_JD_URL_HINTS = (
    "job",
    "jobs",
    "jd",
    "job-description",
    "position",
    "role",
    "career",
    "careers",
    "hiring",
    "greenhouse.io",
    "lever.co",
    "ashbyhq.com",
    "workable.com",
    "smartrecruiters.com",
    "linkedin.com/jobs",
    "docs.google.com/document",
    "notion.site",
)
AUTO_MATCH_DEDUPE_MAX = 10_000
# Exclude known-bad resume artifact from auto-match rendering.
AUTO_MATCH_EXCLUDED_RESUME_NAMES = frozenset({"Vladyslav_Stryzhak.pdf"})
IPAddress = ipaddress.IPv4Address | ipaddress.IPv6Address
JobWatchChannel = discord.ForumChannel


class MatchResumeSelectView(discord.ui.View):
    """View containing a resume download select for match results."""

    def __init__(self, options: list[tuple[str, str, str]]) -> None:
        super().__init__(timeout=600)  # 10 minute timeout
        self.add_item(MatchResumeSelect(options))


class MatchResumeSelect(discord.ui.Select):
    """Select menu for downloading a resume from match results."""

    def __init__(self, options: list[tuple[str, str, str]]) -> None:
        discord_options: list[discord.SelectOption] = []
        self._resume_lookup: dict[str, str] = {}

        for contact_name, resume_id, resume_name in options[:25]:
            label = contact_name.strip() or "Unknown"
            if len(label) > 100:
                label = label[:97] + "..."
            description = resume_name.strip() or "Resume"
            if len(description) > 100:
                description = description[:97] + "..."
            discord_options.append(
                discord.SelectOption(
                    label=label,
                    value=resume_id,
                    description=description,
                )
            )
            self._resume_lookup[resume_id] = contact_name

        super().__init__(
            placeholder="Download a resume...",
            min_values=1,
            max_values=1,
            options=discord_options,
            custom_id="match_resume_select",
        )

    async def callback(self, interaction: discord.Interaction) -> None:
        jobs_cog = interaction.client.get_cog("JobsCog")  # type: ignore[attr-defined]
        crm_cog = interaction.client.get_cog("CRMCog")  # type: ignore[attr-defined]
        try:
            if jobs_cog is None or not hasattr(jobs_cog, "_audit_command"):
                await interaction.response.send_message(
                    "❌ Job matching functionality not available.",
                    ephemeral=True,
                )
                return
            download_method = getattr(crm_cog, "_download_and_send_resume", None)
            if crm_cog is None or not callable(download_method):
                await interaction.response.send_message(
                    "❌ CRM functionality not available.",
                    ephemeral=True,
                )
                return

            resume_id = self.values[0]
            contact_name = self._resume_lookup.get(resume_id, "Unknown")

            await interaction.response.defer(ephemeral=True)

            download_ok = await download_method(interaction, contact_name, resume_id)
            try:
                jobs_cog._audit_command(
                    interaction=interaction,
                    action="crm.match_candidates_resume_select",
                    result="success" if download_ok else "error",
                    metadata={"contact_name": contact_name},
                    resource_type="crm_contact",
                    resource_id=resume_id,
                )
            except Exception as audit_exc:
                logger.error("Audit write failed in match resume select: %s", audit_exc)
        except Exception as exc:
            logger.error("Unexpected error in match resume select: %s", exc)
            if jobs_cog is not None and hasattr(jobs_cog, "_audit_command"):
                try:
                    jobs_cog._audit_command(
                        interaction=interaction,
                        action="crm.match_candidates_resume_select",
                        result="error",
                        metadata={"error": str(exc)},
                        resource_type="discord_ui_action",
                        resource_id=self.values[0] if self.values else None,
                    )
                except Exception as audit_exc:
                    logger.error(
                        "Audit write failed in match resume select: %s", audit_exc
                    )
            await interaction.followup.send(
                "❌ An unexpected error occurred while downloading the resume.",
                ephemeral=True,
            )


@dataclass(frozen=True)
class ThreadPost:
    """Thread opener content split from forum tags.

    Auto-match and manual `/match-candidates` both need the same starter message,
    but only prepend tag names at the last moment so the raw text stays reusable.
    """

    starter: discord.Message
    tags: list[str]


class JobsCog(DiscordAuditCogMixin, commands.Cog):
    """Job posting and candidate matching workflows."""

    def __init__(self, bot: commands.Bot) -> None:
        self.bot = bot
        self._init_audit_logger()
        self._jobs_channels_by_guild: dict[int, set[int]] = {}
        self._auto_matched_thread_ids: OrderedDict[int, None] = OrderedDict()
        self._auto_matched_thread_lock = asyncio.Lock()
        self._startup_sync_done = False
        self._startup_sync_lock = asyncio.Lock()

    @staticmethod
    def _resolve_jobs_channel_target(
        interaction: discord.Interaction,
        channel: JobWatchChannel | None,
    ) -> JobWatchChannel | None:
        """Resolve explicit/implicit channel target for job-post registration."""
        if channel is not None:
            return channel

        current = interaction.channel
        if isinstance(current, discord.ForumChannel):
            return current

        if isinstance(current, discord.Thread) and isinstance(
            current.parent, discord.ForumChannel
        ):
            return current.parent

        return None

    async def _refresh_jobs_channel_cache(self, guild_id: int) -> set[int]:
        """Load registered job-post channels for one guild from Postgres."""
        raw_ids = await asyncio.to_thread(
            list_registered_job_post_channels,
            settings,
            guild_id=str(guild_id),
        )
        parsed_ids: set[int] = set()
        for raw_id in raw_ids:
            try:
                parsed_ids.add(int(raw_id))
            except ValueError:
                logger.warning(
                    "Skipping invalid job_post_channels row guild_id=%s channel_id=%s",
                    guild_id,
                    raw_id,
                )
        self._jobs_channels_by_guild[guild_id] = parsed_ids
        return parsed_ids

    def _is_jobs_channel_registered(self, guild_id: int, channel_id: int) -> bool:
        """Return whether a channel is registered for automatic job matching."""
        return channel_id in self._jobs_channels_by_guild.get(guild_id, set())

    async def _refresh_jobs_channel_cache_if_missing(self, guild_id: int) -> bool:
        """Ensure guild cache is loaded, retrying after startup-load failures."""
        if guild_id in self._jobs_channels_by_guild:
            return True
        try:
            await self._refresh_jobs_channel_cache(guild_id)
            return True
        except Exception as exc:
            logger.warning(
                "Failed to refresh jobs-channel cache for guild=%s: %s",
                guild_id,
                exc,
            )
            return False

    async def _mark_thread_auto_matched(self, thread_id: int) -> bool:
        """Deduplicate automatic matching when multiple events race."""
        async with self._auto_matched_thread_lock:
            if thread_id in self._auto_matched_thread_ids:
                self._auto_matched_thread_ids.move_to_end(thread_id)
                return False
            self._auto_matched_thread_ids[thread_id] = None
            if len(self._auto_matched_thread_ids) > AUTO_MATCH_DEDUPE_MAX:
                self._auto_matched_thread_ids.popitem(last=False)
            return True

    async def _unmark_thread_auto_matched(self, thread_id: int) -> None:
        """Allow retry when a thread was marked but processing could not start."""
        async with self._auto_matched_thread_lock:
            self._auto_matched_thread_ids.pop(thread_id, None)

    @staticmethod
    async def _read_thread_post(thread: discord.Thread) -> ThreadPost | None:
        """Read starter message content for a job thread, including forum tags."""
        starter = thread.starter_message
        if starter is None:
            try:
                starter = await thread.fetch_message(thread.id)
            except Exception:
                starter = None

        if starter is None:
            return None

        tags = [t.name for t in thread.applied_tags] if thread.applied_tags else []
        return ThreadPost(starter=starter, tags=tags)

    def _build_job_match_header_and_mentions(
        self,
        *,
        requirements: JobRequirements,
        candidates_count: int,
        guild: discord.Guild | None,
    ) -> tuple[list[str], str | None, list[int], str | None, list[int]]:
        """Build header lines plus discord/locality role mention details."""
        header_parts: list[str] = []
        role_mentions_line: str | None = None
        locality_mentions_line: str | None = None
        role_mentions_role_ids: list[int] = []
        locality_mentions_role_ids: list[int] = []
        excluded_role_names = {
            name.casefold() for name in DISCORD_ROLES_EXCLUDE_FROM_SYNC
        }

        def dedupe_role_names(role_names: list[str]) -> list[str]:
            seen: set[str] = set()
            deduped: list[str] = []
            for role_name in role_names:
                cleaned = role_name.strip()
                if not cleaned:
                    continue
                key = cleaned.casefold()
                if key in seen:
                    continue
                seen.add(key)
                deduped.append(cleaned)
            return deduped

        def build_role_mentions(role_names: list[str]) -> tuple[list[str], list[int]]:
            if not role_names:
                return [], []
            if guild is None:
                return [f"`{r}`" for r in role_names], []

            role_id_map = self._get_role_id_cache().get(guild.id)
            if role_id_map is None:
                self._refresh_role_id_cache(guild)
                role_id_map = self._get_role_id_cache().get(guild.id, {})

            mentions: list[str] = []
            seen_mentions: set[str] = set()
            allowed_role_ids: list[int] = []
            seen_role_ids: set[int] = set()
            for role_name in role_names:
                normalized_role_name = role_name.casefold()
                if normalized_role_name in excluded_role_names:
                    continue
                role_id = role_id_map.get(normalized_role_name)
                if role_id is not None:
                    mention = f"<@&{role_id}>"
                    if mention not in seen_mentions:
                        seen_mentions.add(mention)
                        mentions.append(mention)
                    if role_id not in seen_role_ids:
                        seen_role_ids.add(role_id)
                        allowed_role_ids.append(role_id)
                    continue
                role = None
                for guild_role in guild.roles:
                    guild_role_name = guild_role.name.casefold()
                    if guild_role_name in excluded_role_names:
                        continue
                    if guild_role_name == normalized_role_name:
                        role = guild_role
                        break
                if role is not None:
                    if role.mention not in seen_mentions:
                        seen_mentions.add(role.mention)
                        mentions.append(role.mention)
                    if role.id not in seen_role_ids:
                        seen_role_ids.add(role.id)
                        allowed_role_ids.append(role.id)
                else:
                    mention = f"`{role_name}`"
                    if mention not in seen_mentions:
                        seen_mentions.add(mention)
                        mentions.append(mention)
            return mentions, allowed_role_ids

        if requirements.title:
            header_parts.append(f"**{requirements.title}**")
        if requirements.discord_role_types:
            role_types = dedupe_role_names(requirements.discord_role_types)
            if role_types:
                role_mentions, role_ids = build_role_mentions(role_types)
                if role_mentions:
                    role_mentions_line = "Discord roles: " + ", ".join(role_mentions)
                    role_mentions_role_ids = role_ids

        locality_role_names: list[str] = []
        location_text_parts: list[str] = []
        if requirements.raw_location_text:
            location_text_parts.append(requirements.raw_location_text)
        if requirements.preferred_timezones:
            location_text_parts.extend(requirements.preferred_timezones)
        location_text = " ".join(location_text_parts).casefold()

        # Job postings describe location in loose prose, so we infer the broad
        # locality roles from both the normalized location type and the raw text.
        if requirements.location_type == "us_only" or "united states" in location_text:
            locality_role_names.append("USA")
        if "usa" in location_text:
            locality_role_names.append("USA")
        if (
            "europe" in location_text
            or "emea" in location_text
            or "e.u." in location_text
        ):
            locality_role_names.append("Europe")
        if (
            "americas" in location_text
            or "latin america" in location_text
            or "latam" in location_text
        ):
            locality_role_names.append("Americas")
        if "north america" in location_text or "south america" in location_text:
            locality_role_names.append("Americas")
        if (
            "asia" in location_text
            or "apac" in location_text
            or "asia pacific" in location_text
        ):
            locality_role_names.append("Asia")
        if "japan" in location_text:
            locality_role_names.append("Japan")
        if "taiwan" in location_text:
            locality_role_names.append("Taiwan")
        if "africa" in location_text:
            locality_role_names.append("Africa")

        if requirements.preferred_timezones:
            for tz in requirements.preferred_timezones:
                tz_prefix = (
                    tz.split("/", 1)[0].casefold() if "/" in tz else tz.casefold()
                )
                if tz_prefix == "europe":
                    locality_role_names.append("Europe")
                elif tz_prefix == "america":
                    locality_role_names.append("Americas")
                elif tz_prefix == "asia":
                    locality_role_names.append("Asia")
                elif tz_prefix == "africa":
                    locality_role_names.append("Africa")
                if tz.casefold() == "asia/tokyo":
                    locality_role_names.append("Japan")
                if tz.casefold() == "asia/taipei":
                    locality_role_names.append("Taiwan")

        locality_role_names = [
            role_name
            for role_name in dedupe_role_names(locality_role_names)
            if role_name.casefold() not in excluded_role_names
        ]
        if locality_role_names:
            locality_mentions, role_ids = build_role_mentions(locality_role_names)
            if locality_mentions:
                locality_mentions_line = "Locality: " + ", ".join(locality_mentions)
                locality_mentions_role_ids = role_ids

        if requirements.required_skills:
            header_parts.append(
                "Skills: "
                + ", ".join(f"`{s}`" for s in requirements.required_skills[:8])
            )
        if requirements.seniority:
            header_parts.append(f"Seniority: `{requirements.seniority}`")
        if requirements.location_type == "us_only":
            header_parts.append("📍 US only")
        elif requirements.raw_location_text:
            header_parts.append(f"📍 {requirements.raw_location_text}")

        header_lines: list[str] = ["## Job Match Results"]
        if header_parts:
            header_lines.append(" · ".join(header_parts))
        header_lines.append(f"Found **{candidates_count}** candidate(s).")

        return (
            header_lines,
            role_mentions_line,
            role_mentions_role_ids,
            locality_mentions_line,
            locality_mentions_role_ids,
        )

    @staticmethod
    def _paginate_match_lines(lines: list[str]) -> list[str]:
        """Paginate long match output lines into Discord-sized messages."""
        messages: list[str] = []
        current = ""
        for line in lines:
            candidate_block = line + "\n"
            while len(candidate_block) > 1900:
                if current:
                    messages.append(current.rstrip())
                    current = ""
                messages.append(candidate_block[:1900].rstrip())
                candidate_block = candidate_block[1900:]
            if len(current) + len(candidate_block) > 1900:
                if current:
                    messages.append(current.rstrip())
                current = candidate_block
            else:
                current += candidate_block
        if current.strip():
            messages.append(current.rstrip())
        return messages

    @staticmethod
    def _build_match_candidate_lines(
        *,
        candidates: list[Any],
        crm_base: str,
    ) -> tuple[list[str], list[tuple[str, str, str]]]:
        """Build candidate result lines and resume options for match output."""
        lines: list[str] = []
        resume_options: list[tuple[str, str, str]] = []

        for i, candidate in enumerate(candidates, start=1):
            label = "**[Member]**" if candidate.is_member else "[Prospect]"
            raw_crm_name = (
                candidate.crm_name.strip()
                if isinstance(candidate.crm_name, str) and candidate.crm_name.strip()
                else None
            )
            raw_display_name = (
                candidate.name.strip()
                if isinstance(candidate.name, str) and candidate.name.strip()
                else None
            )
            resolved_name = discord.utils.escape_mentions(
                raw_crm_name or raw_display_name or "Unknown"
            )
            normalized_discord_username = (
                candidate.discord_username.strip()
                if isinstance(candidate.discord_username, str)
                else None
            )
            discord_username = (
                discord.utils.escape_mentions(normalized_discord_username.lstrip("@"))
                if normalized_discord_username
                and normalized_discord_username.lstrip("@")
                else None
            )
            crm_link = (
                f"{crm_base}/#Contact/view/{candidate.crm_contact_id}"
                if candidate.has_crm_link and candidate.crm_contact_id
                else None
            )
            if crm_link:
                display_name = f"[{resolved_name}](<{crm_link}>)"
            else:
                display_name = resolved_name

            parts = [f"{i}. {label} {display_name}"]
            if discord_username:
                parts.append(f"`@{discord_username}`")

            if candidate.linkedin:
                parts.append(f"[LinkedIn](<{candidate.linkedin}>)")
            if (
                candidate.latest_resume_id
                and candidate.latest_resume_name
                and candidate.latest_resume_name not in AUTO_MATCH_EXCLUDED_RESUME_NAMES
            ):
                safe_resume_name = discord.utils.escape_mentions(
                    candidate.latest_resume_name
                )
                resume_options.append(
                    (resolved_name, candidate.latest_resume_id, safe_resume_name)
                )

            skill_info: list[str] = []
            match_score = getattr(candidate, "match_score", None)
            if isinstance(match_score, (int, float)):
                skill_info.append(f"score: {match_score:.1f}")
            if candidate.matched_required_skills:
                skill_info.append(
                    "✅ "
                    + ", ".join(f"`{s}`" for s in candidate.matched_required_skills[:5])
                )
            if candidate.matched_discord_roles:
                skill_info.append(
                    "🏷️ " + ", ".join(f"`{r}`" for r in candidate.matched_discord_roles)
                )
            if candidate.seniority:
                skill_info.append(f"seniority: `{candidate.seniority}`")
            if candidate.timezone:
                skill_info.append(f"tz: `{candidate.timezone}`")
            if skill_info:
                parts.append("   " + " · ".join(skill_info))

            lines.append("\n".join(parts))

        return lines, resume_options

    @staticmethod
    def _resume_file_extension(filename: str | None) -> str:
        if not filename or "." not in filename:
            return ""
        return "." + filename.rsplit(".", 1)[-1].lower()

    def _extract_resume_text(
        self,
        file_content: bytes,
        *,
        filename: str | None,
    ) -> str:
        extension = self._resume_file_extension(filename)
        extracted_text = ""

        try:
            if extension == ".pdf":
                from pdfminer.high_level import extract_text as extract_pdf_text

                extracted_text = extract_pdf_text(io.BytesIO(file_content)).strip()
            elif extension == ".docx":
                from docx import Document

                document = Document(io.BytesIO(file_content))
                chunks: list[str] = []
                for paragraph in document.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        chunks.append(text)
                for table in document.tables:
                    for row in table.rows:
                        row_cells = [
                            cell.text.strip() for cell in row.cells if cell.text.strip()
                        ]
                        if row_cells:
                            chunks.append(" | ".join(row_cells))
                extracted_text = "\n".join(chunks).strip()
            elif extension == ".doc":
                extracted_text = file_content.decode("utf-8", errors="ignore")
                extracted_text = re.sub(r"[^\x20-\x7E\n\r\t]", " ", extracted_text)
                extracted_text = re.sub(r"\s+", " ", extracted_text).strip()
            else:
                extracted_text = file_content.decode("utf-8", errors="ignore").strip()
        except Exception as exc:
            logger.warning(
                "Failed to extract resume text filename=%s extension=%s error=%s",
                filename,
                extension,
                exc,
            )

        if extracted_text:
            return extracted_text
        return file_content.decode("utf-8", errors="ignore")

    @staticmethod
    def _extract_urls_from_text(text: str) -> list[str]:
        if not text:
            return []
        urls: list[str] = []
        seen: set[str] = set()
        for raw in MATCH_CANDIDATES_URL_PATTERN.findall(text):
            normalized = raw.rstrip(".,);]>")
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            urls.append(normalized)
        return urls

    @staticmethod
    def _is_probable_jd_url(url: str) -> bool:
        lowered = url.casefold()
        return any(hint in lowered for hint in MATCH_CANDIDATES_JD_URL_HINTS)

    @staticmethod
    def _strip_html_to_text(raw_html: str) -> str:
        without_scripts = re.sub(
            r"(?is)<(script|style|noscript).*?>.*?</\1>",
            " ",
            raw_html,
        )
        text_only = re.sub(r"(?s)<[^>]+>", " ", without_scripts)
        unescaped = html.unescape(text_only)
        return re.sub(r"\s+", " ", unescaped).strip()

    @staticmethod
    def _parse_ip_literal(value: str) -> IPAddress | None:
        try:
            return ipaddress.ip_address(value)
        except ValueError:
            return None

    @staticmethod
    def _is_public_ip(value: IPAddress) -> bool:
        return not (
            value.is_private
            or value.is_loopback
            or value.is_link_local
            or value.is_multicast
            or value.is_reserved
            or value.is_unspecified
        )

    @classmethod
    def _hostname_resolves_publicly(cls, host: str) -> bool:
        if host in {"localhost", "localhost.localdomain"}:
            return False

        ip_literal = cls._parse_ip_literal(host)
        if ip_literal is not None:
            return cls._is_public_ip(ip_literal)

        try:
            addr_infos = socket.getaddrinfo(host, None, proto=socket.IPPROTO_TCP)
        except socket.gaierror:
            return False
        except Exception:
            return False

        resolved_ips: set[IPAddress] = set()
        for _, _, _, _, sockaddr in addr_infos:
            if not sockaddr:
                continue
            ip_text = str(sockaddr[0]).strip()
            parsed_ip = cls._parse_ip_literal(ip_text)
            if parsed_ip is None:
                continue
            resolved_ips.add(parsed_ip)

        if not resolved_ips:
            return False
        return all(cls._is_public_ip(parsed_ip) for parsed_ip in resolved_ips)

    async def _validate_match_candidates_url(self, candidate_url: str) -> str | None:
        try:
            parsed = urlsplit(candidate_url)
        except Exception:
            return "Job description URL is invalid."

        if parsed.scheme.lower() != "https":
            return "Job description URL must use https."

        if parsed.username or parsed.password:
            return "Job description URL must not include credentials."

        host = (parsed.hostname or "").strip().lower().rstrip(".")
        if not host:
            return "Job description URL must include a hostname."

        resolves_publicly = await asyncio.to_thread(
            self._hostname_resolves_publicly, host
        )
        if not resolves_publicly:
            return "Job description URL host resolves to a non-public address."

        return None

    async def _read_match_candidates_attachment_text(
        self, attachment: discord.Attachment
    ) -> str | None:
        filename = attachment.filename or ""
        extension = self._resume_file_extension(filename)
        content_type = (attachment.content_type or "").strip().lower()
        is_supported_type = (
            extension in MATCH_CANDIDATES_SUPPORTED_ATTACHMENT_EXTENSIONS
            or content_type.startswith("text/")
        )
        if not is_supported_type:
            return None
        if attachment.size and attachment.size > MATCH_CANDIDATES_MAX_ATTACHMENT_BYTES:
            logger.info(
                "Skipping oversized match-candidates attachment filename=%s size=%s",
                filename,
                attachment.size,
            )
            return None

        try:
            file_content = await attachment.read()
        except Exception as exc:
            logger.warning(
                "Failed reading match-candidates attachment filename=%s error=%s",
                filename,
                exc,
            )
            return None

        extracted = self._extract_resume_text(file_content, filename=filename).strip()
        if not extracted:
            return None
        if len(extracted) > MATCH_CANDIDATES_MAX_ATTACHMENT_TEXT_CHARS:
            return (
                extracted[:MATCH_CANDIDATES_MAX_ATTACHMENT_TEXT_CHARS].rstrip()
                + "\n[attachment text truncated]"
            )
        return extracted

    async def _fetch_match_candidates_link_text(self, url: str) -> str | None:
        timeout = aiohttp.ClientTimeout(total=12)
        headers = {"User-Agent": "508-job-match/1.0"}
        current_url = url
        try:
            async with aiohttp.ClientSession(
                timeout=timeout, headers=headers
            ) as session:
                for _ in range(MATCH_CANDIDATES_MAX_LINK_REDIRECTS + 1):
                    validation_error = await self._validate_match_candidates_url(
                        current_url
                    )
                    if validation_error:
                        logger.info(
                            "Skipping JD link fetch url=%s reason=%s",
                            current_url,
                            validation_error,
                        )
                        return None

                    async with session.get(
                        current_url, allow_redirects=False
                    ) as response:
                        if response.status in {301, 302, 303, 307, 308}:
                            redirect_to = response.headers.get("Location")
                            if not redirect_to:
                                return None
                            current_url = urljoin(current_url, redirect_to)
                            continue
                        if response.status >= 400:
                            return None

                        content_type = (
                            response.headers.get("Content-Type", "")
                            .split(";")[0]
                            .strip()
                            .lower()
                        )
                        is_supported_content_type = content_type in {
                            "application/pdf",
                            "application/xhtml+xml",
                        } or content_type.startswith("text/")
                        if content_type and not is_supported_content_type:
                            logger.info(
                                "Skipping unsupported JD content type url=%s content_type=%s",
                                current_url,
                                content_type,
                            )
                            return None

                        content_length = response.headers.get("Content-Length")
                        if content_length:
                            try:
                                content_len_bytes = int(content_length)
                            except (TypeError, ValueError):
                                pass
                            else:
                                if content_len_bytes > MATCH_CANDIDATES_MAX_LINK_BYTES:
                                    logger.info(
                                        "Skipping oversized JD link url=%s bytes=%s",
                                        current_url,
                                        content_len_bytes,
                                    )
                                    return None

                        final_url = str(response.url)
                        raw_chunks = bytearray()
                        async for chunk in response.content.iter_chunked(8192):
                            if not chunk:
                                continue
                            raw_chunks.extend(chunk)
                            if len(raw_chunks) > MATCH_CANDIDATES_MAX_LINK_BYTES:
                                logger.info(
                                    "Skipping oversized JD body url=%s bytes>%s",
                                    final_url,
                                    MATCH_CANDIDATES_MAX_LINK_BYTES,
                                )
                                return None
                        raw = bytes(raw_chunks)
                        break
                else:
                    logger.info("Skipping JD link after too many redirects url=%s", url)
                    return None
        except Exception as exc:
            logger.info("Failed fetching JD link url=%s error=%s", url, exc)
            return None

        if not raw:
            return None

        lower_final = final_url.casefold()
        if content_type == "application/pdf" or lower_final.endswith(".pdf"):
            text = self._extract_resume_text(raw, filename="linked_jd.pdf")
        elif content_type in {"text/plain", "text/markdown"}:
            text = raw.decode("utf-8", errors="ignore")
        else:
            text = self._strip_html_to_text(raw.decode("utf-8", errors="ignore"))

        cleaned = re.sub(r"\s+", " ", text).strip()
        if not cleaned:
            return None
        if len(cleaned) > MATCH_CANDIDATES_MAX_LINK_TEXT_CHARS:
            return cleaned[:MATCH_CANDIDATES_MAX_LINK_TEXT_CHARS].rstrip()
        return cleaned

    async def _build_match_candidates_posting(
        self, starter: discord.Message
    ) -> tuple[str, dict[str, Any]]:
        base_text = starter.content.strip()
        attachment_chunks: list[str] = []
        attachment_urls: list[str] = []
        scanned_attachments = 0

        # We cap attachment and link scanning so one noisy job post does not turn
        # into an unbounded amount of file parsing, network fetches, or LLM input.
        for attachment in starter.attachments[:MATCH_CANDIDATES_MAX_ATTACHMENT_SCAN]:
            scanned_attachments += 1
            extracted = await self._read_match_candidates_attachment_text(attachment)
            if not extracted:
                continue
            display_name = attachment.filename or "attachment"
            attachment_chunks.append(f"Attachment {display_name}:\n{extracted}")
            attachment_urls.extend(self._extract_urls_from_text(extracted))

        candidate_urls: list[str] = []
        candidate_urls.extend(self._extract_urls_from_text(base_text))
        candidate_urls.extend(attachment_urls)
        for embed in starter.embeds:
            if embed.url:
                candidate_urls.append(embed.url)

        deduped_urls: list[str] = []
        seen_urls: set[str] = set()
        for raw_url in candidate_urls:
            parsed = urlsplit(raw_url)
            if not parsed.scheme or not parsed.netloc:
                continue
            normalized = raw_url.strip()
            if not normalized or normalized in seen_urls:
                continue
            seen_urls.add(normalized)
            deduped_urls.append(normalized)

        likely_jd_urls = [u for u in deduped_urls if self._is_probable_jd_url(u)]
        urls_to_fetch = likely_jd_urls[:MATCH_CANDIDATES_MAX_LINK_SCAN]

        fetched_link_chunks: list[str] = []
        fetched_links: list[str] = []
        for url in urls_to_fetch:
            link_text = await self._fetch_match_candidates_link_text(url)
            if not link_text:
                continue
            fetched_links.append(url)
            fetched_link_chunks.append(f"Source {url}:\n{link_text}")

        sections: list[str] = []
        if base_text:
            sections.append(base_text)
        if attachment_chunks:
            sections.append(
                "Attached job description documents (extracted text):\n\n"
                + "\n\n".join(attachment_chunks)
            )
        if deduped_urls:
            sections.append("Referenced links:\n" + "\n".join(deduped_urls))
        if fetched_link_chunks:
            sections.append(
                "Referenced job description pages (extracted text):\n\n"
                + "\n\n".join(fetched_link_chunks)
            )

        posting = "\n\n".join(part for part in sections if part.strip()).strip()
        if len(posting) > MATCH_CANDIDATES_MAX_POSTING_CHARS:
            posting = (
                posting[:MATCH_CANDIDATES_MAX_POSTING_CHARS].rstrip() + "\n[truncated]"
            )

        metadata = {
            "starter_has_text": bool(base_text),
            "attachments_seen": len(starter.attachments),
            "attachments_scanned": scanned_attachments,
            "attachments_extracted": len(attachment_chunks),
            "links_discovered": len(deduped_urls),
            "links_fetched": len(fetched_links),
        }
        return posting, metadata

    async def _publish_match_results(
        self,
        *,
        send: Callable[..., Awaitable[Any]],
        requirements: JobRequirements,
        candidates: list[Any],
        guild: discord.Guild | None,
    ) -> None:
        """Send the formatted match output for both manual and automatic runs."""
        (
            header_lines,
            role_mentions_line,
            role_mentions_role_ids,
            locality_mentions_line,
            locality_mentions_role_ids,
        ) = self._build_job_match_header_and_mentions(
            requirements=requirements,
            candidates_count=len(candidates),
            guild=guild,
        )

        safe_mentions = discord.AllowedMentions(
            roles=False,
            users=False,
            everyone=False,
        )
        for chunk in self._paginate_match_lines(header_lines):
            await send(chunk, allowed_mentions=safe_mentions)
        if role_mentions_line:
            allowed_role_mentions = (
                discord.AllowedMentions(
                    roles=[discord.Object(id=rid) for rid in role_mentions_role_ids],
                    users=False,
                    everyone=False,
                )
                if role_mentions_role_ids
                else safe_mentions
            )
            for chunk in self._paginate_match_lines([role_mentions_line]):
                await send(chunk, allowed_mentions=allowed_role_mentions)
        if locality_mentions_line:
            allowed_locality_mentions = (
                discord.AllowedMentions(
                    roles=[
                        discord.Object(id=rid) for rid in locality_mentions_role_ids
                    ],
                    users=False,
                    everyone=False,
                )
                if locality_mentions_role_ids
                else safe_mentions
            )
            for chunk in self._paginate_match_lines([locality_mentions_line]):
                await send(chunk, allowed_mentions=allowed_locality_mentions)

        crm_base = settings.espo_base_url.rstrip("/")
        lines, resume_options = self._build_match_candidate_lines(
            candidates=candidates,
            crm_base=crm_base,
        )
        for msg in self._paginate_match_lines(lines):
            await send(msg, allowed_mentions=safe_mentions)
        if resume_options:
            await send(
                "Resume download:",
                view=MatchResumeSelectView(resume_options),
            )

    async def _run_auto_match_candidates_for_thread(
        self,
        *,
        thread: discord.Thread,
        trigger: Literal["thread_create", "message_create"],
    ) -> None:
        """Best-effort automatic matching for a newly created job thread."""
        # Thread creation and first-message events can both fire for the same post.
        # We mark first so only one execution path publishes match output.
        if not await self._mark_thread_auto_matched(thread.id):
            return

        guild = thread.guild

        post = await self._read_thread_post(thread)
        if post is None:
            await self._unmark_thread_auto_matched(thread.id)
            await thread.send(
                "⚠️ Could not read thread opening message. "
                "Run `/match-candidates` manually after fixing permissions."
            )
            return

        posting, _posting_metadata = await self._build_match_candidates_posting(
            post.starter
        )
        if not posting.strip():
            await self._unmark_thread_auto_matched(thread.id)
            await thread.send(
                "⚠️ Could not extract a job description from this forum post. "
                "Run `/match-candidates` manually after adding details, attachments, or links."
            )
            return

        if post.tags:
            tag_names = ", ".join(post.tags)
            posting = f"Thread tags: {tag_names}\n\n{posting}"

        try:
            requirements = await asyncio.to_thread(
                extract_job_requirements,
                posting,
                api_key=settings.openai_api_key,
                base_url=settings.openai_base_url or None,
                model=settings.openai_model,
                webhook_url=settings.discord_logs_webhook_url,
            )
        except Exception as exc:
            await self._unmark_thread_auto_matched(thread.id)
            logger.warning(
                "Auto match failed while extracting requirements "
                "(server=%s thread=%s trigger=%s): %s",
                thread.guild.id if thread.guild else "unknown",
                thread.id,
                trigger,
                exc,
            )
            await thread.send(
                "⚠️ Failed to analyze this posting automatically. "
                "Run `/match-candidates` manually in this thread."
            )
            return

        if not requirements.required_skills:
            await self._unmark_thread_auto_matched(thread.id)
            await thread.send(
                "⚠️ No required skills could be extracted automatically. "
                "Run `/match-candidates` manually after updating the posting."
            )
            return

        try:
            candidates = await asyncio.to_thread(
                search_candidates,
                settings,
                requirements,
                guild_id=str(thread.guild.id) if thread.guild else None,
                limit=10,
            )
        except Exception as exc:
            await self._unmark_thread_auto_matched(thread.id)
            logger.warning(
                "Auto candidate search failed (guild=%s thread=%s trigger=%s): %s",
                thread.guild.id if thread.guild else "unknown",
                thread.id,
                trigger,
                exc,
            )
            await thread.send(
                "⚠️ Automatic candidate search failed. "
                "Run `/match-candidates` manually in this thread."
            )
            return

        await self._publish_match_results(
            send=thread.send,
            requirements=requirements,
            candidates=candidates,
            guild=guild,
        )

    async def _bulk_sync_guild_roles(
        self, guild: discord.Guild
    ) -> tuple[int, int, int]:
        """Sync discord_roles for all non-bot guild members.

        Returns (updated, skipped, failed). Per-member failures are logged and
        skipped so one bad record never aborts the full run.
        Roles in DISCORD_ROLES_EXCLUDE_FROM_SYNC (Bots, FixTweet, @everyone)
        are excluded from the stored list.
        """
        updated = 0
        skipped = 0
        failed = 0
        for member in guild.members:
            if member.bot:
                continue
            role_names = [
                r.name
                for r in member.roles
                if r.name not in DISCORD_ROLES_EXCLUDE_FROM_SYNC
            ]
            try:
                await asyncio.to_thread(
                    upsert_discord_member,
                    settings,
                    discord_user_id=str(member.id),
                    guild_id=str(guild.id),
                    discord_username=member.name,
                    display_name=member.display_name,
                    roles=role_names,
                )
                did_update = await asyncio.to_thread(
                    update_person_discord_roles,
                    settings,
                    str(member.id),
                    role_names,
                )
            except Exception as exc:
                failed += 1
                logger.warning(
                    "bulk role sync: failed for user_id=%s: %s", member.id, exc
                )
                continue
            if did_update:
                updated += 1
            else:
                skipped += 1
        return updated, skipped, failed

    def _get_role_id_cache(self) -> dict[int, dict[str, int]]:
        cache = getattr(self, "_role_id_cache", None)
        if cache is None:
            cache = {}
            setattr(self, "_role_id_cache", cache)
        return cache

    def _refresh_role_id_cache(self, guild: discord.Guild) -> None:
        excluded_names = {name.casefold() for name in DISCORD_ROLES_EXCLUDE_FROM_SYNC}
        role_id_map: dict[str, int] = {}
        sorted_roles = sorted(
            guild.roles,
            key=lambda role: (-getattr(role, "position", 0), role.id),
        )
        for role in sorted_roles:
            normalized_name = role.name.casefold()
            if normalized_name in excluded_names or normalized_name in role_id_map:
                continue
            role_id_map[normalized_name] = role.id
        self._get_role_id_cache()[guild.id] = role_id_map

    @commands.Cog.listener()
    async def on_guild_role_create(self, role: discord.Role) -> None:
        self._refresh_role_id_cache(role.guild)

    @commands.Cog.listener()
    async def on_guild_role_delete(self, role: discord.Role) -> None:
        self._refresh_role_id_cache(role.guild)

    @commands.Cog.listener()
    async def on_guild_role_update(
        self, before: discord.Role, after: discord.Role
    ) -> None:
        self._refresh_role_id_cache(after.guild)

    @commands.Cog.listener()
    async def on_guild_remove(self, guild: discord.Guild) -> None:
        self._get_role_id_cache().pop(guild.id, None)

    @commands.Cog.listener()
    async def on_ready(self) -> None:
        """Warm caches and sync role-backed candidate data on startup."""
        async with self._startup_sync_lock:
            if self._startup_sync_done:
                return

            for guild in self.bot.guilds:
                self._refresh_role_id_cache(guild)
                try:
                    channel_ids = await self._refresh_jobs_channel_cache(guild.id)
                    logger.info(
                        "Loaded %d registered jobs channel(s) for guild=%s",
                        len(channel_ids),
                        guild.name,
                    )
                except Exception as exc:
                    logger.warning(
                        "Failed loading jobs channel registrations for guild %s: %s",
                        guild.name,
                        exc,
                    )

                try:
                    updated, skipped, failed = await self._bulk_sync_guild_roles(guild)
                    logger.info(
                        "Startup discord role sync: guild=%s updated=%d skipped=%d failed=%d",
                        guild.name,
                        updated,
                        skipped,
                        failed,
                    )
                except Exception as exc:
                    logger.warning(
                        "Startup discord role sync failed for guild %s: %s",
                        guild.name,
                        exc,
                    )

            self._startup_sync_done = True

    @commands.Cog.listener()
    async def on_thread_create(self, thread: discord.Thread) -> None:
        """Auto-run matching for new forum posts in registered forum channels."""
        guild = thread.guild
        parent = thread.parent
        if guild is None or not isinstance(parent, discord.ForumChannel):
            return
        if not await self._refresh_jobs_channel_cache_if_missing(guild.id):
            return

        if not self._is_jobs_channel_registered(guild.id, parent.id):
            return

        if parent.permissions_for(guild.default_role).view_channel:
            logger.info(
                "Skipping auto-match for publicly visible forum channel %s (%s) in guild %s",
                parent.id,
                parent.name,
                guild.name,
            )
            return

        owner = guild.get_member(thread.owner_id) if thread.owner_id else None
        if owner is None or owner.bot:
            return
        if not check_user_roles_with_hierarchy(owner.roles, ["Member"]):
            return

        await self._run_auto_match_candidates_for_thread(
            thread=thread,
            trigger="thread_create",
        )

    @app_commands.command(
        name="register-jobs-channel",
        description="Register a forum channel for automatic job-post matching.",
    )
    @app_commands.describe(
        channel="Forum channel to watch. Defaults to the current forum or its post thread."
    )
    @require_role("Steering Committee")
    async def register_jobs_channel(
        self,
        interaction: discord.Interaction,
        channel: discord.ForumChannel | None = None,
    ) -> None:
        """Register a forum channel that triggers automatic candidate matching."""
        guild = interaction.guild
        if guild is None:
            await interaction.response.send_message(
                "⚠️ This command must be used inside a server.",
                ephemeral=True,
            )
            return

        target_channel = self._resolve_jobs_channel_target(interaction, channel)
        if target_channel is None:
            await interaction.response.send_message(
                "⚠️ Choose a forum channel or run this inside one of its post threads.",
                ephemeral=True,
            )
            return

        await interaction.response.defer(ephemeral=True)
        try:
            created = await asyncio.to_thread(
                register_job_post_channel,
                settings,
                guild_id=str(guild.id),
                channel_id=str(target_channel.id),
            )
            self._jobs_channels_by_guild.setdefault(guild.id, set()).add(
                target_channel.id
            )
        except Exception as exc:
            logger.warning(
                "Failed to register jobs channel guild=%s channel=%s: %s",
                guild.id,
                target_channel.id,
                exc,
            )
            await interaction.followup.send(
                "❌ Failed to register this channel. Please try again.",
                ephemeral=True,
            )
            return

        try:
            self._audit_command(
                interaction=interaction,
                action="crm.register_jobs_channel",
                result="success",
                metadata={
                    "guild_id": str(guild.id),
                    "channel_id": str(target_channel.id),
                    "channel_name": target_channel.name,
                    "created": created,
                },
            )
        except Exception as exc:
            logger.warning("Audit write failed for crm.register_jobs_channel: %s", exc)

        if created:
            await interaction.followup.send(
                f"✅ Registered <#{target_channel.id}> for automatic job matching.",
                ephemeral=True,
            )
        else:
            await interaction.followup.send(
                f"ℹ️ <#{target_channel.id}> is already registered.",
                ephemeral=True,
            )

    @app_commands.command(
        name="unregister-jobs-channel",
        description="Stop automatic job-post matching for a forum channel.",
    )
    @app_commands.describe(
        channel="Forum channel to stop watching. Defaults to the current forum or its post thread."
    )
    @require_role("Steering Committee")
    async def unregister_jobs_channel(
        self,
        interaction: discord.Interaction,
        channel: discord.ForumChannel | None = None,
    ) -> None:
        """Unregister a forum channel from automatic candidate matching."""
        guild = interaction.guild
        if guild is None:
            await interaction.response.send_message(
                "⚠️ This command must be used inside a server.",
                ephemeral=True,
            )
            return

        target_channel = self._resolve_jobs_channel_target(interaction, channel)
        if target_channel is None:
            await interaction.response.send_message(
                "⚠️ Choose a forum channel or run this inside one of its post threads.",
                ephemeral=True,
            )
            return

        await interaction.response.defer(ephemeral=True)
        try:
            removed = await asyncio.to_thread(
                unregister_job_post_channel,
                settings,
                guild_id=str(guild.id),
                channel_id=str(target_channel.id),
            )
            self._jobs_channels_by_guild.setdefault(guild.id, set()).discard(
                target_channel.id
            )
        except Exception as exc:
            logger.warning(
                "Failed to unregister jobs channel guild=%s channel=%s: %s",
                guild.id,
                target_channel.id,
                exc,
            )
            await interaction.followup.send(
                "❌ Failed to unregister this channel. Please try again.",
                ephemeral=True,
            )
            return

        try:
            self._audit_command(
                interaction=interaction,
                action="crm.unregister_jobs_channel",
                result="success",
                metadata={
                    "guild_id": str(guild.id),
                    "channel_id": str(target_channel.id),
                    "channel_name": target_channel.name,
                    "removed": removed,
                },
            )
        except Exception as exc:
            logger.warning(
                "Audit write failed for crm.unregister_jobs_channel: %s",
                exc,
            )

        if removed:
            await interaction.followup.send(
                f"✅ Unregistered <#{target_channel.id}> from automatic job matching.",
                ephemeral=True,
            )
        else:
            await interaction.followup.send(
                f"ℹ️ <#{target_channel.id}> was not registered.",
                ephemeral=True,
            )

    @app_commands.command(
        name="match-candidates",
        description="Reads this thread's opening message, attachments, and JD links, then returns ranked matching candidates.",
    )
    @require_role("Member")
    async def match_candidates(
        self,
        interaction: discord.Interaction,
    ) -> None:
        """Parse the thread's starter message and find matching candidates ranked by fit.

        Must be invoked inside a thread. The starter message is used as the job posting text.
        The response is posted publicly in the thread.
        """
        if not isinstance(interaction.channel, discord.Thread) or not isinstance(
            interaction.channel.parent, discord.ForumChannel
        ):
            self._audit_command_safe(
                interaction=interaction,
                action="crm.match_candidates",
                result="error",
                metadata={"stage": "not_thread"},
            )
            await interaction.response.send_message(
                "⚠️ This command must be used inside a forum post thread.",
                ephemeral=True,
            )
            return

        thread: discord.Thread = interaction.channel
        starter = thread.starter_message
        fetch_error = None
        if starter is None:
            try:
                starter = await thread.fetch_message(thread.id)
            except Exception as exc:
                fetch_error = exc
                starter = None

        if starter is None:
            metadata = {"stage": "starter_message_unavailable"}
            if fetch_error is not None:
                error_text = (
                    str(fetch_error).replace("\r", " ").replace("\n", " ").strip()
                )
                if len(error_text) > 300:
                    error_text = f"{error_text[:297]}..."
                metadata["error"] = error_text
            self._audit_command_safe(
                interaction=interaction,
                action="crm.match_candidates",
                result="error",
                metadata=metadata,
            )
            await interaction.response.send_message(
                "⚠️ Could not read the thread's opening message. "
                "Make sure the thread was created from a job posting message.",
                ephemeral=True,
            )
            return

        await interaction.response.defer(ephemeral=False)

        posting, posting_metadata = await self._build_match_candidates_posting(starter)
        if not posting.strip():
            self._audit_command_safe(
                interaction=interaction,
                action="crm.match_candidates",
                result="error",
                metadata={
                    "stage": "starter_message_empty_after_scan",
                    **posting_metadata,
                },
            )
            await interaction.followup.send(
                "⚠️ Could not extract a job description from the thread opener, "
                "its attachments, or linked pages.",
                ephemeral=True,
            )
            return

        if thread.applied_tags:
            tag_names = ", ".join(t.name for t in thread.applied_tags)
            posting = f"Thread tags: {tag_names}\n\n{posting}"

        try:
            requirements = await asyncio.to_thread(
                extract_job_requirements,
                posting,
                api_key=settings.openai_api_key,
                base_url=settings.openai_base_url or None,
                model=settings.openai_model,
                webhook_url=settings.discord_logs_webhook_url,
            )
        except Exception as exc:
            self._audit_command_safe(
                interaction=interaction,
                action="crm.match_candidates",
                result="error",
                metadata={"stage": "extract_requirements", "error": str(exc)},
            )
            await interaction.followup.send(
                f"❌ Failed to analyze the job posting: {exc}",
                ephemeral=True,
            )
            return

        if not requirements.required_skills:
            self._audit_command_safe(
                interaction=interaction,
                action="crm.match_candidates",
                result="error",
                metadata={"stage": "no_required_skills_extracted"},
            )
            await interaction.followup.send(
                "⚠️ No required skills could be extracted from this posting. "
                "Please include explicit skill requirements and try again.",
                ephemeral=True,
            )
            return

        try:
            candidates = await asyncio.to_thread(
                search_candidates,
                settings,
                requirements,
                guild_id=str(interaction.guild.id) if interaction.guild else None,
                limit=20,
                min_match_score=8.0,
            )
        except Exception as exc:
            logger.error("Candidate search failed: %s", exc)
            self._audit_command_safe(
                interaction=interaction,
                action="crm.match_candidates",
                result="error",
                metadata={"stage": "search_candidates", "error": str(exc)},
            )
            await interaction.followup.send(
                "❌ Candidate search failed. Please try again later.",
                ephemeral=True,
            )
            return

        await self._publish_match_results(
            send=interaction.followup.send,
            requirements=requirements,
            candidates=candidates,
            guild=interaction.guild,
        )

        self._audit_command_safe(
            interaction=interaction,
            action="crm.match_candidates",
            result="success",
            metadata={
                "title": requirements.title,
                "required_skills_count": len(requirements.required_skills),
                "preferred_skills_count": len(requirements.preferred_skills),
                "discord_role_types": requirements.discord_role_types,
                "candidates_returned": len(candidates),
                **posting_metadata,
            },
        )

    @app_commands.command(
        name="sync-discord-roles",
        description="Re-sync all server members' Discord roles into the candidate database.",
    )
    @require_role("Steering Committee")
    async def sync_discord_roles(
        self,
        interaction: discord.Interaction,
    ) -> None:
        """Manually trigger a full guild role sync (also runs automatically on startup)."""
        guild = interaction.guild
        if guild is None:
            await interaction.response.send_message(
                "⚠️ This command must be used inside a server.",
                ephemeral=True,
            )
            return

        await interaction.response.defer(ephemeral=True)

        updated, skipped, failed = await self._bulk_sync_guild_roles(guild)

        self._audit_command_safe(
            interaction=interaction,
            action="crm.sync_discord_roles",
            result="success",
            metadata={
                "updated": updated,
                "skipped_no_db_match": skipped,
                "failed": failed,
                "total_members_scanned": updated + skipped + failed,
            },
        )

        await interaction.followup.send(
            f"✅ Discord role sync complete.\n"
            f"Updated: **{updated}** · No DB match (skipped): **{skipped}** · Failed: **{failed}**",
            ephemeral=True,
        )

    @commands.Cog.listener()
    async def on_member_update(
        self,
        before: discord.Member,
        after: discord.Member,
    ) -> None:
        """Automatically sync discord roles/names on member updates."""
        if after.guild is None or after.bot:
            return
        roles_changed = before.roles != after.roles
        name_changed = (
            before.display_name != after.display_name or before.name != after.name
        )
        if not roles_changed and not name_changed:
            return

        role_names = [
            r.name for r in after.roles if r.name not in DISCORD_ROLES_EXCLUDE_FROM_SYNC
        ]

        try:
            await asyncio.to_thread(
                upsert_discord_member,
                settings,
                discord_user_id=str(after.id),
                guild_id=str(after.guild.id),
                discord_username=after.name,
                display_name=after.display_name,
                roles=role_names,
            )
            if roles_changed:
                await asyncio.to_thread(
                    update_person_discord_roles,
                    settings,
                    str(after.id),
                    role_names,
                )
        except Exception as exc:
            logger.warning(
                "on_member_update: failed to sync roles for user %s: %s",
                after.id,
                exc,
            )


async def setup(bot: commands.Bot) -> None:
    """Add the jobs cog to the bot."""
    await bot.add_cog(JobsCog(bot))
