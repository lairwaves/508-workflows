"""
CRM integration cog for the 508.dev Discord bot.

This cog provides commands for interacting with EspoCRM through Discord slash commands.
It allows team members to quickly access CRM data without leaving Discord.
"""

import asyncio
import ast
import io
import json
import logging
from datetime import date, datetime, timezone
import re
from typing import Any, Literal

import aiohttp
import discord
from discord import app_commands
from discord.ext import commands

from five08.discord_bot.config import settings
from five08.clients import espo
from five08.skills import normalize_skill_list
from five08.resume_extractor import (
    ResumeExtractedProfile,
    ResumeProfileExtractor,
    is_reserved_resume_name_token,
    normalize_resume_name_token,
)
from five08.discord_bot.utils.audit import DiscordAuditLogger
from five08.discord_bot.utils.role_decorators import (
    require_role,
    check_user_roles_with_hierarchy,
)
from five08.job_match import extract_job_requirements, DISCORD_ROLES_EXCLUDE_FROM_SYNC
from five08.candidate_search import search_candidates
from five08.audit import update_person_discord_roles, upsert_discord_member

logger = logging.getLogger(__name__)

ID_VERIFIED_AT_FIELD = "cIdVerifiedAt"
ID_VERIFIED_BY_FIELD = "cIdVerifiedBy"
ID_VERIFIED_TYPE_FIELD = "cVerifiedIdType"
ONBOARDING_STATUS_FIELD_CANDIDATES = (
    "cOnboardingState",
    "cOnboardingStatus",
    "cOnboarding",
)
ONBOARDER_FIELD_CANDIDATES = (
    "cOnboarder",
    "cOnboardingCoordinator",
)
EXCLUDED_ONBOARDING_STATES = frozenset({"onboarded", "waitlist", "rejected"})
ONBOARDING_QUEUE_MAX_SIZE = 200
ONBOARDING_QUEUE_PAGE_SIZE = 1
EspoAPI = espo.EspoAPI
EspoAPIError = espo.EspoAPIError


def _configured_linkedin_field_from_settings() -> str:
    value = getattr(settings, "crm_linkedin_field", None)
    if isinstance(value, str):
        value = value.strip()
        if value:
            return value
    return "cLinkedIn"


class ResumeButtonView(discord.ui.View):
    """View containing resume download buttons for contact search results."""

    def __init__(self) -> None:
        super().__init__(timeout=300)  # 5 minute timeout

    def add_resume_button(self, contact_name: str, resume_id: str) -> None:
        """Add a resume download button for a contact."""
        if len(self.children) >= 5:  # Discord limit of 5 buttons per row
            return

        button = ResumeDownloadButton(contact_name, resume_id)
        self.add_item(button)


class ResumeDownloadButton(discord.ui.Button[discord.ui.View]):
    """Button for downloading a specific contact's resume."""

    def __init__(self, contact_name: str, resume_id: str) -> None:
        self.contact_name = contact_name
        self.resume_id = resume_id

        # Truncate long names for button label
        label = f"📄 Resume: {contact_name}"
        if len(label) > 80:  # Discord button label limit
            # Account for "📄 Resume: " (11 chars) + "..." (3 chars) = 14 chars
            max_name_length = 80 - 14
            label = f"📄 Resume: {contact_name[:max_name_length]}..."

        super().__init__(
            label=label,
            style=discord.ButtonStyle.secondary,
            custom_id=f"resume_{resume_id}",
        )

    async def callback(self, interaction: discord.Interaction) -> None:
        """Handle resume download button click."""
        try:
            # Get the CRM cog to access the API
            cog = interaction.client.get_cog("CRMCog")  # type: ignore[attr-defined]
            if not cog:
                await interaction.response.send_message(
                    "❌ CRM functionality not available.", ephemeral=True
                )
                return

            # Check if user has Member role
            if not cog._check_member_role(interaction):
                cog._audit_command(
                    interaction=interaction,
                    action="crm.resume_download_button",
                    result="denied",
                    metadata={"reason": "missing_member_role"},
                    resource_type="discord_ui_action",
                    resource_id=self.resume_id,
                )
                await interaction.response.send_message(
                    "❌ You must have the Member role to download resumes.",
                    ephemeral=True,
                )
                return

            await interaction.response.defer(ephemeral=True)

            # Use shared download method
            download_ok = await cog._download_and_send_resume(
                interaction, self.contact_name, self.resume_id
            )
            cog._audit_command(
                interaction=interaction,
                action="crm.resume_download_button",
                result="success" if download_ok else "error",
                metadata={"contact_name": self.contact_name},
                resource_type="crm_contact",
                resource_id=self.resume_id,
            )

        except Exception as e:
            logger.error(f"Unexpected error in resume button callback: {e}")
            if "cog" in locals() and cog:
                cog._audit_command(
                    interaction=interaction,
                    action="crm.resume_download_button",
                    result="error",
                    metadata={"error": str(e)},
                    resource_type="discord_ui_action",
                    resource_id=self.resume_id,
                )
            await interaction.followup.send(
                "❌ An unexpected error occurred while downloading the resume."
            )


class MatchResumeSelectView(discord.ui.View):
    """View containing a resume download select for match results."""

    def __init__(self, options: list[tuple[str, str, str]]) -> None:
        super().__init__(timeout=600)  # 10 minute timeout
        self.add_item(MatchResumeSelect(options))


class MatchResumeSelect(discord.ui.Select):
    """Select menu for downloading a resume from match results."""

    def __init__(self, options: list[tuple[str, str, str]]) -> None:
        discord_options: list[discord.SelectOption] = []
        self._resume_lookup: dict[str, str] = {}

        for contact_name, resume_id, resume_name in options[:25]:
            label = contact_name.strip() or "Unknown"
            if len(label) > 100:
                label = label[:97] + "..."
            description = resume_name.strip() or "Resume"
            if len(description) > 100:
                description = description[:97] + "..."
            discord_options.append(
                discord.SelectOption(
                    label=label,
                    value=resume_id,
                    description=description,
                )
            )
            self._resume_lookup[resume_id] = contact_name

        super().__init__(
            placeholder="Download a resume...",
            min_values=1,
            max_values=1,
            options=discord_options,
            custom_id="match_resume_select",
        )

    async def callback(self, interaction: discord.Interaction) -> None:
        try:
            cog = interaction.client.get_cog("CRMCog")  # type: ignore[attr-defined]
            if not cog:
                await interaction.response.send_message(
                    "❌ CRM functionality not available.", ephemeral=True
                )
                return

            resume_id = self.values[0]
            contact_name = self._resume_lookup.get(resume_id, "Unknown")

            await interaction.response.defer(ephemeral=True)

            download_ok = await cog._download_and_send_resume(
                interaction, contact_name, resume_id
            )
            try:
                cog._audit_command(
                    interaction=interaction,
                    action="crm.match_candidates_resume_select",
                    result="success" if download_ok else "error",
                    metadata={"contact_name": contact_name},
                    resource_type="crm_contact",
                    resource_id=resume_id,
                )
            except Exception as audit_exc:
                logger.error("Audit write failed in match resume select: %s", audit_exc)
        except Exception as exc:
            logger.error("Unexpected error in match resume select: %s", exc)
            if "cog" in locals() and cog:
                try:
                    cog._audit_command(
                        interaction=interaction,
                        action="crm.match_candidates_resume_select",
                        result="error",
                        metadata={"error": str(exc)},
                        resource_type="discord_ui_action",
                        resource_id=self.values[0] if self.values else None,
                    )
                except Exception as audit_exc:
                    logger.error(
                        "Audit write failed in match resume select: %s", audit_exc
                    )
            await interaction.followup.send(
                "❌ An unexpected error occurred while downloading the resume."
            )


class ContactSelectionView(discord.ui.View):
    """View containing contact selection buttons for Discord linking."""

    def __init__(self, user: discord.Member, search_term: str) -> None:
        super().__init__(timeout=300)  # 5 minute timeout
        self.user = user
        self.search_term = search_term

    def add_contact_button(self, contact: dict[str, Any]) -> None:
        """Add a contact selection button."""
        if len(self.children) >= 5:  # Discord limit of 5 buttons per row
            return

        button = ContactSelectionButton(contact, self.user)
        self.add_item(button)


class ContactSelectionButton(discord.ui.Button[ContactSelectionView]):
    """Button for selecting a contact to link to Discord user."""

    def __init__(self, contact: dict[str, Any], user: discord.Member) -> None:
        # Create button label from contact name (truncate if too long)
        contact_name = contact.get("name", "Unknown")
        label = contact_name[:80] if len(contact_name) > 80 else contact_name

        super().__init__(style=discord.ButtonStyle.primary, label=label, emoji="🔗")
        self.contact = contact
        self.user = user

    async def callback(self, interaction: discord.Interaction) -> None:
        """Handle contact selection and perform the Discord linking."""
        try:
            # Check if user has required role
            if not hasattr(
                interaction.user, "roles"
            ) or not check_user_roles_with_hierarchy(
                interaction.user.roles, ["Steering Committee"]
            ):
                await interaction.response.send_message(
                    "❌ You must have Steering Committee role or higher to use this command.",
                    ephemeral=True,
                )
                return

            await interaction.response.defer(ephemeral=True)

            # Get the CRM cog to perform the linking
            if not self.view:
                await interaction.followup.send("❌ View not found.")
                return

            from discord.ext import commands

            bot = interaction.client
            assert isinstance(bot, commands.Bot)
            cog = bot.get_cog("CRMCog")
            if not cog or not isinstance(cog, CRMCog):
                await interaction.followup.send("❌ CRM cog not found.")
                return

            # Perform the Discord linking
            success = await cog._perform_discord_linking(
                interaction, self.user, self.contact
            )

            if success and self.view:
                # Disable all buttons in the view
                for item in self.view.children:
                    if isinstance(item, discord.ui.Button):
                        item.disabled = True

                # Update the original message to show selection was made
                embed = discord.Embed(
                    title="✅ Contact Selected",
                    description=f"Selected **{self.contact.get('name', 'Unknown')}** for linking.",
                    color=0x00FF00,
                )

                # Edit the original message with disabled buttons
                if interaction.message:
                    try:
                        await interaction.message.edit(embed=embed, view=self.view)
                    except discord.NotFound:
                        # Message was deleted or not found, ignore this error
                        logger.debug(
                            "Original message not found when trying to update button view"
                        )
                    except discord.HTTPException as e:
                        # Other Discord API errors
                        logger.warning(f"Failed to update original message: {e}")

        except Exception as e:
            logger.error(f"Error in contact selection callback: {e}")
            await interaction.followup.send(
                "❌ An error occurred while linking the contact."
            )


class MarkIdVerifiedSelectionButton(discord.ui.Button["MarkIdVerifiedSelectionView"]):
    """Button for selecting a contact to mark ID verification on."""

    def __init__(
        self,
        contact: dict[str, Any],
        verified_by: str,
        verified_at: str,
        id_type: str | None,
        requester_id: int,
    ) -> None:
        contact_name = contact.get("name", "Unknown")
        label = contact_name[:80] if len(contact_name) > 80 else contact_name
        super().__init__(style=discord.ButtonStyle.success, label=label, emoji="✅")
        self.contact = contact
        self.verified_by = verified_by
        self.verified_at = verified_at
        self.id_type = id_type
        self.requester_id = requester_id

    async def callback(self, interaction: discord.Interaction) -> None:
        """Handle contact selection and perform the ID verification."""
        try:
            if not self.view:
                await interaction.response.send_message("❌ View not found.")
                return
            if interaction.user.id != self.requester_id:
                await interaction.response.send_message(
                    "❌ Only the command requester can confirm this action.",
                    ephemeral=True,
                )
                return

            await interaction.response.defer(ephemeral=True)
            await self.view.crm_cog._mark_id_verified_for_contact(
                interaction=interaction,
                contact=self.contact,
                verified_by=self.verified_by,
                verified_at=self.verified_at,
                id_type=self.id_type,
            )

            for item in self.view.children:
                if isinstance(item, discord.ui.Button):
                    item.disabled = True

            if interaction.message:
                try:
                    await interaction.message.edit(view=self.view)
                except discord.NotFound:
                    pass
                except discord.HTTPException as exc:
                    logger.warning(
                        f"Failed to update ID verification selection view: {exc}"
                    )
        except Exception as exc:
            logger.error(f"Error in ID verified selection callback: {exc}")
            await interaction.followup.send(
                "❌ An error occurred while marking ID verification."
            )


class MarkIdVerifiedSelectionView(discord.ui.View):
    """View containing contact selection buttons for ID verification."""

    def __init__(
        self,
        crm_cog: "CRMCog",
        requester_id: int,
        verified_by: str,
        verified_at: str,
        id_type: str | None,
    ) -> None:
        super().__init__(timeout=300)  # 5 minute timeout
        self.crm_cog = crm_cog
        self.requester_id = requester_id
        self.verified_by = verified_by
        self.verified_at = verified_at
        self.id_type = id_type

    def add_contact_button(
        self,
        contact: dict[str, Any],
    ) -> None:
        """Add a contact selection button."""
        if len(self.children) >= 5:
            return
        button = MarkIdVerifiedSelectionButton(
            contact=contact,
            verified_by=self.verified_by,
            verified_at=self.verified_at,
            id_type=self.id_type,
            requester_id=self.requester_id,
        )
        self.add_item(button)


class MarkIdVerifiedOverwriteConfirmationView(discord.ui.View):
    """View for confirming overwrite of existing ID verification values."""

    def __init__(
        self,
        crm_cog: "CRMCog",
        interaction: discord.Interaction,
        contact: dict[str, Any],
        verified_by: str,
        verified_at: str,
        id_type: str | None,
    ) -> None:
        super().__init__(timeout=300)  # 5 minute timeout
        self.crm_cog = crm_cog
        self.original_interaction = interaction
        self.contact = contact
        self.verified_by = verified_by
        self.verified_at = verified_at
        self.id_type = id_type
        self.requester_id = interaction.user.id

    async def interaction_check(self, interaction: discord.Interaction) -> bool:
        """Allow only the original requester to confirm/cancel."""
        if interaction.user.id != self.requester_id:
            await interaction.response.send_message(
                "❌ Only the command requester can confirm this action.",
                ephemeral=True,
            )
            return False
        return True

    @discord.ui.button(label="Overwrite", style=discord.ButtonStyle.danger, emoji="⚠️")
    async def confirm_overwrite(
        self,
        interaction: discord.Interaction,
        button: discord.ui.Button["MarkIdVerifiedOverwriteConfirmationView"],
    ) -> None:
        """Overwrite existing verification metadata and continue."""
        await interaction.response.defer(ephemeral=True)
        await self.crm_cog._mark_id_verified_for_contact(
            interaction=interaction,
            contact=self.contact,
            verified_by=self.verified_by,
            verified_at=self.verified_at,
            id_type=self.id_type,
            allow_overwrite=True,
        )
        for item in self.children:
            if isinstance(item, discord.ui.Button):
                item.disabled = True

        if interaction.message:
            try:
                await interaction.message.edit(view=self)
            except discord.NotFound:
                pass

    @discord.ui.button(label="Cancel", style=discord.ButtonStyle.secondary, emoji="❌")
    async def cancel_overwrite(
        self,
        interaction: discord.Interaction,
        button: discord.ui.Button["MarkIdVerifiedOverwriteConfirmationView"],
    ) -> None:
        """Cancel overwrite and leave contact unchanged."""
        await interaction.response.send_message(
            "✅ ID verification overwrite cancelled. No changes were made.",
            ephemeral=True,
        )
        for item in self.children:
            if isinstance(item, discord.ui.Button):
                item.disabled = True

        if interaction.message:
            try:
                await interaction.message.edit(view=self)
            except discord.NotFound:
                pass


class ResumeConfirmationView(discord.ui.View):
    """View for confirming resume upload when duplicate is detected."""

    def __init__(
        self,
        crm_cog: "CRMCog",
        interaction: discord.Interaction,
        file: discord.Attachment,
        contact_id: str,
        contact_name: str,
        existing_resume_id: str,
        overwrite: bool = False,
    ) -> None:
        super().__init__(timeout=300)  # 5 minute timeout
        self.crm_cog = crm_cog
        self.original_interaction = interaction
        self.file = file
        self.contact_id = contact_id
        self.contact_name = contact_name
        self.existing_resume_id = existing_resume_id
        self.overwrite = overwrite

    @discord.ui.button(
        label="Yes, Upload Anyway", style=discord.ButtonStyle.primary, emoji="📄"
    )
    async def confirm_upload(
        self,
        interaction: discord.Interaction,
        button: discord.ui.Button["ResumeConfirmationView"],
    ) -> None:
        """Proceed with the upload despite duplicate."""
        await interaction.response.defer()

        try:
            # Download file content from Discord
            file_content = await self.file.read()

            # Upload file to EspoCRM
            attachment = self.crm_cog.espo_api.upload_file(
                file_content=file_content,
                filename=self.file.filename,
                related_type="Contact",
                related_id=self.contact_id,
                field="resume",
            )

            attachment_id = attachment.get("id")
            if not attachment_id:
                await interaction.followup.send("❌ Failed to upload file to CRM.")
                return

            # Update contact's resume field (use original overwrite setting)
            if await self.crm_cog._update_contact_resume(
                self.contact_id, attachment_id, self.overwrite
            ):
                # Create success embed
                embed = discord.Embed(
                    title="✅ Resume Uploaded Successfully",
                    description="Resume has been uploaded and linked to the contact.",
                    color=0x00FF00,
                )
                embed.add_field(name="👤 Contact", value=self.contact_name, inline=True)
                embed.add_field(name="📄 File", value=self.file.filename, inline=True)
                embed.add_field(
                    name="📁 Size", value=f"{self.file.size / 1024:.1f} KB", inline=True
                )

                # Add CRM link
                profile_url = f"{self.crm_cog.base_url}/#Contact/view/{self.contact_id}"
                embed.add_field(
                    name="🔗 CRM Profile",
                    value=f"[View in CRM]({profile_url})",
                    inline=False,
                )

                await interaction.followup.send(embed=embed)

                logger.info(
                    f"Resume uploaded for {self.contact_name} (ID: {self.contact_id}) "
                    f"by {self.original_interaction.user.name}: {self.file.filename}"
                )
            else:
                await interaction.followup.send(
                    "⚠️ File uploaded but failed to link to contact. Please check CRM manually."
                )

        except Exception as e:
            logger.error(f"Error during confirmed resume upload: {e}")
            await interaction.followup.send(
                "❌ An error occurred while uploading the resume."
            )

        # Disable all buttons
        for item in self.children:
            if isinstance(item, discord.ui.Button):
                item.disabled = True

        # Update the original message
        try:
            if interaction.message:
                await interaction.message.edit(view=self)
        except discord.NotFound:
            pass

    @discord.ui.button(
        label="No, Cancel", style=discord.ButtonStyle.secondary, emoji="❌"
    )
    async def cancel_upload(
        self,
        interaction: discord.Interaction,
        button: discord.ui.Button["ResumeConfirmationView"],
    ) -> None:
        """Cancel the upload."""
        await interaction.response.send_message(
            "📄 Resume upload cancelled. No changes were made.", ephemeral=True
        )

        # Disable all buttons
        for item in self.children:
            if isinstance(item, discord.ui.Button):
                item.disabled = True

        # Update the original message
        try:
            if interaction.message:
                await interaction.message.edit(view=self)
        except discord.NotFound:
            pass


class ResumeUpdateConfirmationView(discord.ui.View):
    """Confirm extracted profile updates before writing to CRM."""

    _EMBED_FIELD_LIMIT = 1024
    _APPLIED_VALUE_LIMIT = 150
    _APPLIED_FIELD_TOTAL_LIMIT = 900

    _FIELD_LABELS: dict[str, str] = {
        "emailAddressData": "Email Addresses",
        "cGitHubUsername": "GitHub",
        "phoneNumber": "Phone",
        "skills": "Skills",
        "cSkillAttrs": "Skill Strengths",
        "cWebsiteLink": "Website",
        "cSocialLinks": "Social Links",
        "cSeniority": "Seniority",
        "cDiscordUserID": "Discord User ID",
        "cDiscordUsername": "Discord Username",
    }

    def __init__(
        self,
        *,
        crm_cog: "CRMCog",
        requester_id: int,
        contact_id: str,
        contact_name: str,
        proposed_updates: dict[str, Any],
        link_discord: dict[str, str] | None = None,
    ) -> None:
        super().__init__(timeout=300)
        self.crm_cog = crm_cog
        self.requester_id = requester_id
        self.contact_id = contact_id
        self.contact_name = contact_name
        self.proposed_updates = proposed_updates
        self.link_discord = link_discord

    @classmethod
    def _field_label(cls, field: str) -> str:
        linkedin_field = _configured_linkedin_field_from_settings()
        if field == linkedin_field:
            return "LinkedIn"
        return cls._FIELD_LABELS.get(field, field)

    @staticmethod
    def _truncate_embed_field(value: str, limit: int = 1024) -> str:
        if len(value) <= limit:
            return value
        if limit <= 3:
            return value[:limit]
        return value[: limit - 3] + "..."

    @staticmethod
    def _decode_json_like_mapping(value: Any) -> dict[str, Any] | None:
        candidate = value
        if isinstance(candidate, str):
            raw = candidate.strip()
            if not raw:
                return None
            try:
                candidate = json.loads(raw)
            except Exception:
                try:
                    candidate = ast.literal_eval(raw)
                except Exception:
                    return None
            if isinstance(candidate, str):
                nested = candidate.strip()
                if not nested:
                    return None
                try:
                    candidate = json.loads(nested)
                except Exception:
                    try:
                        candidate = ast.literal_eval(nested)
                    except Exception:
                        return None
        if not isinstance(candidate, dict):
            return None
        return {str(key): item_value for key, item_value in candidate.items()}

    @classmethod
    def _format_field_value(cls, field: str, value: Any) -> str:
        if value is None:
            return "None"

        if field == "cGitHubUsername":
            username = str(value).strip().lstrip("@")
            return f"@{username}" if username else "None"

        if field == "skills":
            if isinstance(value, str):
                return cls._truncate_embed_field(value, cls._APPLIED_VALUE_LIMIT)
            if isinstance(value, (list, tuple, set)):
                items = [str(item).strip() for item in value if str(item).strip()]
                joined = ", ".join(items) if items else "None"
                return cls._truncate_embed_field(joined, cls._APPLIED_VALUE_LIMIT)

        if field == "cSkillAttrs":
            parsed = cls._decode_json_like_mapping(value)
            if parsed:
                formatted: list[str] = []
                for raw_skill, raw_payload in parsed.items():
                    strength_value = (
                        raw_payload.get("strength")
                        if isinstance(raw_payload, dict)
                        else raw_payload
                    )
                    if strength_value is None:
                        strength = 0
                        skill = str(raw_skill).strip()
                        if not skill:
                            continue
                        formatted.append(skill)
                        continue
                    try:
                        strength = int(float(strength_value))
                    except Exception:
                        strength = 0
                    skill = str(raw_skill).strip()
                    if not skill:
                        continue
                    if 1 <= strength <= 5:
                        formatted.append(f"{skill} ({strength})")
                    else:
                        formatted.append(skill)
                joined = ", ".join(formatted) if formatted else "None"
                return cls._truncate_embed_field(joined, cls._APPLIED_VALUE_LIMIT)

        if isinstance(value, (list, tuple, set)):
            items = [str(item).strip() for item in value if str(item).strip()]
            joined = ", ".join(items) if items else "None"
            return cls._truncate_embed_field(joined, cls._APPLIED_VALUE_LIMIT)
        if isinstance(value, dict):
            try:
                encoded = json.dumps(value, sort_keys=True, separators=(",", ":"))
            except Exception:
                encoded = str(value)
            return cls._truncate_embed_field(encoded, cls._APPLIED_VALUE_LIMIT)

        text = str(value).strip()
        return cls._truncate_embed_field(text or "None", cls._APPLIED_VALUE_LIMIT)

    def _build_applied_updates_lines(
        self,
        *,
        updated_fields: list[str],
        updated_values: dict[str, Any],
    ) -> list[str]:
        has_skills_field = "skills" in updated_fields
        has_skill_attrs_field = "cSkillAttrs" in updated_fields
        lines: list[str] = []
        if has_skills_field or has_skill_attrs_field:
            skills_value = updated_values.get(
                "skills", self.proposed_updates.get("skills")
            )
            attrs_value = updated_values.get(
                "cSkillAttrs",
                self.proposed_updates.get("cSkillAttrs"),
            )
            combined_skills = self._format_combined_skills_value(
                skills_value=skills_value,
                attrs_value=attrs_value,
            )
            truncated_combined_skills = self._truncate_embed_field(
                combined_skills, self._APPLIED_VALUE_LIMIT
            )
            lines.append(f"**Skills**: `{truncated_combined_skills}`")

        for field in self._collapse_updated_fields(updated_fields):
            if field == "skills":
                continue
            label = self._field_label(field)
            value = updated_values.get(field, self.proposed_updates.get(field))
            formatted = self._format_field_value(field, value)
            lines.append(f"**{label}**: `{formatted}`")
        return lines

    @classmethod
    def _format_updated_fields_value(cls, labeled_fields: list[str]) -> str:
        if not labeled_fields:
            return "No field changes"

        full = ", ".join(labeled_fields)
        if len(full) <= cls._EMBED_FIELD_LIMIT:
            return full

        kept: list[str] = []
        for index, field in enumerate(labeled_fields):
            kept.append(field)
            remaining = len(labeled_fields) - index - 1
            suffix = f", and {remaining} more" if remaining > 0 else ""
            candidate = ", ".join(kept) + suffix
            if len(candidate) > cls._EMBED_FIELD_LIMIT:
                kept.pop()
                break

        if not kept:
            return cls._truncate_embed_field(full, cls._EMBED_FIELD_LIMIT)

        remaining = len(labeled_fields) - len(kept)
        if remaining > 0:
            candidate = ", ".join(kept) + f", and {remaining} more"
            return cls._truncate_embed_field(candidate, cls._EMBED_FIELD_LIMIT)
        return cls._truncate_embed_field(", ".join(kept), cls._EMBED_FIELD_LIMIT)

    @classmethod
    def _format_applied_updates_value(cls, applied_lines: list[str]) -> str:
        if not applied_lines:
            return "No applied updates"

        kept: list[str] = []
        total = 0
        for index, line in enumerate(applied_lines[:8]):
            line_len = len(line) + (1 if kept else 0)
            remaining = len(applied_lines[:8]) - index - 1
            suffix = f"... and {remaining} more" if remaining > 0 else ""
            projected = total + line_len
            if suffix:
                projected += len(suffix) + 1
            if projected > cls._APPLIED_FIELD_TOTAL_LIMIT:
                break
            kept.append(line)
            total += line_len

        if not kept:
            joined = "\n".join(applied_lines[:8])
            return cls._truncate_embed_field(joined, cls._APPLIED_FIELD_TOTAL_LIMIT)

        remaining = len(applied_lines[:8]) - len(kept)
        if remaining > 0:
            kept.append(f"... and {remaining} more")
        joined = "\n".join(kept)
        return cls._truncate_embed_field(joined, cls._APPLIED_FIELD_TOTAL_LIMIT)

    @staticmethod
    def _collapse_updated_fields(updated_fields: list[str]) -> list[str]:
        """Collapse skill fields into a single logical skills entry."""
        collapsed: list[str] = []
        seen: set[str] = set()
        has_skills = "skills" in updated_fields
        has_skill_attrs = "cSkillAttrs" in updated_fields

        for field in updated_fields:
            normalized_field = field
            if field == "cSkillAttrs":
                if has_skills:
                    continue
                if has_skill_attrs:
                    normalized_field = "skills"
            key = normalized_field.casefold()
            if key in seen:
                continue
            seen.add(key)
            collapsed.append(normalized_field)
        return collapsed

    @classmethod
    def _format_combined_skills_value(
        cls,
        *,
        skills_value: Any,
        attrs_value: Any,
    ) -> str:
        skills: list[str] = []
        if isinstance(skills_value, str):
            skills = [item.strip() for item in skills_value.split(",") if item.strip()]
        elif isinstance(skills_value, (list, tuple, set)):
            skills = [str(item).strip() for item in skills_value if str(item).strip()]

        parsed_attrs = cls._decode_json_like_mapping(attrs_value) or {}
        strengths: dict[str, int] = {}
        display_by_key: dict[str, str] = {}
        for raw_skill, raw_payload in parsed_attrs.items():
            skill_name = str(raw_skill).strip()
            if not skill_name:
                continue
            key = skill_name.casefold()
            display_by_key.setdefault(key, skill_name)
            strength_value = (
                raw_payload.get("strength")
                if isinstance(raw_payload, dict)
                else raw_payload
            )
            if strength_value is None:
                continue
            try:
                strength = int(float(strength_value))
            except Exception:
                continue
            if 1 <= strength <= 5:
                strengths[key] = strength

        ordered: list[str] = []
        seen_order: set[str] = set()
        for raw_skill in skills:
            key = raw_skill.casefold()
            if key in seen_order:
                continue
            seen_order.add(key)
            ordered.append(raw_skill)
            display_by_key.setdefault(key, raw_skill)
        for key, display_name in display_by_key.items():
            if key in seen_order:
                continue
            seen_order.add(key)
            ordered.append(display_name)

        if not ordered:
            return "None"

        formatted: list[str] = []
        for skill_name in ordered:
            key = skill_name.casefold()
            skill_strength = strengths.get(key)
            if skill_strength is not None:
                formatted.append(f"{skill_name} ({skill_strength})")
            else:
                formatted.append(skill_name)
        return ", ".join(formatted) if formatted else "None"

    async def interaction_check(self, interaction: discord.Interaction) -> bool:
        """Allow only the original requester to confirm/cancel."""
        if interaction.user.id != self.requester_id:
            await interaction.response.send_message(
                "❌ Only the command requester can confirm these updates.",
                ephemeral=True,
            )
            return False
        return True

    @discord.ui.button(label="Confirm Updates", style=discord.ButtonStyle.primary)
    async def confirm_updates(
        self,
        interaction: discord.Interaction,
        button: discord.ui.Button["ResumeUpdateConfirmationView"],
    ) -> None:
        """Apply confirmed updates through the worker."""
        await interaction.response.defer(thinking=True, ephemeral=True)

        def _audit_apply_event(result: str, metadata: dict[str, Any]) -> None:
            try:
                self.crm_cog._audit_command(
                    interaction=interaction,
                    action="crm.upload_resume.apply",
                    result=result,
                    metadata=metadata,
                    resource_type="crm_contact",
                    resource_id=self.contact_id,
                )
            except Exception as exc:
                logger.warning(
                    "Failed to write resume apply audit event for contact_id=%s: %s",
                    self.contact_id,
                    exc,
                )

        try:
            apply_job_id = await self.crm_cog._enqueue_resume_apply_job(
                contact_id=self.contact_id,
                updates=self.proposed_updates,
                link_discord=self.link_discord,
            )
        except Exception as exc:
            logger.error("Failed to enqueue resume apply job: %s", exc)
            _audit_apply_event(
                "error",
                {
                    "contact_id": self.contact_id,
                    "stage": "apply_enqueue",
                    "error": str(exc),
                    "updated_fields": [],
                    "proposed_updates_count": len(self.proposed_updates),
                    "link_member_requested": bool(self.link_discord),
                    "link_discord_applied": None,
                },
            )
            await interaction.followup.send(
                "❌ Failed to enqueue CRM apply job. Please try again.",
                ephemeral=True,
            )
            return

        await interaction.followup.send(
            "🛠️ Applying confirmed updates to CRM...",
            ephemeral=True,
        )
        try:
            apply_result = await self.crm_cog._wait_for_backend_job_result(apply_job_id)
        except Exception as exc:
            logger.error(
                "Worker polling failed for apply_job_id=%s contact_id=%s error=%s",
                apply_job_id,
                self.contact_id,
                exc,
            )
            _audit_apply_event(
                "error",
                {
                    "contact_id": self.contact_id,
                    "stage": "apply_polling_failed",
                    "job_id": apply_job_id,
                    "error": str(exc),
                    "updated_fields": [],
                    "link_discord_applied": None,
                },
            )
            await interaction.followup.send(
                "⚠️ Resume apply polling failed. Please retry or check CRM manually.",
                ephemeral=True,
            )
            return

        if not apply_result:
            _audit_apply_event(
                "error",
                {
                    "contact_id": self.contact_id,
                    "stage": "apply_timeout",
                    "job_id": apply_job_id,
                    "updated_fields": [],
                    "link_discord_applied": None,
                },
            )
            await interaction.followup.send(
                "⚠️ Timed out waiting for apply job. Please check again shortly.",
                ephemeral=True,
            )
            return

        status = str(apply_result.get("status", "unknown"))
        result = apply_result.get("result")
        updated_fields: list[str] = []
        updated_values: dict[str, Any] = {}
        link_discord_applied: bool | None = None
        if isinstance(result, dict):
            raw_fields = result.get("updated_fields")
            if isinstance(raw_fields, list):
                updated_fields = [str(field) for field in raw_fields]
            raw_values = result.get("updated_values")
            if isinstance(raw_values, dict):
                updated_values = {
                    str(field): value for field, value in raw_values.items()
                }
            raw_link_applied = result.get("link_discord_applied")
            if isinstance(raw_link_applied, bool):
                link_discord_applied = raw_link_applied

        if status != "succeeded":
            result_error = str(apply_result.get("last_error", ""))
            result_success = None
            if isinstance(result, dict):
                result_success = result.get("success")

            if result_success is False:
                error_message = str(
                    result.get("error") if isinstance(result, dict) else ""
                )
                if not error_message:
                    error_message = result_error or "Unknown error"
                _audit_apply_event(
                    "error",
                    {
                        "contact_id": self.contact_id,
                        "stage": "apply_failed",
                        "job_id": apply_job_id,
                        "job_status": status,
                        "last_error": result_error,
                        "apply_error": error_message,
                        "updated_fields": updated_fields,
                        "link_discord_applied": link_discord_applied,
                    },
                )
                await interaction.followup.send(
                    f"❌ Apply job failed (status: {status}). Error: {error_message}",
                    ephemeral=True,
                )
                return

            _audit_apply_event(
                "error",
                {
                    "contact_id": self.contact_id,
                    "stage": "apply_failed",
                    "job_id": apply_job_id,
                    "job_status": status,
                    "last_error": str(apply_result.get("last_error", "")),
                    "updated_fields": updated_fields,
                    "link_discord_applied": link_discord_applied,
                },
            )
            await interaction.followup.send(
                f"❌ Apply job failed (status: {status}). "
                f"Error: {apply_result.get('last_error') or 'Unknown error'}",
                ephemeral=True,
            )
            return

        if isinstance(result, dict) and result.get("success") is False:
            error_message = str(result.get("error") or "")
            if not error_message:
                error_message = str(apply_result.get("last_error", "Unknown error"))
            _audit_apply_event(
                "error",
                {
                    "contact_id": self.contact_id,
                    "stage": "apply_failed",
                    "job_id": apply_job_id,
                    "job_status": status,
                    "updated_fields": updated_fields,
                    "link_discord_applied": link_discord_applied,
                },
            )
            await interaction.followup.send(
                "❌ Apply completed but returned a failed result. "
                f"Error: {error_message}",
                ephemeral=True,
            )
            return

        if (
            isinstance(result, dict)
            and result.get("success") is True
            and not updated_fields
        ):
            _audit_apply_event(
                "error",
                {
                    "contact_id": self.contact_id,
                    "stage": "apply_no_updates",
                    "job_id": apply_job_id,
                    "job_status": status,
                    "updated_fields": updated_fields,
                    "link_discord_applied": link_discord_applied,
                },
            )
            await interaction.followup.send(
                "❌ Apply reported success but no fields were updated. "
                "Please verify your permissions or contact field mapping and try again.",
                ephemeral=True,
            )
            return

        embed = discord.Embed(
            title="✅ CRM Updated",
            description=f"Applied updates for **{self.contact_name}**.",
            color=0x00FF00,
        )
        display_fields = self._collapse_updated_fields(updated_fields)
        labeled_fields = [self._field_label(field) for field in display_fields]
        updated_fields_value = self._format_updated_fields_value(labeled_fields)
        embed.add_field(
            name="Updated Fields",
            value=updated_fields_value,
            inline=False,
        )
        applied_lines = self._build_applied_updates_lines(
            updated_fields=updated_fields,
            updated_values=updated_values,
        )
        if applied_lines:
            applied_updates_value = self._format_applied_updates_value(applied_lines)
            embed.add_field(
                name="Applied Updates",
                value=applied_updates_value,
                inline=False,
            )
        profile_url = f"{self.crm_cog.base_url}/#Contact/view/{self.contact_id}"
        embed.add_field(name="🔗 CRM Profile", value=f"[View in CRM]({profile_url})")
        _audit_apply_event(
            "success",
            {
                "contact_id": self.contact_id,
                "stage": "apply_succeeded",
                "job_id": apply_job_id,
                "updated_fields": updated_fields,
                "proposed_updates_count": len(self.proposed_updates),
                "link_member_requested": bool(self.link_discord),
                "link_discord_applied": link_discord_applied,
            },
        )
        await interaction.followup.send(embed=embed, ephemeral=True)

        for item in self.children:
            if isinstance(item, discord.ui.Button):
                item.disabled = True
        if interaction.message:
            try:
                await interaction.message.edit(view=self)
            except discord.NotFound:
                pass
            except discord.HTTPException as exc:
                logger.warning("Failed to update confirmation view: %s", exc)

    @discord.ui.button(label="Cancel", style=discord.ButtonStyle.secondary)
    async def cancel_updates(
        self,
        interaction: discord.Interaction,
        button: discord.ui.Button["ResumeUpdateConfirmationView"],
    ) -> None:
        """Cancel CRM updates after preview."""
        self.crm_cog._audit_command(
            interaction=interaction,
            action="crm.upload_resume",
            result="denied",
            metadata={
                "stage": "apply_cancelled",
                "proposed_updates_count": len(self.proposed_updates),
                "link_member_requested": bool(self.link_discord),
            },
            resource_type="crm_contact",
            resource_id=self.contact_id,
        )
        await interaction.response.send_message(
            "No CRM profile updates were applied.", ephemeral=True
        )
        for item in self.children:
            if isinstance(item, discord.ui.Button):
                item.disabled = True
        if interaction.message:
            try:
                await interaction.message.edit(view=self)
            except discord.NotFound:
                pass
            except discord.HTTPException as exc:
                logger.warning("Failed to update confirmation view: %s", exc)


class ResumeReprocessConfirmationView(discord.ui.View):
    """Confirm reprocessing a contact's resume from existing attachment."""

    def __init__(
        self,
        *,
        crm_cog: "CRMCog",
        interaction: discord.Interaction,
        contact_id: str,
        contact_name: str,
        attachment_id: str,
        filename: str,
    ) -> None:
        super().__init__(timeout=180)
        self.crm_cog = crm_cog
        self.requester_id = getattr(interaction.user, "id", 0)
        self.contact_id = contact_id
        self.contact_name = contact_name
        self.attachment_id = attachment_id
        self.filename = filename

    async def interaction_check(self, interaction: discord.Interaction) -> bool:
        """Allow only the original requester to confirm reprocessing."""
        if interaction.user.id != self.requester_id:
            await interaction.response.send_message(
                "❌ Only the command requester can confirm reprocessing.",
                ephemeral=True,
            )
            return False
        return True

    @discord.ui.button(label="Reprocess Resume", style=discord.ButtonStyle.primary)
    async def confirm_reprocess(
        self,
        interaction: discord.Interaction,
        button: discord.ui.Button["ResumeReprocessConfirmationView"],
    ) -> None:
        """Re-run resume extraction for the selected resume attachment."""
        await interaction.response.defer(ephemeral=True, thinking=True)
        await self.crm_cog._run_resume_extract_and_preview(
            interaction=interaction,
            contact_id=self.contact_id,
            contact_name=self.contact_name,
            attachment_id=self.attachment_id,
            filename=self.filename,
            link_member=None,
            action="crm.reprocess_resume",
            status_message="🔄 Reprocessing resume and extracting profile fields now...",
        )
        for item in self.children:
            if isinstance(item, discord.ui.Button):
                item.disabled = True
        if interaction.message:
            try:
                await interaction.message.edit(view=self)
            except discord.NotFound:
                pass
            except discord.HTTPException as exc:
                logger.warning("Failed to update confirmation view: %s", exc)

    @discord.ui.button(label="Cancel", style=discord.ButtonStyle.secondary)
    async def cancel_reprocess(
        self,
        interaction: discord.Interaction,
        button: discord.ui.Button["ResumeReprocessConfirmationView"],
    ) -> None:
        """Cancel the reprocess request."""
        self.crm_cog._audit_command(
            interaction=interaction,
            action="crm.reprocess_resume",
            result="denied",
            metadata={
                "contact_id": self.contact_id,
                "contact_name": self.contact_name,
                "stage": "reprocess_cancelled",
            },
            resource_type="crm_contact",
            resource_id=self.contact_id,
        )
        await interaction.response.send_message(
            "Reprocess cancelled. No changes were made.",
            ephemeral=True,
        )
        for item in self.children:
            if isinstance(item, discord.ui.Button):
                item.disabled = True
        if interaction.message:
            try:
                await interaction.message.edit(view=self)
            except discord.NotFound:
                pass
            except discord.HTTPException as exc:
                logger.warning("Failed to update confirmation view: %s", exc)


class ResumeCreateContactView(discord.ui.View):
    """Prompt to create a new contact from parsed resume data."""

    def __init__(
        self,
        crm_cog: "CRMCog",
        interaction: discord.Interaction,
        file_content: bytes,
        filename: str,
        file_size: int,
        search_term: str | None,
        overwrite: bool,
        link_user: discord.Member | None,
        inferred_contact_meta: dict[str, Any] | None,
        target_scope: str,
        create_payload_override: dict[str, str] | None = None,
        created_target_scope: str = "created",
    ) -> None:
        super().__init__(timeout=180)
        self.crm_cog = crm_cog
        self.original_interaction = interaction
        self.file_content = file_content
        self.filename = filename
        self.file_size = file_size
        self.search_term = search_term
        self.overwrite = overwrite
        self.link_user = link_user
        self.inferred_contact_meta = inferred_contact_meta
        self.target_scope = target_scope
        self.create_payload_override = create_payload_override
        self.created_target_scope = created_target_scope

    async def interaction_check(self, interaction: discord.Interaction) -> bool:
        if interaction.user.id != self.original_interaction.user.id:
            await interaction.response.send_message(
                "❌ Only the command requester can confirm contact creation.",
                ephemeral=True,
            )
            return False
        return True

    async def _finalize(
        self, interaction: discord.Interaction, *, error_message: str | None = None
    ) -> None:
        for item in self.children:
            if isinstance(item, discord.ui.Button):
                item.disabled = True
        if interaction.message:
            try:
                if error_message:
                    await interaction.followup.send(error_message, ephemeral=True)
                await interaction.message.edit(view=self)
            except discord.NotFound:
                pass
            except discord.HTTPException as exc:
                logger.warning("Failed to update create-contact view: %s", exc)

    @discord.ui.button(label="Create Contact", style=discord.ButtonStyle.primary)
    async def confirm_create(
        self,
        interaction: discord.Interaction,
        button: discord.ui.Button["ResumeCreateContactView"],
    ) -> None:
        """Create the inferred contact and continue resume upload."""
        await interaction.response.defer(ephemeral=True)
        create_payload: dict[str, str] | None = None
        try:
            if self.create_payload_override:
                create_payload = dict(self.create_payload_override)
            else:
                create_payload = (
                    await self.crm_cog._build_resume_create_contact_payload_async(
                        file_content=self.file_content,
                        filename=self.filename,
                    )
                )
            self.crm_cog._populate_name_fields(
                create_payload,
                source_name=str(create_payload.get("name", "")).strip(),
            )
            target_contact = self.crm_cog.espo_api.request(
                "POST", "Contact", create_payload
            )
            contact_id = (
                target_contact.get("id") if isinstance(target_contact, dict) else None
            )

            if not contact_id:
                raise ValueError("Created contact had no valid ID.")

            logger.info(
                "Created new contact %s from resume for %s",
                contact_id,
                interaction.user.name,
            )

            await self.crm_cog._upload_resume_attachment_to_contact(
                interaction=interaction,
                file_content=self.file_content,
                filename=self.filename,
                file_size=self.file_size,
                contact=target_contact,
                target_scope=self.created_target_scope,
                search_term=self.search_term,
                overwrite=self.overwrite,
                link_user=self.link_user,
                inferred_contact_meta=self.inferred_contact_meta,
            )
        except Exception as exc:
            status_code = getattr(self.crm_cog.espo_api, "status_code", None)
            error_detail = self.crm_cog._sanitize_error_message_for_discord(exc)
            status_note = f" (status {status_code})" if status_code else ""
            logger.exception(
                "Failed to create contact from resume filename=%s target_scope=%s inferred_meta=%s "
                "status_code=%s payload=%s error=%s",
                self.filename,
                self.target_scope,
                self.inferred_contact_meta,
                status_code,
                create_payload,
                error_detail,
            )
            audit_metadata: dict[str, Any] = {
                "filename": self.filename,
                "target_scope": self.target_scope,
                "reason": "contact_create_failed",
                "error": error_detail,
                "status_code": status_code,
            }
            if create_payload:
                audit_metadata["create_payload_keys"] = sorted(create_payload.keys())
            self.crm_cog._audit_command(
                interaction=interaction,
                action="crm.upload_resume",
                result="error",
                metadata=audit_metadata,
            )
            await interaction.followup.send(
                f"⚠️ Could not create a contact from this resume: `{error_detail}`{status_note}. "
                "Please provide `search_term` or `link_user`.",
                ephemeral=True,
            )
        finally:
            await self._finalize(interaction)

    @discord.ui.button(label="Cancel", style=discord.ButtonStyle.secondary)
    async def cancel_create(
        self,
        interaction: discord.Interaction,
        button: discord.ui.Button["ResumeCreateContactView"],
    ) -> None:
        """Cancel contact creation and keep the queue untouched."""
        self.crm_cog._audit_command(
            interaction=interaction,
            action="crm.upload_resume",
            result="denied",
            metadata={
                "filename": self.filename,
                "target_scope": self.target_scope,
                "reason": "create_contact_cancelled",
            },
            resource_type="crm_contact",
        )
        await interaction.response.send_message(
            "Contact creation cancelled. No changes were made.",
            ephemeral=True,
        )
        await self._finalize(interaction)


class OnboardingQueuePagerView(discord.ui.View):
    """View for paging through onboarding queue entries one person at a time."""

    def __init__(
        self,
        crm_cog: "CRMCog",
        interaction: discord.Interaction,
        queue_rows: list[dict[str, str]],
        *,
        page_size: int = ONBOARDING_QUEUE_PAGE_SIZE,
    ) -> None:
        super().__init__(timeout=300)
        self.crm_cog = crm_cog
        self.requester_id = getattr(interaction.user, "id", 0)
        self.queue_rows = queue_rows
        self.page_size = max(1, page_size)
        self.page_index = 0
        self._message: discord.Message | None = None
        self.total_pages = (
            (len(self.queue_rows) - 1) // self.page_size + 1 if self.queue_rows else 0
        )
        self._update_button_states()

    async def interaction_check(self, interaction: discord.Interaction) -> bool:
        """Allow only the original command requester to page through the queue."""
        if interaction.user.id != self.requester_id:
            await interaction.response.send_message(
                "❌ Only the command requester can page through this queue.",
                ephemeral=True,
            )
            return False
        return True

    def _build_embed(self) -> discord.Embed:
        return self.crm_cog._build_onboarding_queue_page_embed(
            self.queue_rows, page_index=self.page_index, page_size=self.page_size
        )

    def _set_message(self, message: discord.Message | None) -> None:
        self._message = message

    async def on_timeout(self) -> None:
        """Disable controls after pager timeout."""
        for item in self.children:
            if isinstance(item, discord.ui.Button):
                item.disabled = True
        if self._message:
            try:
                await self._message.edit(view=self)
            except discord.NotFound:
                pass
            except discord.HTTPException as exc:
                logger.warning("Failed to disable onboarding queue pager view: %s", exc)

    def _update_button_states(self) -> None:
        if not self.children:
            return
        if len(self.children) >= 1:
            next_button = self.children[0]
            if isinstance(next_button, discord.ui.Button):
                next_button.disabled = self.page_index >= self.total_pages - 1

    @discord.ui.button(label="Next", style=discord.ButtonStyle.primary, emoji="▶️")
    async def next_page(
        self,
        interaction: discord.Interaction,
        button: discord.ui.Button["OnboardingQueuePagerView"],
    ) -> None:
        """Show next contact in the onboarding queue."""
        if self.page_index >= self.total_pages - 1:
            return

        self.page_index += 1
        self._update_button_states()
        await interaction.response.edit_message(embed=self._build_embed(), view=self)


class CRMCog(commands.Cog):
    """CRM integration cog for EspoCRM operations."""

    def __init__(self, bot: commands.Bot) -> None:
        self.bot = bot
        # Construct API URL from base URL
        api_url = settings.espo_base_url.rstrip("/") + "/api/v1"
        self.espo_api = EspoAPI(api_url, settings.espo_api_key)
        # Store base URL for profile links
        self.base_url = settings.espo_base_url.rstrip("/")
        self.resume_extractor = ResumeProfileExtractor(
            api_key=settings.openai_api_key,
            base_url=settings.openai_base_url,
            model=settings.openai_model,
        )
        self._resume_profile_cache: (
            tuple[tuple[int, str], ResumeExtractedProfile] | None
        ) = None
        self.audit_logger = DiscordAuditLogger(
            base_url=settings.audit_api_base_url,
            shared_secret=settings.api_shared_secret,
            timeout_seconds=settings.audit_api_timeout_seconds,
            discord_logs_webhook_url=settings.discord_logs_webhook_url,
            discord_logs_webhook_wait=settings.discord_logs_webhook_wait,
        )

    @staticmethod
    def _configured_linkedin_field() -> str:
        """Return the configured field for LinkedIn profile values."""
        return _configured_linkedin_field_from_settings()

    @staticmethod
    def _sanitize_error_message_for_discord(
        raw_error: Any,
        max_length: int = 1900,
    ) -> str:
        """Normalize and truncate error text for safe Discord/log output."""
        text = str(raw_error).strip()
        if not text:
            return "Unknown error"

        text = text.replace("`", "'")
        text = re.sub(
            r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]",
            " ",
            text,
        )
        text = re.sub(r"\s+", " ", text).strip()

        if len(text) <= max_length:
            return text

        if max_length <= 1:
            return text[:max_length]

        return text[: max_length - 1].rstrip() + "…"

    def _audit_command(
        self,
        *,
        interaction: discord.Interaction,
        action: str,
        result: str,
        metadata: dict[str, Any] | None = None,
        resource_type: str | None = "discord_command",
        resource_id: str | None = None,
    ) -> None:
        """Queue a best-effort audit write for CRM command activity."""
        self.audit_logger.log_command(
            interaction=interaction,
            action=action,
            result=result,
            metadata=metadata,
            resource_type=resource_type,
            resource_id=resource_id,
        )

    def _backend_headers(self) -> dict[str, str]:
        """Build auth headers for internal backend API calls."""
        if not settings.api_shared_secret:
            raise ValueError("API_SHARED_SECRET is required for backend API requests.")
        return {
            "X-API-Secret": settings.api_shared_secret,
            "Content-Type": "application/json",
        }

    async def _backend_request_json(
        self,
        method: Literal["GET", "POST"],
        path: str,
        *,
        expected_status: int,
        payload: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        timeout = aiohttp.ClientTimeout(total=30)
        request_kwargs: dict[str, Any] = {
            "headers": self._backend_headers(),
            "timeout": timeout,
        }
        if payload is not None:
            request_kwargs["json"] = payload

        async with aiohttp.ClientSession() as session:
            async with session.request(
                method,
                self._backend_url(path),
                **request_kwargs,
            ) as response:
                data = await response.json()
                if response.status != expected_status:
                    raise ValueError(f"Backend {method} {path} failed: {data}")
                if not isinstance(data, dict):
                    raise ValueError(
                        f"Backend {method} {path} returned a non-object response."
                    )
                return data

    def _backend_url(self, path: str) -> str:
        return f"{settings.backend_api_base_url.rstrip('/')}{path}"

    async def _enqueue_resume_extract_job(
        self, *, contact_id: str, attachment_id: str, filename: str
    ) -> str:
        payload = {
            "contact_id": contact_id,
            "attachment_id": attachment_id,
            "filename": filename,
        }
        data = await self._backend_request_json(
            "POST",
            "/jobs/resume-extract",
            payload=payload,
            expected_status=202,
        )
        job_id = data.get("job_id")
        if not isinstance(job_id, str) or not job_id:
            raise ValueError("Missing backend extract job_id in response.")
        return job_id

    async def _enqueue_resume_apply_job(
        self,
        *,
        contact_id: str,
        updates: dict[str, Any],
        link_discord: dict[str, str] | None = None,
    ) -> str:
        payload = {
            "contact_id": contact_id,
            "updates": updates,
            "link_discord": link_discord,
        }
        data = await self._backend_request_json(
            "POST",
            "/jobs/resume-apply",
            payload=payload,
            expected_status=202,
        )
        job_id = data.get("job_id")
        if not isinstance(job_id, str) or not job_id:
            raise ValueError("Missing backend apply job_id in response.")
        return job_id

    async def _get_backend_job_status(self, job_id: str) -> dict[str, Any]:
        return await self._backend_request_json(
            "GET",
            f"/jobs/{job_id}",
            expected_status=200,
        )

    async def _wait_for_backend_job_result(
        self, job_id: str, *, timeout_seconds: int = 180, poll_seconds: float = 2.0
    ) -> dict[str, Any] | None:
        """Poll backend job status until terminal or timeout."""
        terminal = {"succeeded", "dead", "canceled"}
        max_attempts = max(1, int(timeout_seconds / poll_seconds))

        for _ in range(max_attempts):
            job = await self._get_backend_job_status(job_id)
            status = str(job.get("status", ""))
            if status in terminal:
                return job
            await asyncio.sleep(poll_seconds)

        return None

    def _resolve_field_name(
        self, contact: dict[str, Any], *, candidates: tuple[str, ...]
    ) -> str | None:
        """Return the first matching field name that exists on a contact."""
        for field_name in candidates:
            if field_name in contact:
                return field_name
        return None

    def _normalize_onboarding_state(self, value: Any) -> str:
        """Normalize onboarding state for comparisons."""
        if value is None:
            return ""
        return str(value).strip().lower()

    def _format_onboarding_updated_at(self, raw_value: Any) -> str:
        """Normalize the onboarding updated-at value for display."""
        if raw_value is None:
            return "Unknown"

        if isinstance(raw_value, (int, float)):
            try:
                return datetime.fromtimestamp(raw_value, tz=timezone.utc).strftime(
                    "%Y-%m-%d %H:%M UTC"
                )
            except (OSError, OverflowError, ValueError):
                return str(raw_value)

        raw_value_text = str(raw_value).strip()
        if not raw_value_text:
            return "Unknown"

        try:
            parsed = datetime.fromisoformat(raw_value_text.replace("Z", "+00:00"))
        except ValueError:
            return raw_value_text

        if parsed.tzinfo is None:
            if parsed.time() and parsed.time() != datetime.min.time():
                return parsed.strftime("%Y-%m-%d %H:%M")
            return parsed.strftime("%Y-%m-%d")

        parsed_utc = parsed.astimezone(timezone.utc)
        if parsed_utc.time() and parsed_utc.time() != datetime.min.time():
            return parsed_utc.strftime("%Y-%m-%d %H:%M UTC")
        return parsed_utc.strftime("%Y-%m-%d")

    async def _resolve_onboarder_username(
        self, interaction: discord.Interaction, raw_onboarder: str
    ) -> str | None:
        """Resolve onboarder input into a normalized 508 username."""
        candidate = (raw_onboarder or "").strip()
        if not candidate:
            return None

        mention_match = re.fullmatch(r"<@!?(?P<user_id>\d+)>", candidate)
        if mention_match:
            discord_id = mention_match.group("user_id")
            try:
                contact = await self._find_contact_by_discord_id(discord_id)
            except ValueError:
                contact = None

            if contact:
                linked_username = self._normalize_508_username(
                    contact.get("c508Email") or ""
                )
                if linked_username:
                    return linked_username
            return None

        return self._normalize_508_username(candidate)

    def _format_contact_card(
        self,
        contact: dict[str, Any],
        interaction: discord.Interaction | None = None,
        *,
        additional_fields: list[tuple[str, str]] | None = None,
    ) -> str:
        """Build a reusable contact card text with optional additional lines."""
        email = contact.get("emailAddress", "No email")
        contact_type = contact.get("type", "Unknown")
        email_508 = contact.get("c508Email", "None")
        discord_username = contact.get("cDiscordUsername") or "No Discord"
        discord_user_id = contact.get("cDiscordUserID")
        contact_id = contact.get("id", "")

        clean_discord_username = discord_username
        if discord_username and " (ID: " in discord_username:
            clean_discord_username = discord_username.split(" (ID: ")[0]

        discord_display = clean_discord_username
        if (
            discord_user_id
            and discord_user_id != "No Discord"
            and interaction
            and interaction.guild
        ):
            try:
                member = interaction.guild.get_member(int(discord_user_id))
                if member:
                    discord_display = f"{member.mention} ({clean_discord_username})"
            except (ValueError, AttributeError):
                pass

        contact_info = f"📧 {email}\n🏷️ Type: {contact_type}"
        if contact_type in ["Candidate / Member", "Member"]:
            contact_info += (
                f"\n🏢 508 Email: {email_508}\n💬 Discord: {discord_display}"
            )

        if additional_fields:
            for label, value in additional_fields:
                normalized_value = str(value).strip() if value is not None else ""
                if normalized_value:
                    contact_info += f"\n{label}: {normalized_value}"

        if contact_id:
            profile_url = f"{self.base_url}/#Contact/view/{contact_id}"
            contact_info = f"🔗 [View in CRM]({profile_url})\n{contact_info}"

        return contact_info

    def _build_onboarding_queue_row(
        self, contact_record: dict[str, Any], status: str
    ) -> dict[str, str]:
        """Build a compact dictionary for one onboarding queue row."""
        email = str(contact_record.get("emailAddress") or "No email")
        name = str(contact_record.get("name") or "Unknown")
        contact_id = str(contact_record.get("id") or "")

        discord_user_id = contact_record.get("cDiscordUserID")
        discord_username = contact_record.get("cDiscordUsername")
        clean_discord_username = "No Discord"
        if isinstance(discord_username, str) and discord_username.strip():
            clean_discord_username = (
                discord_username.split(" (ID: ")[0]
                if " (ID: " in discord_username
                else discord_username.strip()
            )

        onboarder_field = self._resolve_field_name(
            contact_record, candidates=ONBOARDER_FIELD_CANDIDATES
        )
        onboarder_value = (
            str(contact_record.get(onboarder_field, "")).strip()
            if onboarder_field
            else ""
        )

        return {
            "name": name,
            "email": email,
            "status": status or "Unknown",
            "onboarder": onboarder_value or "Unassigned",
            "discord_user": clean_discord_username,
            "discord_user_id": str(discord_user_id or ""),
            "onboarding_updated_at": self._format_onboarding_updated_at(
                contact_record.get("cOnboardingUpdatedAt")
            ),
            "crm_url": f"{self.base_url}/#Contact/view/{contact_id}"
            if contact_id
            else "",
            "id": contact_id,
        }

    def _build_onboarding_queue_rows(
        self,
        interaction: discord.Interaction,
        queue_entries: list[tuple[dict[str, Any], str]],
    ) -> list[dict[str, str]]:
        """Build compact per-row onboarding data."""
        rows: list[dict[str, str]] = []
        for contact_record, status in queue_entries:
            row = self._build_onboarding_queue_row(contact_record, status)

            discord_user_id = row.get("discord_user_id")
            discord_display = row.get("discord_user", "No Discord")
            if discord_user_id and interaction.guild:
                try:
                    member = interaction.guild.get_member(int(discord_user_id))
                    if member:
                        discord_display = f"{member.mention} ({discord_display})"
                except (TypeError, ValueError):
                    pass
            row["discord_user"] = discord_display
            rows.append(row)
        return rows

    def _build_onboarding_queue_page_embed(
        self,
        queue_rows: list[dict[str, str]],
        *,
        page_index: int,
        page_size: int,
    ) -> discord.Embed:
        """Build one-page onboarding queue embed."""
        if not queue_rows:
            return discord.Embed(
                title="📋 Onboarding Queue",
                description="No onboarding contacts were found.",
                color=0x0099FF,
            )

        total_rows = len(queue_rows)
        start = max(0, page_index) * page_size
        end = min(start + page_size, total_rows)
        shown_rows = queue_rows[start:end]

        embed = discord.Embed(
            title="📋 Onboarding Queue",
            description=(
                "Contacts currently outside `onboarded`, `waitlist`, and "
                f"`rejected` states. Showing {start + 1}-{end} of {total_rows}."
            ),
            color=0x0099FF,
        )

        for row in shown_rows:
            details = [
                f"📧 **Email:** {row.get('email', 'No email')}",
                f"👤 **Name:** {row.get('name', 'Unknown')}",
                f"💬 **Linked Discord:** {row.get('discord_user', 'No Discord')}",
                f"🕘 **cOnboardingUpdatedAt:** {row.get('onboarding_updated_at', 'Unknown')}",
                f"📌 **cOnboardingState:** {row.get('status', 'Unknown')}",
                f"🧑‍💼 **cOnboarder:** {row.get('onboarder', 'Unassigned')}",
            ]
            crm_url = row.get("crm_url", "")
            if crm_url:
                details.append(f"🔗 [View in CRM]({crm_url})")
            else:
                details.append("🔗 **CRM:** Unavailable")

            embed.add_field(
                name=f"Contact: {row.get('name', 'Unknown')}",
                value="\n".join(details),
                inline=False,
            )

        total_pages = (total_rows - 1) // page_size + 1
        embed.set_footer(text=f"Page {page_index + 1} of {total_pages}")
        return embed

    def _build_resume_preview_embed(
        self,
        *,
        contact_id: str,
        contact_name: str,
        result: dict[str, Any],
        link_member: discord.Member | None,
    ) -> tuple[discord.Embed, dict[str, Any]]:
        """Render backend extraction result as a Discord preview embed."""
        proposed_updates_raw = result.get("proposed_updates")
        proposed_updates: dict[str, Any] = {}
        if isinstance(proposed_updates_raw, dict):
            proposed_updates = {
                str(field): value
                for field, value in proposed_updates_raw.items()
                if value is not None
                and not (
                    isinstance(value, (dict, list, tuple, set)) and len(value) == 0
                )
                and (not isinstance(value, str) or value.strip())
            }

        changes = result.get("proposed_changes")
        new_skills = result.get("new_skills")
        skipped = result.get("skipped")
        extracted_profile = result.get("extracted_profile")

        embed = discord.Embed(
            title="🧾 Resume Parsed",
            description=f"Review extracted updates for **{contact_name}**.",
            color=0x0099FF,
        )

        if isinstance(changes, list) and changes:
            lines: list[str] = []
            for change in changes[:8]:
                if not isinstance(change, dict):
                    continue
                label = str(change.get("label", change.get("field", "Field")))
                current = str(change.get("current", "None"))
                proposed = str(change.get("proposed", ""))
                lines.append(f"**{label}**: `{current}` → `{proposed}`")
            embed.add_field(
                name="Proposed Changes",
                value="\n".join(lines) if lines else "No changes",
                inline=False,
            )
        else:
            embed.add_field(
                name="Proposed Changes",
                value="No CRM field updates were extracted.",
                inline=False,
            )

        if isinstance(new_skills, list) and new_skills:
            formatted_skills = ", ".join(str(skill) for skill in new_skills[:25])
            embed.add_field(
                name="New Skills",
                value=formatted_skills,
                inline=False,
            )

        if isinstance(skipped, list) and skipped:
            skip_lines: list[str] = []
            for item in skipped[:4]:
                if not isinstance(item, dict):
                    continue
                field = str(item.get("field", "field"))
                reason = str(item.get("reason", "Skipped"))
                value = str(item.get("value", ""))
                skip_lines.append(f"`{field}`: `{value}` ({reason})")
            if skip_lines:
                embed.add_field(
                    name="Skipped",
                    value="\n".join(skip_lines),
                    inline=False,
                )

        if isinstance(extracted_profile, dict):
            confidence = extracted_profile.get("confidence")
            source = extracted_profile.get("source")
            if confidence is not None or source:
                embed.add_field(
                    name="Extraction",
                    value=f"Source: `{source or 'unknown'}` | Confidence: `{confidence}`",
                    inline=False,
                )

        if link_member:
            embed.add_field(
                name="Discord Link",
                value=f"Will link contact to {link_member.mention}",
                inline=False,
            )

        profile_url = f"{self.base_url}/#Contact/view/{contact_id}"
        embed.add_field(name="🔗 CRM Profile", value=f"[View in CRM]({profile_url})")
        return embed, proposed_updates

    async def _run_resume_extract_and_preview(
        self,
        interaction: discord.Interaction,
        contact_id: str,
        contact_name: str,
        attachment_id: str,
        filename: str,
        link_member: discord.Member | None,
        *,
        action: str = "crm.upload_resume",
        status_message: str | None = None,
    ) -> None:
        """Kick off worker extraction and show confirmation preview."""
        action_name = action
        status_text = (
            status_message or "📥 Resume uploaded. Extracting profile fields now..."
        )
        try:
            job_id = await self._enqueue_resume_extract_job(
                contact_id=contact_id,
                attachment_id=attachment_id,
                filename=filename,
            )
        except Exception as exc:
            logger.error("Failed to enqueue resume extract job: %s", exc)
            self._audit_command(
                interaction=interaction,
                action=action_name,
                result="error",
                metadata={
                    "filename": filename,
                    "attachment_id": attachment_id,
                    "stage": "extract_enqueue",
                    "error": str(exc),
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await interaction.followup.send(
                "⚠️ Resume uploaded, but extraction job could not be enqueued.",
                ephemeral=True,
            )
            return

        await interaction.followup.send(
            status_text,
            ephemeral=True,
        )

        try:
            job = await self._wait_for_backend_job_result(job_id)
        except Exception as exc:
            logger.error("Worker polling failed for job_id=%s error=%s", job_id, exc)
            self._audit_command(
                interaction=interaction,
                action=action_name,
                result="error",
                metadata={
                    "filename": filename,
                    "attachment_id": attachment_id,
                    "job_id": job_id,
                    "stage": "extract_polling",
                    "error": str(exc),
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await interaction.followup.send(
                "⚠️ Resume uploaded, but extraction polling failed.",
                ephemeral=True,
            )
            return
        if not job:
            self._audit_command(
                interaction=interaction,
                action=action_name,
                result="error",
                metadata={
                    "filename": filename,
                    "attachment_id": attachment_id,
                    "job_id": job_id,
                    "stage": "extract_timeout",
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await interaction.followup.send(
                "⚠️ Timed out waiting for extraction result. Try again in a moment.",
                ephemeral=True,
            )
            return

        status = str(job.get("status", "unknown"))
        if status != "succeeded":
            self._audit_command(
                interaction=interaction,
                action=action_name,
                result="error",
                metadata={
                    "filename": filename,
                    "attachment_id": attachment_id,
                    "job_id": job_id,
                    "stage": "extract_failed",
                    "job_status": status,
                    "last_error": str(job.get("last_error", "")),
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await interaction.followup.send(
                f"❌ Extraction job failed (status: {status}). "
                f"Error: {job.get('last_error') or 'Unknown error'}",
                ephemeral=True,
            )
            return

        result = job.get("result")
        if not isinstance(result, dict):
            self._audit_command(
                interaction=interaction,
                action=action_name,
                result="error",
                metadata={
                    "filename": filename,
                    "attachment_id": attachment_id,
                    "job_id": job_id,
                    "stage": "extract_malformed_result",
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await interaction.followup.send(
                "❌ Extraction result was empty or malformed.",
                ephemeral=True,
            )
            return

        if not result.get("success", False):
            self._audit_command(
                interaction=interaction,
                action=action_name,
                result="error",
                metadata={
                    "filename": filename,
                    "attachment_id": attachment_id,
                    "job_id": job_id,
                    "stage": "extract_unsuccessful",
                    "error": str(result.get("error", "")),
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await interaction.followup.send(
                f"❌ Resume extraction failed: {result.get('error') or 'Unknown error'}",
                ephemeral=True,
            )
            return

        embed, proposed_updates = self._build_resume_preview_embed(
            contact_id=contact_id,
            contact_name=contact_name,
            result=result,
            link_member=link_member,
        )

        if not proposed_updates and not link_member:
            self._audit_command(
                interaction=interaction,
                action=action_name,
                result="success",
                metadata={
                    "filename": filename,
                    "attachment_id": attachment_id,
                    "job_id": job_id,
                    "stage": "preview_no_changes",
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await interaction.followup.send(embed=embed, ephemeral=True)
            return

        link_discord_payload: dict[str, str] | None = None
        if link_member:
            link_discord_payload = {
                "user_id": str(link_member.id),
                "username": str(link_member),
            }

        view = ResumeUpdateConfirmationView(
            crm_cog=self,
            requester_id=interaction.user.id,
            contact_id=contact_id,
            contact_name=contact_name,
            proposed_updates=proposed_updates,
            link_discord=link_discord_payload,
        )
        self._audit_command(
            interaction=interaction,
            action=action_name,
            result="success",
            metadata={
                "filename": filename,
                "attachment_id": attachment_id,
                "job_id": job_id,
                "stage": "preview_ready",
                "proposed_updates_count": len(proposed_updates),
                "link_member_requested": bool(link_member),
            },
            resource_type="crm_contact",
            resource_id=str(contact_id),
        )
        await interaction.followup.send(embed=embed, view=view, ephemeral=True)

    async def _download_and_send_resume(
        self, interaction: discord.Interaction, contact_name: str, resume_id: str
    ) -> bool:
        """Download and send a resume file as a Discord attachment."""
        try:
            # Download the resume file
            file_content = self.espo_api.download_file(f"Attachment/file/{resume_id}")

            # Get file metadata to determine filename
            file_info = self.espo_api.request("GET", f"Attachment/{resume_id}")
            filename = file_info.get("name", f"{contact_name}_resume.pdf")

            # Create Discord file object
            file_buffer = io.BytesIO(file_content)
            discord_file = discord.File(file_buffer, filename=filename)

            await interaction.followup.send(
                f"📄 Resume for **{contact_name}**:", file=discord_file
            )
            return True

        except EspoAPIError as e:
            logger.error(f"Failed to download resume {resume_id}: {e}")
            await interaction.followup.send(f"❌ Failed to download resume: {str(e)}")
            return False

    def _check_member_role(self, interaction: discord.Interaction) -> bool:
        """Check if user has Member role or higher for resume access."""
        if not hasattr(interaction.user, "roles"):
            return False
        return check_user_roles_with_hierarchy(interaction.user.roles, ["Member"])

    def _parse_contact_skill_attrs(self, value: Any) -> dict[str, int]:
        """Parse skill attributes from a contact record into normalized strengths."""
        if value is None:
            return {}

        candidate = value
        if isinstance(candidate, str):
            raw_value = candidate.strip()
            if not raw_value:
                return {}
            parsed_payload = self._parse_json_object_with_recovery(raw_value)
            if parsed_payload is None:
                return {}
            candidate = parsed_payload

        if not isinstance(candidate, dict):
            return {}

        parsed: dict[str, int] = {}
        for raw_skill, payload in candidate.items():
            normalized_skill = self._normalize_skill(str(raw_skill))
            if not normalized_skill:
                continue

            if not isinstance(payload, dict):
                continue

            raw_strength = payload.get("strength")
            if raw_strength is None:
                continue
            try:
                strength = int(float(raw_strength))
            except (TypeError, ValueError):
                continue

            if not 1 <= strength <= 5:
                continue

            parsed[normalized_skill.casefold()] = strength

        return parsed

    def _normalize_skill(self, value: str) -> str:
        """Normalize one skill name from source data."""
        normalized = normalize_skill_list([value])
        return normalized[0] if normalized else ""

    def _parse_skill_updates(
        self, skills: str
    ) -> tuple[list[str], dict[str, int], list[str]]:
        """Parse comma-separated skills with optional `skill:level` syntax."""
        parsed_skills: list[str] = []
        requested_strengths: dict[str, int] = {}
        invalid_entries: list[str] = []
        seen: set[str] = set()

        for raw_token in skills.replace(";", ",").split(","):
            token = raw_token.strip()
            if not token:
                continue

            token_skill = token
            strength_value: int | None = None
            if ":" in token:
                token_skill, raw_strength = token.rsplit(":", 1)
                token_skill = token_skill.strip()
                raw_strength = raw_strength.strip()

                if not token_skill or not raw_strength:
                    invalid_entries.append(token)
                    continue

                try:
                    parsed_strength = int(float(raw_strength))
                except (TypeError, ValueError):
                    invalid_entries.append(token)
                    continue

                if not 1 <= parsed_strength <= 5:
                    invalid_entries.append(token)
                    continue

                strength_value = parsed_strength

            normalized_skill = self._normalize_skill(token_skill)
            if not normalized_skill:
                invalid_entries.append(token)
                continue

            key = normalized_skill.casefold()
            if key in seen:
                if strength_value is not None:
                    requested_strengths[key] = strength_value
                continue

            seen.add(key)
            parsed_skills.append(normalized_skill)
            if strength_value is not None:
                requested_strengths[key] = strength_value

        return parsed_skills, requested_strengths, invalid_entries

    def _serialize_skill_attrs(self, attrs: dict[str, int]) -> str:
        """Serialize normalized skill strengths in the CRM-compatible format."""
        payload = {
            skill.casefold(): {"strength": max(1, min(5, int(strength)))}
            for skill, strength in attrs.items()
            if skill
        }
        return json.dumps(payload, separators=(",", ":"), sort_keys=True)

    def _merge_skill_update_payload(
        self,
        contact: dict[str, Any],
        requested_skills: list[str],
        requested_strengths: dict[str, int],
    ) -> tuple[str, str]:
        """Merge requested skills with existing contact skills and attributes."""
        raw_skills = contact.get("skills", "")
        if isinstance(raw_skills, list):
            raw_skill_values = [str(item) for item in raw_skills if str(item).strip()]
        else:
            raw_skill_values = [
                item.strip() for item in str(raw_skills).split(",") if item.strip()
            ]

        existing_skills = normalize_skill_list(raw_skill_values)
        existing_skill_keys = {skill.casefold() for skill in existing_skills}
        merged_skills = list(existing_skills)
        merged_attrs = self._parse_contact_skill_attrs(contact.get("cSkillAttrs"))

        for requested_skill in requested_skills:
            key = requested_skill.casefold()
            if key not in existing_skill_keys:
                merged_skills.append(requested_skill)
                existing_skill_keys.add(key)

            requested_strength = requested_strengths.get(key)
            if requested_strength is None:
                requested_strength = merged_attrs.get(key, 3)

            merged_attrs[key] = requested_strength

        for existing_skill in merged_skills:
            key = existing_skill.casefold()
            merged_attrs.setdefault(key, 3)

        return ", ".join(merged_skills), self._serialize_skill_attrs(merged_attrs)

    def _format_requested_skills(
        self, requested_skills: list[str], contact: dict[str, Any]
    ) -> str:
        """Format requested skills with strength values from contact attributes."""
        if not requested_skills:
            return ""

        skill_attrs = self._parse_contact_skill_attrs(contact.get("cSkillAttrs"))
        if not skill_attrs:
            return ", ".join(requested_skills)

        rendered: list[str] = []
        for skill in requested_skills:
            normalized_skill = self._normalize_skill(skill)
            if not normalized_skill:
                continue

            strength = skill_attrs.get(normalized_skill.casefold())
            if strength is None:
                rendered.append(skill)
            else:
                rendered.append(f"{skill} ({strength})")

        return ", ".join(rendered)

    @app_commands.command(
        name="search-members", description="Search for candidates / members in the CRM"
    )
    @app_commands.describe(
        query="Search term (name, Discord username, email, or 508 email)",
        skills="Comma-separated skills (AND match)",
    )
    @require_role("Member")
    async def search_members(
        self,
        interaction: discord.Interaction,
        query: str | None = None,
        skills: str | None = None,
    ) -> None:
        """Search for contacts in the CRM."""
        try:
            await interaction.response.defer(ephemeral=True)

            query_value = (query or "").strip()
            raw_skills_list = (
                [skill.strip() for skill in skills.split(",") if skill.strip()]
                if skills
                else []
            )
            skills_list = normalize_skill_list(raw_skills_list)

            if not query_value and not skills_list:
                self._audit_command(
                    interaction=interaction,
                    action="crm.search_members",
                    result="denied",
                    metadata={"reason": "missing_query_and_skills"},
                )
                await interaction.followup.send(
                    "❌ Please provide a search term or skills to search by."
                )
                return

            search_parts = []
            if query_value:
                search_parts.append(f"`{query_value}`")
            if skills_list:
                search_parts.append(f"skills: `{', '.join(skills_list)}`")
            search_summary = ", ".join(search_parts)

            # Search contacts using EspoCRM API
            where_filters = []
            if query_value:
                where_filters.append(
                    {
                        "type": "or",
                        "value": [
                            {
                                "type": "contains",
                                "attribute": "name",
                                "value": query_value,
                            },
                            {
                                "type": "contains",
                                "attribute": "cDiscordUsername",
                                "value": query_value,
                            },
                            {
                                "type": "contains",
                                "attribute": "emailAddress",
                                "value": query_value,
                            },
                            {
                                "type": "contains",
                                "attribute": "c508Email",
                                "value": query_value,
                            },
                        ],
                    }
                )

            if skills_list:
                where_filters.append(
                    {
                        "type": "arrayAllOf",
                        "attribute": "skills",
                        "value": skills_list,
                    }
                )

            search_params = {
                "where": where_filters,
                "maxSize": 10,
                "select": "id,name,emailAddress,c508Email,cDiscordUsername,cDiscordUserID,phoneNumber,type,resumeIds,resumeNames,resumeTypes,skills,cSkillAttrs",
            }

            response = self.espo_api.request("GET", "Contact", search_params)
            contacts = response.get("list", [])

            if not contacts:
                self._audit_command(
                    interaction=interaction,
                    action="crm.search_members",
                    result="success",
                    metadata={
                        "query": query_value or None,
                        "skills": skills_list,
                        "contacts_found": 0,
                    },
                )
                await interaction.followup.send(
                    f"🔍 No contacts found for: {search_summary}"
                )
                return

            logger.info(f"Found {len(contacts)} contacts for: {search_summary}")

            # Create embed with results
            embed = discord.Embed(
                title="🔍 CRM Contact Search Results",
                description=f"Found {len(contacts)} contact(s) for: {search_summary}",
                color=0x0099FF,
            )

            # Create view with resume download buttons
            view = ResumeButtonView()

            for i, contact in enumerate(contacts):
                name = contact.get("name", "Unknown")
                additional_fields: list[tuple[str, str]] = []

                if skills_list:
                    contact_skills = self._format_requested_skills(skills_list, contact)
                    if contact_skills:
                        additional_fields.append(("🧠 Skills", contact_skills))

                contact_info = self._format_contact_card(
                    contact,
                    interaction=interaction,
                    additional_fields=additional_fields,
                )

                embed.add_field(name=f"👤 {name}", value=contact_info, inline=True)

                # Check for resume data directly from search results
                resume_ids = contact.get("resumeIds", [])
                resume_names = contact.get("resumeNames", {})

                if resume_ids and len(resume_ids) > 0:
                    # Use the last resume ID (newest uploaded)
                    last_resume_id = resume_ids[-1]
                    resume_name = resume_names.get(last_resume_id, f"{name}_resume")
                    logger.info(
                        f"Found resume for {name}: {resume_name} (ID: {last_resume_id})"
                    )
                    view.add_resume_button(name, last_resume_id)
                else:
                    logger.info(f"No resumes found for {name}")

            # Send embed with view only if there are buttons
            if view.children:
                await interaction.followup.send(embed=embed, view=view)
            else:
                await interaction.followup.send(embed=embed)

            self._audit_command(
                interaction=interaction,
                action="crm.search_members",
                result="success",
                metadata={
                    "query": query_value or None,
                    "skills": skills_list,
                    "contacts_found": len(contacts),
                    "resume_button_count": len(view.children),
                },
            )

        except EspoAPIError as e:
            logger.error(f"EspoCRM API error: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.search_members",
                result="error",
                metadata={"error": str(e)},
            )
            await interaction.followup.send(f"❌ CRM API error: {str(e)}")
        except Exception as e:
            logger.error(f"Unexpected error in CRM search: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.search_members",
                result="error",
                metadata={"error": str(e)},
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while searching the CRM."
            )

    @app_commands.command(
        name="assign-onboarder",
        description="Assign an onboarder to a CRM contact (Steering Committee+ only)",
    )
    @app_commands.describe(
        contact="Contact ID, 508 email, or name",
        onboarder="Onboarder 508 username or a Discord mention (mapped automatically)",
    )
    @require_role("Steering Committee")
    async def assign_onboarder(
        self, interaction: discord.Interaction, contact: str, onboarder: str
    ) -> None:
        """Assign an onboarder and set onboarding state to selected if still pending."""
        try:
            await interaction.response.defer(ephemeral=True)

            onboarder_username = await self._resolve_onboarder_username(
                interaction, onboarder
            )
            if not onboarder_username:
                self._audit_command(
                    interaction=interaction,
                    action="crm.assign_onboarder",
                    result="error",
                    metadata={"contact": contact, "onboarder": onboarder},
                )
                await interaction.followup.send(
                    "❌ Could not resolve a valid 508 onboarder username. "
                    "Use a 508 username directly or a linked Discord mention."
                )
                return

            contacts = await self._search_contact_for_linking(contact)
            if not contacts:
                self._audit_command(
                    interaction=interaction,
                    action="crm.assign_onboarder",
                    result="error",
                    metadata={"contact": contact, "onboarder": onboarder_username},
                )
                await interaction.followup.send(f"❌ No contact found for: `{contact}`")
                return

            if len(contacts) > 1:
                embed = discord.Embed(
                    title="⚠️ Multiple Contacts Found",
                    description=(
                        f"Found {len(contacts)} contacts for `{contact}`. "
                        "Please rerun with a more specific search term or an exact contact ID."
                    ),
                    color=0xFFA500,
                )
                for i, contact_record in enumerate(contacts[:5], 1):
                    contact_name = contact_record.get("name", "Unknown")
                    contact_info = self._format_contact_card(
                        contact_record,
                        interaction=interaction,
                        additional_fields=[
                            ("🆔 ID", str(contact_record.get("id", ""))),
                        ],
                    )
                    embed.add_field(
                        name=f"{i}. {contact_name}",
                        value=contact_info,
                        inline=False,
                    )
                await interaction.followup.send(embed=embed)
                self._audit_command(
                    interaction=interaction,
                    action="crm.assign_onboarder",
                    result="error",
                    metadata={
                        "contact": contact,
                        "onboarder": onboarder_username,
                        "contacts_found": len(contacts),
                    },
                )
                return

            target_contact = contacts[0]
            contact_id = target_contact.get("id")
            if not contact_id:
                self._audit_command(
                    interaction=interaction,
                    action="crm.assign_onboarder",
                    result="error",
                    metadata={"contact": contact, "onboarder": onboarder_username},
                )
                await interaction.followup.send("❌ Selected contact is missing an ID.")
                return

            full_contact = self.espo_api.request("GET", f"Contact/{contact_id}")
            onboarder_field = self._resolve_field_name(
                full_contact, candidates=ONBOARDER_FIELD_CANDIDATES
            )
            if not onboarder_field:
                self._audit_command(
                    interaction=interaction,
                    action="crm.assign_onboarder",
                    result="error",
                    metadata={
                        "contact_id": str(contact_id),
                        "onboarder": onboarder_username,
                        "reason": "missing_onboarder_field",
                    },
                    resource_type="crm_contact",
                    resource_id=str(contact_id),
                )
                await interaction.followup.send(
                    "❌ Could not locate a known onboarder field for this CRM contact."
                )
                return

            state_field = self._resolve_field_name(
                full_contact, candidates=ONBOARDING_STATUS_FIELD_CANDIDATES
            )
            current_state = self._normalize_onboarding_state(
                full_contact.get(state_field) if state_field else None
            )

            update_payload: dict[str, str] = {onboarder_field: onboarder_username}
            state_updated = False
            if state_field and current_state == "pending":
                update_payload[state_field] = "selected"
                state_updated = True

            self.espo_api.request("PUT", f"Contact/{contact_id}", update_payload)

            contact_name = full_contact.get("name", "Unknown")
            status_line = (
                "onboarding state set to `selected`"
                if state_updated
                else "onboarding state left unchanged"
            )
            await interaction.followup.send(
                f"✅ Assigned **{onboarder_username}** as onboarder for "
                f"**{contact_name}** (`{contact_id}`); {status_line}."
            )
            self._audit_command(
                interaction=interaction,
                action="crm.assign_onboarder",
                result="success",
                metadata={
                    "contact_id": str(contact_id),
                    "contact_name": contact_name,
                    "onboarder": onboarder_username,
                    "state_field": state_field,
                    "state_updated": state_updated,
                    "previous_state": current_state or None,
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )

        except EspoAPIError as e:
            logger.error(f"EspoCRM API error in assign_onboarder: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.assign_onboarder",
                result="error",
                metadata={"contact": contact, "onboarder": onboarder, "error": str(e)},
            )
            await interaction.followup.send(f"❌ CRM API error: {str(e)}")
        except Exception as e:
            logger.error(f"Unexpected error in assign_onboarder: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.assign_onboarder",
                result="error",
                metadata={"contact": contact, "onboarder": onboarder, "error": str(e)},
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while assigning the onboarder."
            )

    @app_commands.command(
        name="view-onboarding-queue",
        description="View contacts in onboarding queue (Steering Committee+ only)",
    )
    @require_role("Steering Committee")
    async def view_onboarding_queue(self, interaction: discord.Interaction) -> None:
        """Show contacts in onboarding queue (excluding onboarded, waitlist, rejected)."""
        try:
            await interaction.response.defer(ephemeral=True)

            response = self.espo_api.request(
                "GET",
                "Contact",
                {
                    "maxSize": ONBOARDING_QUEUE_MAX_SIZE,
                    "select": (
                        "id,name,emailAddress,cDiscordUsername,cDiscordUserID,"
                        "cOnboardingState,cOnboardingStatus,cOnboarding,"
                        "cOnboarder,cOnboardingCoordinator,cOnboardingUpdatedAt"
                    ),
                },
            )
            contacts = response.get("list", [])

            queue_entries: list[tuple[dict[str, Any], str]] = []
            for contact_record in contacts:
                state_field = self._resolve_field_name(
                    contact_record, candidates=ONBOARDING_STATUS_FIELD_CANDIDATES
                )
                status = (
                    self._normalize_onboarding_state(contact_record.get(state_field))
                    if state_field
                    else ""
                )
                if status in EXCLUDED_ONBOARDING_STATES:
                    continue
                queue_entries.append((contact_record, status))

            if not queue_entries:
                self._audit_command(
                    interaction=interaction,
                    action="crm.view_onboarding_queue",
                    result="success",
                    metadata={"count": 0},
                )
                await interaction.followup.send(
                    "✅ No contacts found in onboarding queue."
                )
                return

            queue_entries.sort(
                key=lambda item: (item[1] or "unknown", str(item[0].get("name", "")))
            )

            queue_rows = self._build_onboarding_queue_rows(interaction, queue_entries)
            view = OnboardingQueuePagerView(
                crm_cog=self,
                interaction=interaction,
                queue_rows=queue_rows,
                page_size=ONBOARDING_QUEUE_PAGE_SIZE,
            )
            embed = view._build_embed()

            self._audit_command(
                interaction=interaction,
                action="crm.view_onboarding_queue",
                result="success",
                metadata={
                    "count": len(queue_entries),
                    "output_format": "embed_paged",
                },
            )
            if view.total_pages > 1:
                message = await interaction.followup.send(
                    embed=embed, view=view, wait=True
                )
                view._set_message(message)
            else:
                await interaction.followup.send(embed=embed)

        except EspoAPIError as e:
            logger.error(f"EspoCRM API error in view_onboarding_queue: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.view_onboarding_queue",
                result="error",
                metadata={"error": str(e)},
            )
            await interaction.followup.send(f"❌ CRM API error: {str(e)}")
        except Exception as e:
            logger.error(f"Unexpected error in view_onboarding_queue: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.view_onboarding_queue",
                result="error",
                metadata={"error": str(e)},
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while loading onboarding queue."
            )

    @app_commands.command(
        name="crm-status", description="Check CRM API connection status"
    )
    async def crm_status(self, interaction: discord.Interaction) -> None:
        """Check if the CRM API is accessible."""
        try:
            await interaction.response.defer(ephemeral=True)

            # Try a simple API call to check connectivity
            response = self.espo_api.request("GET", "App/user")
            user_name = response.get("user", {}).get("name", "Unknown")

            embed = discord.Embed(
                title="✅ CRM Status",
                description="Connection to EspoCRM is working!",
                color=0x00FF00,
            )
            embed.add_field(name="Connected as", value=user_name, inline=True)
            embed.add_field(name="Base URL", value=settings.espo_base_url, inline=True)

            await interaction.followup.send(embed=embed)
            self._audit_command(
                interaction=interaction,
                action="crm.status",
                result="success",
                metadata={"connected_as": user_name},
            )

        except EspoAPIError as e:
            logger.error(f"EspoCRM API error: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.status",
                result="error",
                metadata={"error": str(e)},
            )
            embed = discord.Embed(
                title="❌ CRM Status",
                description=f"Failed to connect to EspoCRM: {str(e)}",
                color=0xFF0000,
            )
            await interaction.followup.send(embed=embed)
        except Exception as e:
            logger.error(f"Unexpected error in CRM status: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.status",
                result="error",
                metadata={"error": str(e)},
            )
            embed = discord.Embed(
                title="❌ CRM Status",
                description="An unexpected error occurred while checking CRM status.",
                color=0xFF0000,
            )
            await interaction.followup.send(embed=embed)

    @app_commands.command(
        name="get-resume", description="Download and send a contact's resume"
    )
    @app_commands.describe(
        query="Email address, 508 email (username, username@, or username@508.dev), or Discord username"
    )
    @require_role("Member")
    async def get_resume(self, interaction: discord.Interaction, query: str) -> None:
        """Download and send a contact's resume as a file attachment."""
        try:
            await interaction.response.defer(ephemeral=True)
            contacts = await self._search_contact_for_linking(
                query,
                max_size=1,
                select="id,name,emailAddress,c508Email,cDiscordUsername,resumeIds,resumeNames,resumeTypes",
                include_discord_username_search=True,
            )

            if not contacts:
                self._audit_command(
                    interaction=interaction,
                    action="crm.get_resume",
                    result="success",
                    metadata={"query": query, "contact_found": False},
                )
                await interaction.followup.send(f"❌ No contact found for: `{query}`")
                return

            contact = contacts[0]
            contact_name = contact.get("name", "Unknown")

            # Get resume data directly from search results
            resume_ids = contact.get("resumeIds", [])
            resume_names = contact.get("resumeNames", {})

            if not resume_ids or len(resume_ids) == 0:
                self._audit_command(
                    interaction=interaction,
                    action="crm.get_resume",
                    result="success",
                    metadata={
                        "query": query,
                        "contact_found": True,
                        "has_resume": False,
                    },
                )
                await interaction.followup.send(
                    f"❌ No resume found for {contact_name}"
                )
                return

            # Use the last resume (newest uploaded)
            resume_id = resume_ids[-1]
            resume_name = resume_names.get(resume_id, f"{contact_name}_resume")

            logger.info(
                f"Downloading resume for {contact_name}: {resume_name} (ID: {resume_id})"
            )

            # Use shared download method
            download_ok = await self._download_and_send_resume(
                interaction, contact_name, resume_id
            )
            self._audit_command(
                interaction=interaction,
                action="crm.get_resume",
                result="success" if download_ok else "error",
                metadata={
                    "query": query,
                    "contact_found": True,
                    "has_resume": True,
                    "download_ok": download_ok,
                },
                resource_type="crm_contact",
                resource_id=str(contact.get("id", "")),
            )

        except EspoAPIError as e:
            logger.error(f"EspoCRM API error in get_resume: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.get_resume",
                result="error",
                metadata={"query": query, "error": str(e)},
            )
            await interaction.followup.send(f"❌ CRM API error: {str(e)}")
        except Exception as e:
            logger.error(f"Unexpected error in get_resume: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.get_resume",
                result="error",
                metadata={"query": query, "error": str(e)},
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while fetching the resume."
            )

    def _is_hex_string(self, s: str) -> bool:
        """Check if string looks like a hex contact ID."""
        return len(s) >= 15 and all(c in "0123456789abcdefABCDEF" for c in s)

    def _build_contact_search_filters(self, search_term: str) -> list[dict[str, Any]]:
        """Build a shared list of CRM search filters for contact lookup."""
        normalized = search_term.strip()
        if not normalized:
            return []

        mention_user_id = self._extract_discord_id_from_mention(normalized)
        if mention_user_id:
            return [
                {
                    "type": "equals",
                    "attribute": "cDiscordUserID",
                    "value": mention_user_id,
                }
            ]

        search_filters: list[dict[str, Any]] = []
        has_space = " " in normalized
        has_at = "@" in normalized

        if has_at:
            local_part, _, domain = normalized.partition("@")
            if local_part and (not domain or domain.lower() in {"", "508", "508.dev"}):
                normalized = f"{local_part}@508.dev"

            search_filters.extend(
                [
                    {
                        "type": "equals",
                        "attribute": "emailAddress",
                        "value": normalized,
                    },
                    {"type": "equals", "attribute": "c508Email", "value": normalized},
                ]
            )
            return search_filters

        search_filters.append(
            {"type": "contains", "attribute": "name", "value": normalized}
        )
        if not has_space:
            search_filters.append(
                {
                    "type": "equals",
                    "attribute": "c508Email",
                    "value": f"{normalized}@508.dev",
                }
            )
        return search_filters

    async def _search_contact_for_linking(
        self,
        search_term: str,
        *,
        max_size: int | None = None,
        select: str = "id,name,emailAddress,c508Email,cDiscordUsername",
        include_discord_username_search: bool = False,
    ) -> list[dict[str, Any]]:
        """Search for contacts using multiple shared criteria."""
        # Check if it looks like a hex contact ID
        if self._is_hex_string(search_term):
            try:
                response = self.espo_api.request("GET", f"Contact/{search_term}")
                if response and response.get("id"):
                    return [response]
            except EspoAPIError:
                pass  # If direct ID lookup fails, fall through to regular search

        search_filters = self._build_contact_search_filters(search_term)
        if not search_filters:
            return []

        if include_discord_username_search and "@" not in search_term:
            search_filters.append(
                {
                    "type": "contains",
                    "attribute": "cDiscordUsername",
                    "value": search_term.strip(),
                }
            )

        has_at = "@" in search_term.strip()
        has_space = " " in search_term.strip()
        if max_size is None:
            max_size = 1 if has_space and not has_at else 10
            if self._extract_discord_id_from_mention(search_term.strip()):
                max_size = 1

        search_params = {
            "where": [{"type": "or", "value": search_filters}],
            "maxSize": max_size,
            "select": select,
        }

        response = self.espo_api.request("GET", "Contact", search_params)
        contacts: list[dict[str, Any]] = response.get("list", [])

        # Deduplicate contacts by ID to avoid showing duplicates
        seen_ids = set()
        deduplicated_contacts = []
        for contact in contacts:
            contact_id = contact.get("id")
            if contact_id and contact_id not in seen_ids:
                seen_ids.add(contact_id)
                deduplicated_contacts.append(contact)

        return deduplicated_contacts

    @staticmethod
    def _resume_file_extension(filename: str | None) -> str:
        if not filename or "." not in filename:
            return ""
        return "." + filename.rsplit(".", 1)[-1].lower()

    @staticmethod
    def _is_valid_resume_name_candidate(value: str) -> bool:
        normalized = value.strip()
        if len(normalized) < 2:
            return False
        if not any(char.isalpha() for char in normalized):
            return False
        normalized_token = normalize_resume_name_token(normalized)
        if is_reserved_resume_name_token(normalized):
            return False
        if normalized.endswith(":") and len(normalized_token.split()) <= 3:
            return False
        return True

    def _extract_resume_text(
        self,
        file_content: bytes,
        *,
        filename: str | None,
    ) -> str:
        extension = self._resume_file_extension(filename)
        extracted_text = ""

        try:
            if extension == ".pdf":
                from pdfminer.high_level import extract_text as extract_pdf_text

                extracted_text = extract_pdf_text(io.BytesIO(file_content)).strip()
            elif extension == ".docx":
                from docx import Document

                document = Document(io.BytesIO(file_content))
                chunks: list[str] = []
                for paragraph in document.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        chunks.append(text)
                for table in document.tables:
                    for row in table.rows:
                        row_cells = [
                            cell.text.strip() for cell in row.cells if cell.text.strip()
                        ]
                        if row_cells:
                            chunks.append(" | ".join(row_cells))
                extracted_text = "\n".join(chunks).strip()
            elif extension == ".doc":
                extracted_text = file_content.decode("utf-8", errors="ignore")
                extracted_text = re.sub(r"[^\x20-\x7E\n\r\t]", " ", extracted_text)
                extracted_text = re.sub(r"\s+", " ", extracted_text).strip()
            else:
                extracted_text = file_content.decode("utf-8", errors="ignore").strip()
        except Exception as exc:
            logger.warning(
                "Failed to extract resume text filename=%s extension=%s error=%s",
                filename,
                extension,
                exc,
            )

        if extracted_text:
            return extracted_text
        return file_content.decode("utf-8", errors="ignore")

    def _extract_resume_contact_hints(
        self,
        file_content: bytes,
        *,
        filename: str | None = None,
    ) -> dict[str, Any]:
        """Extract contact-identifying signals and shared resume fields from bytes."""
        profile = self._extract_resume_profile(file_content, filename=filename)
        emails: list[str] = []
        if profile.email:
            emails.append(profile.email)
        for raw_email in getattr(profile, "additional_emails", []) or []:
            if raw_email and raw_email not in emails:
                emails.append(raw_email)
        return {
            "emails": emails,
            "github_usernames": [profile.github_username]
            if profile.github_username
            else [],
            "linkedin_urls": [profile.linkedin_url] if profile.linkedin_url else [],
            "phone": profile.phone,
            "name": profile.name,
            "address_country": profile.address_country,
            "timezone": profile.timezone,
            "address_city": profile.address_city,
            "description": profile.description,
            "primary_roles": profile.primary_roles,
            "seniority_level": profile.seniority_level,
            "skills": profile.skills,
            "availability": profile.availability,
            "rate_range": profile.rate_range,
            "referred_by": profile.referred_by,
        }

    async def _extract_resume_contact_hints_async(
        self,
        file_content: bytes,
        *,
        filename: str | None = None,
    ) -> dict[str, Any]:
        """Extract contact hints without blocking the event loop."""
        return await asyncio.to_thread(
            self._extract_resume_contact_hints,
            file_content,
            filename=filename,
        )

    def _extract_resume_profile(
        self,
        file_content: bytes,
        *,
        filename: str | None = None,
    ) -> Any:
        """Extract resume profile fields and cache per-file-content results."""
        cache = self._resume_profile_cache
        cache_key = (hash(file_content), self._resume_file_extension(filename))
        if cache and cache[0] == cache_key:
            return cache[1]

        text = self._extract_resume_text(file_content, filename=filename)
        profile = self.resume_extractor.extract(text)
        self._resume_profile_cache = (cache_key, profile)
        return profile

    def _extract_resume_name_fallback(
        self,
        file_content: bytes,
        *,
        filename: str | None = None,
    ) -> str:
        """Simple name heuristic fallback when extraction did not return a name."""
        text = self._extract_resume_text(file_content, filename=filename)
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        for line in lines[:40]:
            candidate = line.strip()
            if not candidate:
                continue
            if len(candidate) < 2:
                continue
            if "@" in candidate or "http" in candidate.lower():
                continue
            if not self._is_valid_resume_name_candidate(candidate):
                continue
            if len(candidate) <= 70:
                return candidate
        return "Unknown Contact"

    def _format_inferred_attempts(self, attempts: list[dict[str, Any]] | None) -> str:
        if not attempts:
            return ""

        formatted: list[str] = []
        for attempt in attempts:
            if not isinstance(attempt, dict):
                continue
            method = str(attempt.get("method", "")).strip()
            value = str(attempt.get("value", "")).strip()
            if not method or not value:
                continue
            formatted.append(f"{method}: `{value}`")

        return ", ".join(formatted)

    @staticmethod
    def _normalize_timezone(value: Any) -> str | None:
        if not isinstance(value, str):
            return None

        raw = value.strip().replace(" ", "")
        if not raw:
            return None

        utc_pattern = re.search(
            r"(?i)\b(?:utc|gmt)\s*([+-]\d{1,2}(?:[:.]?[0-5]?\d)?)\b", raw
        )
        if utc_pattern:
            raw = utc_pattern.group(1)
        if raw.lower() in {"utc", "gmt"}:
            return "UTC+00:00"

        if raw[0] not in {"+", "-"}:
            return None
        match = re.match(r"([+-])(\d{1,2})(?::?([0-5]?\d))?$", raw)
        if not match:
            return None

        sign = match.group(1)
        try:
            hours = int(match.group(2))
        except Exception:
            return None
        if not 0 <= hours <= 14:
            return None

        minutes = match.group(3)
        if minutes is None:
            minutes_value = 0
        else:
            try:
                minutes_value = int(minutes)
            except Exception:
                return None
            if minutes_value > 59:
                return None

        return f"UTC{sign}{hours:02d}:{minutes_value:02d}"

    def _build_inference_lookup_summary(
        self,
        *,
        file_content: bytes,
        attempts: list[dict[str, Any]] | None,
        filename: str | None = None,
    ) -> str:
        """Build a user-facing description of resume-derived lookup values."""
        attempts_text = self._format_inferred_attempts(attempts)
        if attempts_text:
            return f"\nTried contact lookups: {attempts_text}"

        hints_raw = self._extract_resume_contact_hints(file_content, filename=filename)
        if isinstance(hints_raw, dict):
            hints: dict[str, Any] = hints_raw
        else:
            hints = {}

        def _to_values(raw_values: Any) -> list[str]:
            values: list[str] = []
            if not isinstance(raw_values, list):
                return values
            for item in raw_values:
                if not isinstance(item, str):
                    continue
                normalized = item.strip()
                if normalized and normalized not in values:
                    values.append(normalized)
            return values

        email_values = _to_values(hints.get("emails"))
        github_usernames = _to_values(hints.get("github_usernames"))
        linkedin_urls = _to_values(hints.get("linkedin_urls"))

        summary_parts: list[str] = []
        if email_values:
            summary_parts.append(
                "emails: " + ", ".join(f"`{value}`" for value in email_values)
            )
        if github_usernames:
            summary_parts.append(
                "github usernames: "
                + ", ".join(f"`{value}`" for value in github_usernames)
            )
        if linkedin_urls:
            summary_parts.append(
                "linkedin URLs: " + ", ".join(f"`{value}`" for value in linkedin_urls)
            )

        if not summary_parts:
            return ""

        return "\nParsed resume identifiers: " + "; ".join(summary_parts)

    async def _build_inference_lookup_summary_async(
        self,
        *,
        file_content: bytes,
        attempts: list[dict[str, Any]] | None,
        filename: str | None = None,
    ) -> str:
        """Build lookup summary without blocking the event loop."""
        return await asyncio.to_thread(
            self._build_inference_lookup_summary,
            file_content=file_content,
            attempts=attempts,
            filename=filename,
        )

    def _build_resume_parsed_identity_summary(
        self, file_content: bytes, *, filename: str | None = None
    ) -> str:
        """Build a short display summary of parsed contact identity fields."""
        hints = self._extract_resume_contact_hints(file_content, filename=filename)
        parsed_name = str(hints.get("name") or "").strip()
        if not self._is_valid_resume_name_candidate(parsed_name):
            parsed_name = self._extract_resume_name_fallback(
                file_content, filename=filename
            )

        emails = hints.get("emails", [])
        if not isinstance(emails, list):
            emails = []
        primary_email = "No email parsed"
        if emails:
            raw_email = str(emails[0]).strip()
            if raw_email:
                primary_email = raw_email

        return (
            f"\nParsed contact details: name=`{parsed_name}`, email=`{primary_email}`"
        )

    async def _build_resume_parsed_identity_summary_async(
        self, file_content: bytes, *, filename: str | None = None
    ) -> str:
        """Build parsed identity summary without blocking the event loop."""
        return await asyncio.to_thread(
            self._build_resume_parsed_identity_summary,
            file_content,
            filename=filename,
        )

    def _extract_resume_name_hint(
        self, file_content: bytes, *, filename: str | None = None
    ) -> str:
        """Best-effort contact name extraction from resume text."""
        hints = self._extract_resume_contact_hints(file_content, filename=filename)
        extracted_name = str(hints.get("name") or "").strip()
        if self._is_valid_resume_name_candidate(extracted_name):
            return extracted_name
        return self._extract_resume_name_fallback(file_content, filename=filename)

    def _populate_name_fields(
        self, payload: dict[str, str], *, source_name: str
    ) -> None:
        """Populate firstName and lastName fields for CRM contact creation payloads."""
        first_name, last_name = self.resume_extractor.split_name(
            full_name=source_name,
            first_name_hint=str(payload.get("firstName", "")).strip() or None,
            last_name_hint=str(payload.get("lastName", "")).strip() or None,
        )
        payload["firstName"] = first_name
        payload["lastName"] = last_name

    def _build_resume_create_contact_payload(
        self,
        file_content: bytes,
        *,
        filename: str | None = None,
    ) -> dict[str, str]:
        """Build a minimal contact create payload from resume hints."""
        hints = self._extract_resume_contact_hints(file_content, filename=filename)
        name = self._extract_resume_name_hint(file_content, filename=filename)
        contact_name = name if name != "Unknown Contact" else "Resume Candidate"
        emails = hints.get("emails", [])
        github_usernames = hints.get("github_usernames", [])
        linkedin_urls = hints.get("linkedin_urls", [])
        skills = hints.get("skills", [])
        description = str(hints.get("description", "")).strip()
        if not isinstance(emails, list):
            emails = []
        if not isinstance(github_usernames, list):
            github_usernames = []
        if not isinstance(linkedin_urls, list):
            linkedin_urls = []
        if not isinstance(skills, list):
            skills = []

        payload: dict[str, Any] = {
            "type": "Prospect",
            "name": contact_name,
        }
        self._populate_name_fields(payload, source_name=contact_name)
        if emails:
            primary_email = emails[0]
            if primary_email.endswith("@508.dev"):
                payload["c508Email"] = primary_email
            else:
                payload["emailAddress"] = primary_email
        if github_usernames:
            payload["cGitHubUsername"] = github_usernames[0]
        if linkedin_urls:
            payload[self._configured_linkedin_field()] = linkedin_urls[0]
        phone = hints.get("phone")
        if isinstance(phone, str) and phone.strip():
            payload["phoneNumber"] = phone.strip()
        primary_roles = hints.get("primary_roles")
        if isinstance(primary_roles, list):
            normalized_roles = [
                str(role).strip()
                for role in primary_roles
                if isinstance(role, str) and role.strip()
            ]
            if normalized_roles:
                payload["cRoles"] = normalized_roles
        address_country = str(hints.get("address_country", "")).strip()
        if address_country:
            payload["addressCountry"] = address_country
        timezone = self._normalize_timezone(hints.get("timezone"))
        if timezone:
            payload["cTimezone"] = timezone
        address_city = str(hints.get("address_city", "")).strip()
        if address_city:
            payload["addressCity"] = address_city
        seniority = str(hints.get("seniority_level", "")).strip()
        if seniority:
            payload["cSeniority"] = seniority
        if description:
            payload["description"] = description
        if skills:
            normalized_skills = [
                str(item).strip() for item in skills if str(item).strip()
            ]
            if normalized_skills:
                payload["skills"] = ", ".join(normalized_skills)

        return payload

    async def _build_resume_create_contact_payload_async(
        self,
        file_content: bytes,
        *,
        filename: str | None = None,
    ) -> dict[str, str]:
        """Build contact payload without blocking the event loop."""
        return await asyncio.to_thread(
            self._build_resume_create_contact_payload,
            file_content,
            filename=filename,
        )

    def _discord_display_name(self, user: discord.Member) -> str:
        """Format Discord username for CRM fields."""
        username = str(getattr(user, "name", "")).strip()
        discriminator = getattr(user, "discriminator", "0")
        if (
            isinstance(discriminator, str)
            and discriminator.strip()
            and discriminator.strip() != "0"
            and discriminator.strip().isdigit()
        ):
            return f"{username}#{discriminator.strip()}"
        return username

    def _discord_link_fields(self, user: discord.Member) -> dict[str, str]:
        """Build CRM fields used to persist Discord linkage."""
        return {
            "cDiscordUsername": self._discord_display_name(user),
            "cDiscordUserID": str(user.id),
        }

    def _fallback_contact_name_for_discord_user(self, user: discord.Member) -> str:
        display_name = str(getattr(user, "display_name", "")).strip()
        if display_name:
            return display_name
        username = str(getattr(user, "name", "")).strip()
        if username:
            return username
        return f"Discord User {user.id}"

    def _build_contact_payload_for_link_user(
        self,
        *,
        user: discord.Member,
        file_content: bytes,
        filename: str | None = None,
    ) -> dict[str, str]:
        """Build contact payload from resume hints plus explicit Discord linkage."""
        payload = self._build_resume_create_contact_payload(
            file_content=file_content,
            filename=filename,
        )
        parsed_name = str(payload.get("name", "")).strip()
        if not parsed_name or parsed_name == "Resume Candidate":
            payload["name"] = self._fallback_contact_name_for_discord_user(user)
            payload.pop("firstName", None)
            payload.pop("lastName", None)
        self._populate_name_fields(
            payload, source_name=str(payload.get("name", "")).strip()
        )
        payload.update(self._discord_link_fields(user))
        return payload

    async def _build_contact_payload_for_link_user_async(
        self,
        *,
        user: discord.Member,
        file_content: bytes,
        filename: str | None = None,
    ) -> dict[str, str]:
        """Build link-user payload without blocking the event loop."""
        return await asyncio.to_thread(
            self._build_contact_payload_for_link_user,
            user=user,
            file_content=file_content,
            filename=filename,
        )

    async def _search_contacts_by_field(
        self, *, field: str, value: str, max_size: int = 10
    ) -> list[dict[str, Any]]:
        """Search contacts using an exact field equals match."""
        select_fields = [
            "id",
            "name",
            "emailAddress",
            "c508Email",
            "cDiscordUsername",
            "cGitHubUsername",
        ]
        if field not in select_fields:
            select_fields.append(field)
        search_params = {
            "where": [{"type": "equals", "attribute": field, "value": value}],
            "maxSize": max_size,
            "select": ",".join(select_fields),
        }

        response = self.espo_api.request("GET", "Contact", search_params)
        contacts: list[dict[str, Any]] = response.get("list", [])

        deduplicated_contacts: list[dict[str, Any]] = []
        seen_ids = set()
        for contact in contacts:
            contact_id = contact.get("id")
            if contact_id and contact_id not in seen_ids:
                seen_ids.add(contact_id)
                deduplicated_contacts.append(contact)

        return deduplicated_contacts

    async def _infer_contact_from_resume(
        self, file_content: bytes, *, filename: str | None = None
    ) -> tuple[dict[str, Any] | None, dict[str, Any] | None]:
        """Infer target contact from resume identifiers."""
        hints = await self._extract_resume_contact_hints_async(
            file_content,
            filename=filename,
        )
        attempts: list[dict[str, Any]] = []
        emails = hints.get("emails", [])
        if not isinstance(emails, list):
            emails = []
        for email in emails:
            attempts.append({"method": "email", "value": email})
            contacts = await self._search_contact_for_linking(email)
            if len(contacts) == 1:
                return contacts[0], {
                    "method": "email",
                    "value": email,
                    "attempts": attempts,
                }
            if len(contacts) > 1:
                return None, {
                    "method": "email",
                    "value": email,
                    "reason": "multiple_matches",
                    "attempts": attempts,
                }

        github_usernames = hints.get("github_usernames", [])
        if not isinstance(github_usernames, list):
            github_usernames = []
        for github_username in github_usernames:
            attempts.append({"method": "github", "value": github_username})
            contacts = await self._search_contacts_by_field(
                field="cGitHubUsername", value=github_username
            )
            if len(contacts) == 1:
                return contacts[0], {
                    "method": "github",
                    "value": github_username,
                    "attempts": attempts,
                }
            if len(contacts) > 1:
                return None, {
                    "method": "github",
                    "value": github_username,
                    "reason": "multiple_matches",
                    "attempts": attempts,
                }

        linkedin_urls = hints.get("linkedin_urls", [])
        if not isinstance(linkedin_urls, list):
            linkedin_urls = []
        for linkedin_url in linkedin_urls:
            attempts.append({"method": "linkedin", "value": linkedin_url})
            contacts = await self._search_contacts_by_field(
                field=self._configured_linkedin_field(), value=linkedin_url
            )
            if len(contacts) == 1:
                return contacts[0], {
                    "method": "linkedin",
                    "value": linkedin_url,
                    "attempts": attempts,
                }
            if len(contacts) > 1:
                return None, {
                    "method": "linkedin",
                    "value": linkedin_url,
                    "reason": "multiple_matches",
                    "attempts": attempts,
                }

        return None, {"reason": "no_matching_contact", "attempts": attempts}

    async def _perform_discord_linking(
        self,
        interaction: discord.Interaction,
        user: discord.Member,
        contact: dict[str, Any],
    ) -> bool:
        """Shared method to perform Discord user linking to a contact."""
        try:
            contact_id = contact.get("id")
            contact_name = contact.get("name", "Unknown")

            # Prepare the Discord username for storage (without ID) and display
            discord_display = self._discord_display_name(user)

            # Update the contact's Discord username and user ID
            update_data = self._discord_link_fields(user)

            update_response = self.espo_api.request(
                "PUT", f"Contact/{contact_id}", update_data
            )

            if update_response:
                # Create success embed
                embed = discord.Embed(
                    title="✅ Discord User Linked",
                    description="Successfully linked Discord user to CRM contact (updated username and user ID)",
                    color=0x00FF00,
                )
                embed.add_field(
                    name="👤 Contact", value=f"{contact_name}", inline=False
                )
                embed.add_field(
                    name="📧 Email",
                    value=f"{contact.get('c508Email') or contact.get('emailAddress', 'N/A')}",
                    inline=True,
                )
                embed.add_field(
                    name="💬 Discord User",
                    value=f"{user.mention} ({discord_display})",
                    inline=True,
                )
                # Add CRM link
                if contact_id:
                    profile_url = f"{self.base_url}/#Contact/view/{contact_id}"
                    embed.add_field(
                        name="🔗 CRM Profile",
                        value=f"[View in CRM]({profile_url})",
                        inline=True,
                    )

                await interaction.followup.send(embed=embed)

                logger.info(
                    f"Discord user {user.name} (ID: {user.id}) linked to CRM contact "
                    f"{contact_name} (ID: {contact_id}) by {interaction.user.name}"
                )
                self._audit_command(
                    interaction=interaction,
                    action="crm.link_discord_user.execute",
                    result="success",
                    metadata={
                        "linked_user_id": str(user.id),
                        "linked_username": user.name,
                        "contact_name": contact_name,
                    },
                    resource_type="crm_contact",
                    resource_id=str(contact_id),
                )
                return True
            else:
                self._audit_command(
                    interaction=interaction,
                    action="crm.link_discord_user.execute",
                    result="error",
                    metadata={
                        "linked_user_id": str(user.id),
                        "contact_name": contact_name,
                        "error": "crm_update_failed",
                    },
                    resource_type="crm_contact",
                    resource_id=str(contact_id),
                )
                await interaction.followup.send(
                    "❌ Failed to update contact in CRM. Please try again."
                )
                return False

        except EspoAPIError as e:
            logger.error(f"EspoCRM API error in _perform_discord_linking: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.link_discord_user.execute",
                result="error",
                metadata={"linked_user_id": str(user.id), "error": str(e)},
            )
            await interaction.followup.send(f"❌ CRM API error: {str(e)}")
            return False
        except Exception as e:
            logger.error(f"Unexpected error in _perform_discord_linking: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.link_discord_user.execute",
                result="error",
                metadata={"linked_user_id": str(user.id), "error": str(e)},
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while linking the user."
            )
            return False

    async def _show_contact_choices(
        self,
        interaction: discord.Interaction,
        user: discord.Member,
        search_term: str,
        contacts: list[dict[str, Any]],
    ) -> None:
        """Show contact choices when multiple results found."""
        embed = discord.Embed(
            title="🔍 Multiple Contacts Found",
            description=f"Found {len(contacts)} contacts for `{search_term}`. Click a button below to link the Discord user.",
            color=0xFFA500,
        )

        # Create view with contact selection buttons
        view = ContactSelectionView(user, search_term)

        for i, contact in enumerate(contacts[:5], 1):  # Show max 5
            name = contact.get("name", "Unknown")
            email = contact.get("emailAddress", "No email")
            email_508 = contact.get("c508Email", "No 508 email")
            contact_id = contact.get("id", "")

            contact_info = (
                f"📧 {email}\n🏢 508 Email: {email_508}\n🆔 ID: `{contact_id}`"
            )
            embed.add_field(name=f"{i}. {name}", value=contact_info, inline=True)

            # Add button for this contact
            view.add_contact_button(contact)

        embed.add_field(
            name="💡 Tip",
            value="Click the button for the contact you want to link, or use the contact ID for exact matching.",
            inline=False,
        )

        await interaction.followup.send(embed=embed, view=view)

    def _normalize_508_username(self, value: str | None) -> str | None:
        """Normalize a 508 username candidate."""
        if not value:
            return None

        normalized = value.strip()
        if not normalized:
            return None

        normalized = normalized.lstrip("@").strip()
        normalized = " ".join(normalized.split())
        if not normalized:
            return None

        if "@" in normalized:
            username, _, domain = normalized.partition("@")
            if not username:
                return None
            if domain.lower() in {"", "508", "508.dev"}:
                return username.lower()
            return username.lower()

        return normalized.lower()

    async def _resolve_verified_by(
        self, interaction: discord.Interaction, verified_by: str
    ) -> str | None:
        """Resolve verifier.

        If a value is provided:
            - If it looks like a Discord mention, try to resolve from CRM using the
              linked Discord user ID.
            - Otherwise normalize directly as a 508 username.

        If no value is provided, resolve the invoker from CRM first by Discord ID, then by
        Discord username.
        """
        if verified_by.strip():
            match = re.match(r"^<@!?(\d+)>$", verified_by.strip())
            if match and interaction.guild:
                member = interaction.guild.get_member(int(match.group(1)))
                if member:
                    contact = await self._find_contact_by_discord_id(str(member.id))
                    if contact:
                        candidate = self._normalize_508_username(
                            contact.get("c508Email") or ""
                        )
                        if candidate:
                            return candidate
                    return self._normalize_508_username(member.name)

            return self._normalize_508_username(verified_by)

        return await self._resolve_verified_by_from_interaction_user(interaction)

    async def _resolve_verified_by_from_interaction_user(
        self, interaction: discord.Interaction
    ) -> str | None:
        """Resolve verifier using the invoking Discord user."""
        user = interaction.user
        if not user:
            return None

        if getattr(user, "id", None):
            contact = await self._find_contact_by_discord_id(str(user.id))
            if contact:
                candidate = self._normalize_508_username(contact.get("c508Email") or "")
                if candidate:
                    return candidate

        discord_username = self._normalize_508_username(str(user.name))
        if discord_username:
            contact = await self._find_contact_by_discord_username(discord_username)
            if contact:
                candidate = self._normalize_508_username(contact.get("c508Email") or "")
                if candidate:
                    return candidate
            return discord_username

        return None

    async def _find_contact_by_discord_username(
        self, discord_username: str
    ) -> dict[str, Any] | None:
        """Find a contact by Discord username."""
        search_params = {
            "where": [
                {
                    "type": "equals",
                    "attribute": "cDiscordUsername",
                    "value": discord_username,
                }
            ],
            "maxSize": 1,
            "select": "id,name,emailAddress,c508Email,cDiscordUsername,cDiscordUserID",
        }

        response = self.espo_api.request("GET", "Contact", search_params)
        contacts = response.get("list", [])
        return contacts[0] if contacts else None

    async def _parse_verified_at(self, raw_verified_at: str | None) -> str:
        """Parse an ID verification date or default to today."""
        if not raw_verified_at or not raw_verified_at.strip():
            return date.today().isoformat()

        value = raw_verified_at.strip()
        normalized = " ".join(value.replace(",", " ").split())

        for fmt in (
            "%B %d %Y",
            "%b %d %Y",
            "%d %B %Y",
            "%d %b %Y",
        ):
            try:
                return datetime.strptime(normalized, fmt).date().isoformat()
            except ValueError:
                continue

        normalized_numeric = re.sub(r"[./\s]+", "-", value)
        for fmt in (
            "%Y-%m-%d",
            "%d-%m-%Y",
            "%d-%m-%y",
            "%m-%d-%Y",
            "%m-%d-%y",
        ):
            try:
                return datetime.strptime(normalized_numeric, fmt).date().isoformat()
            except ValueError:
                continue

        raise ValueError(f"Invalid verified_at format: '{raw_verified_at}'.")

    async def _search_contacts_for_mark_id_verification(
        self, search_term: str
    ) -> list[dict[str, Any]]:
        """Search for contacts for ID verification."""
        contacts = await self._search_contact_for_linking(search_term)
        return contacts

    async def _mark_id_verified_for_contact(
        self,
        interaction: discord.Interaction,
        contact: dict[str, Any],
        verified_by: str,
        verified_at: str,
        id_type: str | None,
        allow_overwrite: bool = False,
    ) -> bool:
        """Persist ID verification metadata to CRM."""
        contact_id = contact.get("id")
        contact_name = contact.get("name", "Unknown")

        if not contact_id:
            self._audit_command(
                interaction=interaction,
                action="crm.mark_id_verified",
                result="error",
                metadata={
                    "verified_by": verified_by,
                    "verified_at": verified_at,
                    "error": "contact_id_missing",
                },
            )
            await interaction.followup.send("❌ Contact ID not found.")
            return False

        try:
            current_contact = self.espo_api.request("GET", f"Contact/{contact_id}")
        except EspoAPIError as exc:
            logger.error(
                f"Failed to fetch contact before marking ID verification: {exc}"
            )
            self._audit_command(
                interaction=interaction,
                action="crm.mark_id_verified",
                result="error",
                metadata={
                    "contact_id": str(contact_id),
                    "verified_by": verified_by,
                    "verified_at": verified_at,
                    "error": "contact_lookup_failed",
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await interaction.followup.send(
                "❌ Failed to load current verification data from CRM."
            )
            return False

        existing_verified_by = str(
            current_contact.get(ID_VERIFIED_BY_FIELD, "") or ""
        ).strip()
        existing_verified_at = str(
            current_contact.get(ID_VERIFIED_AT_FIELD, "") or ""
        ).strip()
        normalized_verified_by = verified_by.strip()
        normalized_verified_at = verified_at.strip()

        verified_by_conflict = (
            bool(existing_verified_by)
            and existing_verified_by != normalized_verified_by
        )
        verified_at_conflict = (
            bool(existing_verified_at)
            and existing_verified_at != normalized_verified_at
        )
        needs_confirmation = verified_by_conflict or verified_at_conflict

        if needs_confirmation and not allow_overwrite:
            self._audit_command(
                interaction=interaction,
                action="crm.mark_id_verified",
                result="denied",
                metadata={
                    "contact_id": str(contact_id),
                    "contact_name": contact_name,
                    "verified_by": verified_by,
                    "verified_at": verified_at,
                    "existing_verified_by": existing_verified_by,
                    "existing_verified_at": existing_verified_at,
                    "reason": "overwrite_confirmation_needed",
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            confirm_view = MarkIdVerifiedOverwriteConfirmationView(
                crm_cog=self,
                interaction=interaction,
                contact=current_contact,
                verified_by=normalized_verified_by,
                verified_at=normalized_verified_at,
                id_type=id_type,
            )
            await interaction.followup.send(
                (
                    "⚠️ This contact is already ID verified.\n"
                    f"- Current verifier: {existing_verified_by}\n"
                    f"- Current date: {existing_verified_at}\n"
                    f"- New verifier: {normalized_verified_by}\n"
                    f"- New date: {normalized_verified_at}\n\n"
                    "Select **Overwrite** only if you want to replace existing values."
                ),
                view=confirm_view,
            )
            return False

        payload = {
            ID_VERIFIED_AT_FIELD: verified_at,
            ID_VERIFIED_BY_FIELD: verified_by,
        }
        if id_type:
            payload[ID_VERIFIED_TYPE_FIELD] = id_type

        try:
            update_response = self.espo_api.request(
                "PUT", f"Contact/{contact_id}", payload
            )
            if update_response:
                embed = discord.Embed(
                    title="✅ ID Verified",
                    description=f"Marked **{contact_name}** as ID verified.",
                    color=0x00FF00,
                )
                embed.add_field(name="📅 Verified at", value=verified_at, inline=True)
                embed.add_field(name="✅ Verified by", value=verified_by, inline=True)
                if id_type:
                    embed.add_field(name="🆔 ID type", value=id_type, inline=True)
                profile_url = f"{self.base_url}/#Contact/view/{contact_id}"
                embed.add_field(
                    name="🔗 CRM Profile",
                    value=f"[View in CRM]({profile_url})",
                    inline=True,
                )
                await interaction.followup.send(embed=embed)

                self._audit_command(
                    interaction=interaction,
                    action="crm.mark_id_verified",
                    result="success",
                    metadata={
                        "contact_id": str(contact_id),
                        "verified_by": verified_by,
                        "verified_at": verified_at,
                    },
                    resource_type="crm_contact",
                    resource_id=str(contact_id),
                )
                return True

            self._audit_command(
                interaction=interaction,
                action="crm.mark_id_verified",
                result="error",
                metadata={
                    "contact_id": str(contact_id),
                    "verified_by": verified_by,
                    "verified_at": verified_at,
                    "error": "crm_update_failed",
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await interaction.followup.send(
                "❌ Failed to update contact in CRM. Please try again."
            )
            return False

        except EspoAPIError as exc:
            logger.error(f"Failed to update contact ID verification: {exc}")
            self._audit_command(
                interaction=interaction,
                action="crm.mark_id_verified",
                result="error",
                metadata={
                    "contact_id": str(contact_id),
                    "verified_by": verified_by,
                    "verified_at": verified_at,
                    "error": str(exc),
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await interaction.followup.send(f"❌ CRM API error: {str(exc)}")
            return False
        except Exception as exc:
            logger.error(f"Unexpected error marking ID verification: {exc}")
            self._audit_command(
                interaction=interaction,
                action="crm.mark_id_verified",
                result="error",
                metadata={
                    "contact_id": str(contact_id),
                    "verified_by": verified_by,
                    "verified_at": verified_at,
                    "error": str(exc),
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while marking ID verification."
            )
            return False

    async def _show_mark_id_verified_contact_choices(
        self,
        interaction: discord.Interaction,
        search_term: str,
        contacts: list[dict[str, Any]],
        verified_by: str,
        verified_at: str,
        id_type: str | None,
    ) -> None:
        """Show contact choices when multiple candidates are found."""
        embed = discord.Embed(
            title="🔍 Multiple Contacts Found",
            description=(
                f"Found {len(contacts)} contacts for `{search_term}`. "
                "Select the correct person to mark as ID verified."
            ),
            color=0xFFA500,
        )

        view = MarkIdVerifiedSelectionView(
            crm_cog=self,
            requester_id=interaction.user.id,
            verified_by=verified_by,
            verified_at=verified_at,
            id_type=id_type,
        )

        for i, contact in enumerate(contacts[:5], 1):
            name = contact.get("name", "Unknown")
            email = contact.get("emailAddress", "No email")
            email_508 = contact.get("c508Email", "No 508 email")
            contact_id = contact.get("id", "")
            contact_info = (
                f"📧 {email}\n🏢 508 Email: {email_508}\n🆔 ID: `{contact_id}`"
            )
            embed.add_field(name=f"{i}. {name}", value=contact_info, inline=True)
            view.add_contact_button(contact)

        embed.add_field(
            name="💡 Tip",
            value="Select the contact button to continue, or rerun with a more specific term.",
            inline=False,
        )

        await interaction.followup.send(embed=embed, view=view)

    @app_commands.command(
        name="mark-id-verified",
        description="Mark a contact as ID verified (Admin only).",
    )
    @app_commands.describe(
        search_term="Email, 508 username, or name.",
        verified_by="Verifier 508 username or @Discord mention.",
        id_type="Type of ID used for verification (e.g. passport, driver's license).",
        verified_at=(
            "Date verified (e.g. YYYY-MM-DD, DD/MM/YYYY, March 5, 2026). "
            "Defaults to today."
        ),
    )
    @require_role("Admin")
    async def mark_id_verified(
        self,
        interaction: discord.Interaction,
        search_term: str,
        verified_by: str,
        id_type: str | None = None,
        verified_at: str | None = None,
    ) -> None:
        """Mark a contact as ID verified and record verifier, ID type, and date."""
        try:
            await interaction.response.defer(ephemeral=True)

            resolved_verified_by = await self._resolve_verified_by(
                interaction, verified_by
            )
            if not resolved_verified_by:
                self._audit_command(
                    interaction=interaction,
                    action="crm.mark_id_verified",
                    result="denied",
                    metadata={
                        "search_term": search_term,
                        "verified_by": verified_by,
                        "verified_at": verified_at,
                        "id_type": id_type,
                        "reason": "verified_by_not_resolved",
                    },
                )
                await interaction.followup.send(
                    "❌ Unable to resolve verifier from `verified_by`."
                )
                return

            if id_type is not None and verified_at is None:
                try:
                    await self._parse_verified_at(id_type)
                except ValueError:
                    pass
                else:
                    verified_at = id_type
                    id_type = None

            resolved_verified_at = await self._parse_verified_at(verified_at)

            contacts = await self._search_contacts_for_mark_id_verification(search_term)

            if not contacts:
                self._audit_command(
                    interaction=interaction,
                    action="crm.mark_id_verified",
                    result="success",
                    metadata={
                        "search_term": search_term,
                        "verified_by": resolved_verified_by,
                        "verified_at": resolved_verified_at,
                        "id_type": id_type,
                        "contacts_found": 0,
                    },
                )
                await interaction.followup.send(
                    f"❌ No contact found for: `{search_term}`"
                )
                return

            if len(contacts) > 1:
                self._audit_command(
                    interaction=interaction,
                    action="crm.mark_id_verified",
                    result="success",
                    metadata={
                        "search_term": search_term,
                        "verified_by": resolved_verified_by,
                        "verified_at": resolved_verified_at,
                        "id_type": id_type,
                        "contacts_found": len(contacts),
                        "requires_selection": True,
                    },
                )
                await self._show_mark_id_verified_contact_choices(
                    interaction=interaction,
                    search_term=search_term,
                    contacts=contacts,
                    verified_by=resolved_verified_by,
                    verified_at=resolved_verified_at,
                    id_type=id_type,
                )
                return

            target_contact = contacts[0]
            self._audit_command(
                interaction=interaction,
                action="crm.mark_id_verified",
                result="success",
                metadata={
                    "search_term": search_term,
                    "verified_by": resolved_verified_by,
                    "verified_at": resolved_verified_at,
                    "id_type": id_type,
                    "contacts_found": 1,
                },
                resource_type="crm_contact",
                resource_id=str(target_contact.get("id")),
            )
            await self._mark_id_verified_for_contact(
                interaction=interaction,
                contact=target_contact,
                verified_by=resolved_verified_by,
                verified_at=resolved_verified_at,
                id_type=id_type,
            )
        except ValueError as exc:
            logger.error(f"Invalid verified_at value: {exc}")
            self._audit_command(
                interaction=interaction,
                action="crm.mark_id_verified",
                result="denied",
                metadata={"verified_at": verified_at, "reason": str(exc)},
            )
            await interaction.followup.send(f"❌ {exc}")
        except Exception as exc:
            logger.error(f"Unexpected error in mark_id_verified: {exc}")
            self._audit_command(
                interaction=interaction,
                action="crm.mark_id_verified",
                result="error",
                metadata={"search_term": search_term, "error": str(exc)},
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while marking ID verification."
            )

    @app_commands.command(
        name="link-discord-user",
        description="Link a Discord user to a CRM contact (Steering Committee+ only)",
    )
    @app_commands.describe(
        user="Discord user to link (mention them)",
        search_term="Email, 508 email, name, or contact ID to find the contact",
    )
    @require_role("Steering Committee")
    async def link_discord_user(
        self, interaction: discord.Interaction, user: discord.Member, search_term: str
    ) -> None:
        """Link a Discord user to a CRM contact by updating the contact's Discord username."""
        try:
            await interaction.response.defer(ephemeral=True)

            # Determine search strategy based on search_term format
            contacts = await self._search_contact_for_linking(search_term)

            if not contacts:
                self._audit_command(
                    interaction=interaction,
                    action="crm.link_discord_user",
                    result="success",
                    metadata={
                        "search_term": search_term,
                        "linked_user_id": str(user.id),
                        "contacts_found": 0,
                    },
                )
                await interaction.followup.send(
                    f"❌ No contact found for: `{search_term}`"
                )
                return

            # Handle multiple results - show choices
            if len(contacts) > 1:
                self._audit_command(
                    interaction=interaction,
                    action="crm.link_discord_user",
                    result="success",
                    metadata={
                        "search_term": search_term,
                        "linked_user_id": str(user.id),
                        "contacts_found": len(contacts),
                        "requires_selection": True,
                    },
                )
                await self._show_contact_choices(
                    interaction, user, search_term, contacts
                )
                return

            # Single result - proceed with linking
            contact = contacts[0]
            self._audit_command(
                interaction=interaction,
                action="crm.link_discord_user",
                result="success",
                metadata={
                    "search_term": search_term,
                    "linked_user_id": str(user.id),
                    "contacts_found": 1,
                    "requires_selection": False,
                },
                resource_type="crm_contact",
                resource_id=str(contact.get("id", "")),
            )
            await self._perform_discord_linking(interaction, user, contact)

        except EspoAPIError as e:
            logger.error(f"EspoCRM API error in link_discord_user: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.link_discord_user",
                result="error",
                metadata={
                    "search_term": search_term,
                    "linked_user_id": str(user.id),
                    "error": str(e),
                },
            )
            await interaction.followup.send(f"❌ CRM API error: {str(e)}")
        except Exception as e:
            logger.error(f"Unexpected error in link_discord_user: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.link_discord_user",
                result="error",
                metadata={
                    "search_term": search_term,
                    "linked_user_id": str(user.id),
                    "error": str(e),
                },
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while linking the Discord user."
            )

    @app_commands.command(
        name="unlinked-discord-users",
        description="List Discord users with Member role who aren't linked in CRM (Steering Committee+ only)",
    )
    @require_role("Steering Committee")
    async def unlinked_discord_users(self, interaction: discord.Interaction) -> None:
        """Find Discord users with Member role who don't have CRM links."""
        try:
            await interaction.response.defer(ephemeral=True)

            if not interaction.guild:
                self._audit_command(
                    interaction=interaction,
                    action="crm.unlinked_discord_users",
                    result="denied",
                    metadata={"reason": "not_in_guild"},
                )
                await interaction.followup.send(
                    "❌ This command can only be used in a server."
                )
                return

            # Get all contacts that have Discord user IDs set
            linked_user_ids = await self._get_linked_discord_user_ids()

            # Get all guild members with Member role
            member_role_users = []
            for member in interaction.guild.members:
                if not member.bot and hasattr(member, "roles"):
                    if check_user_roles_with_hierarchy(member.roles, ["Member"]):
                        member_role_users.append(member)

            # Find unlinked users
            unlinked_users = []
            for member in member_role_users:
                if str(member.id) not in linked_user_ids:
                    unlinked_users.append(member)

            if not unlinked_users:
                self._audit_command(
                    interaction=interaction,
                    action="crm.unlinked_discord_users",
                    result="success",
                    metadata={"unlinked_count": 0},
                )
                await interaction.followup.send(
                    "✅ **All Members Linked**\nAll Discord users with Member role are linked in the CRM!"
                )
                return

            # Send list of unlinked users
            await self._send_unlinked_users_list(interaction, unlinked_users)
            self._audit_command(
                interaction=interaction,
                action="crm.unlinked_discord_users",
                result="success",
                metadata={"unlinked_count": len(unlinked_users)},
            )

        except EspoAPIError as e:
            logger.error(f"EspoCRM API error in unlinked_discord_users: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.unlinked_discord_users",
                result="error",
                metadata={"error": str(e)},
            )
            await interaction.followup.send(f"❌ CRM API error: {str(e)}")
        except Exception as e:
            logger.error(f"Unexpected error in unlinked_discord_users: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.unlinked_discord_users",
                result="error",
                metadata={"error": str(e)},
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while checking unlinked users."
            )

    async def _get_linked_discord_user_ids(self) -> set[str]:
        """Get all Discord user IDs that are linked in the CRM."""
        search_params = {
            "where": [
                {
                    "type": "isNotNull",
                    "attribute": "cDiscordUserID",
                }
            ],
            "maxSize": 200,  # Adjust based on your needs
            "select": "cDiscordUserID",
        }

        response = self.espo_api.request("GET", "Contact", search_params)
        contacts = response.get("list", [])

        linked_ids = set()
        for contact in contacts:
            discord_id = contact.get("cDiscordUserID")
            if discord_id and discord_id != "No Discord":
                linked_ids.add(discord_id)

        return linked_ids

    async def _send_unlinked_users_list(
        self, interaction: discord.Interaction, unlinked_users: list[discord.Member]
    ) -> None:
        """Send list of unlinked users as simple text."""
        # Create simple text list with mentions and display names
        user_lines = []
        for member in unlinked_users:
            user_lines.append(f"{member.mention} ({member.display_name})")

        # Create message content
        header = f"🔗 **Unlinked Discord Users ({len(unlinked_users)})**\n"
        user_list = "\n".join(user_lines)
        footer = "\n\n💡 Use `/link-discord-user @user <email/name/id>` to link these users to CRM contacts."

        message = header + user_list + footer

        # Discord message limit is 2000 characters, split if needed
        if len(message) <= 2000:
            await interaction.followup.send(message)
        else:
            # Split into multiple messages if too long
            messages = []
            current_message = header

            for user_line in user_lines:
                test_message = current_message + user_line + "\n"
                if len(test_message + footer) <= 2000:
                    current_message = test_message
                else:
                    # Send current message and start new one
                    messages.append(current_message.rstrip())
                    current_message = user_line + "\n"

            # Add the last message with footer
            if current_message.strip():
                messages.append(current_message.rstrip() + footer)

            # Send all messages
            for message in messages:
                await interaction.followup.send(message)

    async def _find_contact_by_discord_id(
        self, discord_user_id: str
    ) -> dict[str, Any] | None:
        """Find a contact by Discord user ID."""
        search_params = {
            "where": [
                {
                    "type": "equals",
                    "attribute": "cDiscordUserID",
                    "value": discord_user_id,
                }
            ],
            "maxSize": 1,
            "select": "id,name,emailAddress,c508Email,cDiscordUsername,cGitHubUsername",
        }

        response = self.espo_api.request("GET", "Contact", search_params)
        contacts = response.get("list", [])
        return contacts[0] if contacts else None

    def _extract_discord_id_from_mention(self, value: str) -> str | None:
        """Extract Discord user ID from @mention syntax."""
        match = re.fullmatch(r"<@!?(\d+)>", value.strip())
        if not match:
            return None
        return match.group(1)

    def _parse_json_object_with_recovery(self, raw: str) -> dict[str, Any] | None:
        """Parse JSON object with lightweight recovery for common malformed payloads."""

        def _load_json_object(candidate: str) -> dict[str, Any] | None:
            try:
                parsed = json.loads(candidate)
            except Exception:
                return None
            if isinstance(parsed, dict):
                return parsed
            return None

        text = raw.strip()
        if not text:
            return None

        attempts: list[str] = []
        attempts.append(text)

        # Keep only the outer object if prefix/suffix noise exists.
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            attempts.append(text[start : end + 1])

        normalized_quotes = (
            text.replace("“", '"').replace("”", '"').replace("’", "'").replace("‘", "'")
        )
        attempts.append(normalized_quotes)

        attempts_with_trimmed_commas = [
            re.sub(r",\s*([}\]])", r"\1", candidate) for candidate in attempts
        ]
        attempts.extend(attempts_with_trimmed_commas)

        for candidate in attempts:
            parsed = _load_json_object(candidate)
            if parsed is not None:
                return parsed

        for candidate in attempts:
            try:
                parsed = ast.literal_eval(candidate)
            except Exception:
                continue
            if isinstance(parsed, dict):
                return parsed

        return None

    def _extract_contact_skills_for_view(
        self, contact: dict[str, Any]
    ) -> tuple[list[tuple[str, int | None]], str]:
        """Return display skills with source priority: structured attrs then multi-enum."""
        parsed_attrs = self._parse_contact_skill_attrs(contact.get("cSkillAttrs"))
        if parsed_attrs:
            ordered = sorted(
                parsed_attrs.items(),
                key=lambda item: (-item[1], item[0].casefold()),
            )
            return [(skill, strength) for skill, strength in ordered], "cSkillAttrs"

        raw_skills = contact.get("skills")
        if isinstance(raw_skills, list):
            skills = normalize_skill_list(
                [str(item) for item in raw_skills if str(item).strip()]
            )
        else:
            skills = normalize_skill_list(
                [
                    item.strip()
                    for item in str(raw_skills or "").split(",")
                    if item.strip()
                ]
            )
        return [(skill, None) for skill in skills], "skills"

    async def _search_contacts_for_view_skills(
        self, search_term: str
    ) -> list[dict[str, Any]]:
        """Resolve search term for view-skills lookup."""
        contacts = await self._search_contact_for_linking(search_term)
        return contacts

    async def _search_contacts_for_reprocess_resume(
        self, search_term: str
    ) -> list[dict[str, Any]]:
        """Resolve search term for resume reprocessing lookup."""
        mention_user_id = self._extract_discord_id_from_mention(search_term)
        if mention_user_id:
            by_discord_id = await self._find_contact_by_discord_id(mention_user_id)
            return [by_discord_id] if by_discord_id else []

        contacts = await self._search_contact_for_linking(search_term)
        if contacts:
            return contacts

        normalized_username = self._normalize_508_username(search_term)
        if normalized_username:
            by_discord_username = await self._find_contact_by_discord_username(
                normalized_username
            )
            if by_discord_username:
                return [by_discord_username]

        if (
            "@" not in search_term
            and " " not in search_term
            and not self._is_hex_string(search_term)
        ):
            contacts = await self._search_contact_for_linking(f"{search_term}@508.dev")
        return contacts

    async def _get_latest_resume_attachment_for_contact(
        self, contact_id: str
    ) -> tuple[str | None, str | None]:
        """Return the most recently attached resume ID and filename."""
        contact_data = self.espo_api.request("GET", f"Contact/{contact_id}")

        resume_ids = contact_data.get("resumeIds")
        if not isinstance(resume_ids, list) or not resume_ids:
            return None, None

        attachment_id = str(resume_ids[-1])
        filename = None
        resume_names = contact_data.get("resumeNames")
        if isinstance(resume_names, dict):
            filename_value = resume_names.get(attachment_id)
            if isinstance(filename_value, str) and filename_value.strip():
                filename = filename_value.strip()

        return attachment_id, filename

    @app_commands.command(
        name="view-skills",
        description="View CRM skills for yourself or a specific member",
    )
    @app_commands.describe(
        search_term="Optional: @mention, email, 508 username, name, or contact ID",
    )
    @require_role("Member")
    async def view_skills(
        self, interaction: discord.Interaction, search_term: str | None = None
    ) -> None:
        """View structured skills (with strengths) or fallback multi-enum skills."""
        try:
            await interaction.response.defer(ephemeral=True)

            query = (search_term or "").strip()
            target_scope = "other" if query else "self"

            target_contact: dict[str, Any] | None = None
            if not query:
                target_contact = await self._find_contact_by_discord_id(
                    str(interaction.user.id)
                )
                if not target_contact:
                    self._audit_command(
                        interaction=interaction,
                        action="crm.view_skills",
                        result="denied",
                        metadata={
                            "target_scope": "self",
                            "reason": "discord_not_linked",
                        },
                    )
                    await interaction.followup.send(
                        "❌ Your Discord account is not linked to a CRM contact. "
                        "Please ask a Steering Committee member to link your account first."
                    )
                    return
            else:
                contacts = await self._search_contacts_for_view_skills(query)
                if not contacts:
                    self._audit_command(
                        interaction=interaction,
                        action="crm.view_skills",
                        result="success",
                        metadata={
                            "search_term": query,
                            "target_scope": target_scope,
                            "contacts_found": 0,
                        },
                    )
                    await interaction.followup.send(
                        f"❌ No contact found for: `{query}`"
                    )
                    return

                if len(contacts) > 1:
                    lines: list[str] = []
                    for contact in contacts[:5]:
                        name = str(contact.get("name", "Unknown"))
                        contact_id = str(contact.get("id", ""))
                        lines.append(f"- **{name}** (`{contact_id}`)")
                    suffix = (
                        f"\n...and {len(contacts) - 5} more."
                        if len(contacts) > 5
                        else ""
                    )
                    self._audit_command(
                        interaction=interaction,
                        action="crm.view_skills",
                        result="success",
                        metadata={
                            "search_term": query,
                            "target_scope": target_scope,
                            "contacts_found": len(contacts),
                            "requires_selection": True,
                        },
                    )
                    await interaction.followup.send(
                        "⚠️ Multiple contacts found. Please refine your search:\n"
                        + "\n".join(lines)
                        + suffix
                    )
                    return

                target_contact = contacts[0]

            assert target_contact is not None
            contact_id = str(target_contact.get("id") or "").strip()
            if not contact_id:
                self._audit_command(
                    interaction=interaction,
                    action="crm.view_skills",
                    result="error",
                    metadata={
                        "search_term": query or None,
                        "error": "missing_contact_id",
                    },
                )
                await interaction.followup.send("❌ Contact ID not found.")
                return

            full_contact = self.espo_api.request("GET", f"Contact/{contact_id}")
            contact_name = str(
                full_contact.get("name") or target_contact.get("name") or "Unknown"
            )
            skills, source = self._extract_contact_skills_for_view(full_contact)

            if not skills:
                self._audit_command(
                    interaction=interaction,
                    action="crm.view_skills",
                    result="success",
                    metadata={
                        "search_term": query or None,
                        "target_scope": target_scope,
                        "skills_count": 0,
                        "source": source,
                    },
                    resource_type="crm_contact",
                    resource_id=contact_id,
                )
                await interaction.followup.send(
                    f"ℹ️ No skills found for **{contact_name}**."
                )
                return

            embed = discord.Embed(
                title="🛠️ CRM Skills",
                description=f"Skills for **{contact_name}**",
                color=0x0099FF,
            )
            skill_lines: list[str] = []
            for skill, strength in skills[:25]:
                if strength is None:
                    skill_lines.append(skill)
                else:
                    skill_lines.append(f"{skill} ({strength})")
            if len(skills) > 25:
                skill_lines.append(f"...and {len(skills) - 25} more.")
            embed.add_field(name="Skills", value=", ".join(skill_lines), inline=False)
            embed.add_field(
                name="🔗 CRM Profile",
                value=f"[View in CRM]({self.base_url}/#Contact/view/{contact_id})",
                inline=False,
            )

            await interaction.followup.send(embed=embed)
            self._audit_command(
                interaction=interaction,
                action="crm.view_skills",
                result="success",
                metadata={
                    "search_term": query or None,
                    "target_scope": target_scope,
                    "skills_count": len(skills),
                    "source": source,
                },
                resource_type="crm_contact",
                resource_id=contact_id,
            )
        except EspoAPIError as e:
            logger.error(f"EspoCRM API error in view_skills: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.view_skills",
                result="error",
                metadata={"search_term": search_term, "error": str(e)},
            )
            await interaction.followup.send(f"❌ CRM API error: {str(e)}")
        except Exception as e:
            logger.error(f"Unexpected error in view_skills: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.view_skills",
                result="error",
                metadata={"search_term": search_term, "error": str(e)},
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while fetching skills."
            )

    @app_commands.command(
        name="update-contact",
        description="Update CRM contact fields (github, linkedin, skills, rate range, and resume)",
    )
    @app_commands.describe(
        github="GitHub username to set",
        linkedin="LinkedIn profile URL to set",
        skills="Comma-separated skills; supports `skill:4` for strength",
        rate_range="Rate range text to set",
        resume="Resume file to upload and analyze",
        overwrite="Replace existing resumes instead of appending",
        search_term="Email, name, or contact ID (optional). Omit to update your own contact.",
    )
    async def update_contact(
        self,
        interaction: discord.Interaction,
        github: str | None = None,
        linkedin: str | None = None,
        skills: str | None = None,
        rate_range: str | None = None,
        resume: discord.Attachment | None = None,
        overwrite: bool = False,
        search_term: str | None = None,
    ) -> None:
        """Update CRM contact fields for yourself or another contact."""
        try:
            await interaction.response.defer(ephemeral=True)

            has_updates = any(
                bool(value) for value in (github, linkedin, skills, rate_range)
            )
            if not has_updates and resume is None:
                self._audit_command(
                    interaction=interaction,
                    action="crm.update_contact",
                    result="denied",
                    metadata={"reason": "no_update_fields"},
                )
                await interaction.followup.send(
                    "❌ Provide at least one of `github`, `linkedin`, `skills`, `rate_range`, or `resume`."
                )
                return

            if resume is not None:
                if not settings.api_shared_secret:
                    self._audit_command(
                        interaction=interaction,
                        action="crm.update_contact",
                        result="error",
                        metadata={
                            "filename": resume.filename,
                            "reason": "api_shared_secret_missing",
                        },
                    )
                    await interaction.followup.send(
                        "❌ API_SHARED_SECRET is not configured for backend API access."
                    )
                    return

                valid_extensions = {".pdf", ".doc", ".docx", ".txt"}
                file_extension = (
                    "." + resume.filename.split(".")[-1].lower()
                    if "." in resume.filename
                    else ""
                )
                if file_extension not in valid_extensions:
                    self._audit_command(
                        interaction=interaction,
                        action="crm.update_contact",
                        result="denied",
                        metadata={
                            "filename": resume.filename,
                            "reason": "invalid_file_type",
                        },
                    )
                    await interaction.followup.send(
                        f"❌ Invalid file type. Upload a PDF, DOC, DOCX, or TXT file.\n"
                        f"You uploaded: `{resume.filename}`"
                    )
                    return

                max_size = 10 * 1024 * 1024
                if resume.size > max_size:
                    self._audit_command(
                        interaction=interaction,
                        action="crm.update_contact",
                        result="denied",
                        metadata={
                            "filename": resume.filename,
                            "size_bytes": resume.size,
                            "reason": "file_too_large",
                        },
                    )
                    await interaction.followup.send(
                        f"❌ File too large. Maximum size is 10MB.\nYour file: {resume.size / (1024 * 1024):.1f}MB"
                    )
                    return

            is_steering = hasattr(
                interaction.user, "roles"
            ) and check_user_roles_with_hierarchy(
                interaction.user.roles, ["Steering Committee"]
            )
            if search_term and not is_steering:
                self._audit_command(
                    interaction=interaction,
                    action="crm.update_contact",
                    result="denied",
                    metadata={
                        "search_term": search_term,
                        "reason": "missing_required_role",
                    },
                )
                await interaction.followup.send(
                    "❌ You must have Steering Committee role or higher to update another contact."
                )
                return

            target_contact = None
            target_scope = "self"

            if search_term:
                contacts = await self._search_contact_for_linking(search_term)
                if not contacts:
                    self._audit_command(
                        interaction=interaction,
                        action="crm.update_contact",
                        result="success",
                        metadata={
                            "search_term": search_term,
                            "contact_found": False,
                            "target_scope": "other",
                        },
                    )
                    await interaction.followup.send(
                        f"❌ No contact found for: `{search_term}`"
                    )
                    return

                if len(contacts) > 1:
                    self._audit_command(
                        interaction=interaction,
                        action="crm.update_contact",
                        result="success",
                        metadata={
                            "search_term": search_term,
                            "contact_found": False,
                            "target_scope": "other",
                            "reason": "multiple_contacts",
                        },
                    )
                    await interaction.followup.send(
                        f"❌ Multiple contacts found for `{search_term}`. Please be more specific or use the contact ID."
                    )
                    return

                target_contact = contacts[0]
                target_scope = "other"
            else:
                target_contact = await self._find_contact_by_discord_id(
                    str(interaction.user.id)
                )
                if not target_contact:
                    self._audit_command(
                        interaction=interaction,
                        action="crm.update_contact",
                        result="denied",
                        metadata={
                            "target_scope": "self",
                            "reason": "discord_not_linked",
                        },
                    )
                    await interaction.followup.send(
                        "❌ Your Discord account is not linked to a CRM contact. "
                        "Please ask a Steering Committee member to link your account first."
                    )
                    return

            assert target_contact is not None
            contact_id = target_contact.get("id")
            if not contact_id:
                self._audit_command(
                    interaction=interaction,
                    action="crm.update_contact",
                    result="error",
                    metadata={
                        "search_term": search_term,
                        "error": "contact_id_missing",
                    },
                )
                await interaction.followup.send("❌ Contact ID not found.")
                return

            contact_name = target_contact.get("name", "Unknown")
            update_data: dict[str, str] = {}
            requested_updates: list[str] = []

            if github is not None:
                clean_github_username = github.strip().lstrip("@")
                if clean_github_username:
                    update_data["cGitHubUsername"] = clean_github_username
                    requested_updates.append("github")

            if linkedin is not None:
                clean_linkedin = linkedin.strip()
                if clean_linkedin:
                    update_data[self._configured_linkedin_field()] = clean_linkedin
                    requested_updates.append("linkedin")

            if rate_range is not None:
                clean_rate_range = rate_range.strip()
                if clean_rate_range:
                    update_data["rateRange"] = clean_rate_range
                    requested_updates.append("rate_range")

            if skills is not None:
                parsed_skills, requested_strengths, invalid_skills = (
                    self._parse_skill_updates(skills)
                )
                if invalid_skills:
                    invalid_message = ", ".join(f"`{item}`" for item in invalid_skills)
                    self._audit_command(
                        interaction=interaction,
                        action="crm.update_contact",
                        result="denied",
                        metadata={
                            "search_term": search_term,
                            "invalid_skills": invalid_skills,
                            "reason": "invalid_skill_format",
                        },
                    )
                    await interaction.followup.send(
                        f"❌ Invalid skill entries: {invalid_message}. "
                        "Use `skill` or `skill:1-5` (e.g. `go`, `python:4`)."
                    )
                    return

                if parsed_skills:
                    merged_skills, merged_attrs = self._merge_skill_update_payload(
                        target_contact, parsed_skills, requested_strengths
                    )
                    update_data["skills"] = merged_skills
                    update_data["cSkillAttrs"] = merged_attrs
                    requested_updates.append("skills")

            if not update_data and resume is None:
                self._audit_command(
                    interaction=interaction,
                    action="crm.update_contact",
                    result="denied",
                    metadata={
                        "search_term": search_term,
                        "reason": "no_effective_updates",
                    },
                )
                await interaction.followup.send(
                    "❌ No valid updatable fields were provided."
                )
                return

            if update_data:
                update_response = self.espo_api.request(
                    "PUT", f"Contact/{contact_id}", update_data
                )

                if update_response:
                    embed = discord.Embed(
                        title="✅ Contact Updated",
                        description="Successfully updated CRM contact fields.",
                        color=0x00FF00,
                    )
                    embed.add_field(
                        name="👤 Contact", value=f"{contact_name}", inline=False
                    )
                    embed.add_field(
                        name="📧 Email",
                        value=f"{target_contact.get('c508Email') or target_contact.get('emailAddress', 'N/A')}",
                        inline=True,
                    )
                    if "github" in requested_updates:
                        embed.add_field(
                            name="🐙 GitHub",
                            value=f"@{update_data['cGitHubUsername']}",
                            inline=True,
                        )
                    if "linkedin" in requested_updates:
                        linkedin_field = self._configured_linkedin_field()
                        embed.add_field(
                            name="🔗 LinkedIn",
                            value=update_data[linkedin_field],
                            inline=True,
                        )
                    if "skills" in requested_updates:
                        embed.add_field(
                            name="🧠 Skills", value=update_data["skills"], inline=False
                        )
                    if "rate_range" in requested_updates:
                        embed.add_field(
                            name="💵 Rate Range",
                            value=update_data["rateRange"],
                            inline=True,
                        )
                    embed.add_field(
                        name="🔎 Updated Fields",
                        value=", ".join(requested_updates),
                        inline=False,
                    )
                    profile_url = f"{self.base_url}/#Contact/view/{contact_id}"
                    embed.add_field(
                        name="🔗 CRM Profile",
                        value=f"[View in CRM]({profile_url})",
                        inline=True,
                    )
                    await interaction.followup.send(embed=embed)

                    logger.info(
                        f"Contact updated for {contact_name} (ID: {contact_id}) fields={requested_updates} by {interaction.user.name}"
                    )
                    self._audit_command(
                        interaction=interaction,
                        action="crm.update_contact",
                        result="success",
                        metadata={
                            "search_term": search_term,
                            "target_scope": target_scope,
                            "updated_fields": requested_updates,
                            "has_resume": resume is not None,
                        },
                        resource_type="crm_contact",
                        resource_id=str(contact_id),
                    )
                else:
                    self._audit_command(
                        interaction=interaction,
                        action="crm.update_contact",
                        result="error",
                        metadata={
                            "search_term": search_term,
                            "error": "crm_update_failed",
                        },
                        resource_type="crm_contact",
                        resource_id=str(contact_id),
                    )
                    await interaction.followup.send(
                        "❌ Failed to update contact in CRM. Please try again."
                    )
                    return

            if resume is not None:
                file_content = await resume.read()
                await self._upload_resume_attachment_to_contact(
                    interaction=interaction,
                    file_content=file_content,
                    filename=resume.filename,
                    file_size=resume.size,
                    contact=target_contact,
                    target_scope=target_scope,
                    search_term=search_term,
                    overwrite=overwrite,
                    link_user=None,
                    inferred_contact_meta=None,
                )

        except EspoAPIError as e:
            logger.error(f"EspoCRM API error in update_contact: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.update_contact",
                result="error",
                metadata={"search_term": search_term, "error": str(e)},
            )
            await interaction.followup.send(f"❌ CRM API error: {str(e)}")
        except Exception as e:
            logger.error(f"Unexpected error in update_contact: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.update_contact",
                result="error",
                metadata={"search_term": search_term, "error": str(e)},
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while updating the contact."
            )

    async def _check_existing_resume(
        self, contact_id: str, filename: str, filesize: int
    ) -> tuple[bool, str | None]:
        """Check if contact already has a resume with the same name and size."""
        try:
            # Get current contact data
            contact_data = self.espo_api.request("GET", f"Contact/{contact_id}")
            current_resume_ids = contact_data.get("resumeIds", [])

            # Check each existing resume
            for resume_id in current_resume_ids:
                try:
                    # Get attachment details
                    attachment_data = self.espo_api.request(
                        "GET", f"Attachment/{resume_id}"
                    )

                    # Compare filename and size
                    if (
                        attachment_data.get("name") == filename
                        and attachment_data.get("size") == filesize
                    ):
                        return True, resume_id

                except EspoAPIError:
                    # If we can't fetch attachment details, skip it
                    continue

            return False, None

        except EspoAPIError as e:
            logger.error(f"Failed to check existing resumes: {e}")
            return False, None

    async def _update_contact_resume(
        self, contact_id: str, attachment_id: str, overwrite: bool = False
    ) -> bool:
        """Update contact's resume field with the attachment ID."""
        try:
            # Get current contact data to preserve existing resume IDs
            contact_data = self.espo_api.request("GET", f"Contact/{contact_id}")
            current_resume_ids = contact_data.get("resumeIds", [])

            # Handle overwrite vs append
            if overwrite:
                # Replace all existing resumes with just this one
                new_resume_ids = [attachment_id]
            else:
                # Add new attachment ID to the end of resume IDs list
                if attachment_id not in current_resume_ids:
                    current_resume_ids.append(attachment_id)
                new_resume_ids = current_resume_ids

            # Update the contact with new resume IDs
            update_data = {"resumeIds": new_resume_ids}

            self.espo_api.request("PUT", f"Contact/{contact_id}", update_data)
            return True
        except EspoAPIError as e:
            logger.error(f"Failed to update contact resume: {e}")
            return False

    async def _upload_resume_attachment_to_contact(
        self,
        *,
        interaction: discord.Interaction,
        file_content: bytes,
        filename: str,
        file_size: int,
        contact: dict[str, Any],
        target_scope: str,
        search_term: str | None,
        overwrite: bool,
        link_user: discord.Member | None,
        inferred_contact_meta: dict[str, Any] | None,
    ) -> None:
        """Upload attachment and launch the worker preview for a given contact."""
        contact_id = contact.get("id")
        contact_name = contact.get("name", "Unknown")

        if not contact_id:
            self._audit_command(
                interaction=interaction,
                action="crm.upload_resume",
                result="error",
                metadata={
                    "search_term": search_term,
                    "filename": filename,
                    "error": "contact_id_missing",
                },
            )
            await interaction.followup.send("❌ Contact ID not found.")
            return

        try:
            attachment = self.espo_api.upload_file(
                file_content=file_content,
                filename=filename,
                related_type="Contact",
                related_id=contact_id,
                field="resume",
            )

            attachment_id = attachment.get("id")
            if not attachment_id:
                self._audit_command(
                    interaction=interaction,
                    action="crm.upload_resume",
                    result="error",
                    metadata={
                        "search_term": search_term,
                        "filename": filename,
                        "error": "attachment_id_missing",
                    },
                    resource_type="crm_contact",
                    resource_id=str(contact_id),
                )
                await interaction.followup.send("❌ Failed to upload file to CRM.")
                return

            if not await self._update_contact_resume(
                contact_id, attachment_id, overwrite
            ):
                self._audit_command(
                    interaction=interaction,
                    action="crm.upload_resume",
                    result="error",
                    metadata={
                        "search_term": search_term,
                        "filename": filename,
                        "attachment_id": attachment_id,
                        "overwrite": overwrite,
                        "target_scope": target_scope,
                        "reason": "resume_link_update_failed",
                    },
                    resource_type="crm_contact",
                    resource_id=str(contact_id),
                )
                await interaction.followup.send(
                    "⚠️ File uploaded, but failed to link in contact resume field."
                )
                return

            logger.info(
                "Resume uploaded for %s (contact_id=%s, attachment_id=%s) by %s",
                contact_name,
                contact_id,
                attachment_id,
                interaction.user.name,
            )
            success_metadata = {
                "search_term": search_term,
                "filename": filename,
                "size_bytes": file_size,
                "overwrite": overwrite,
                "target_scope": target_scope,
                "attachment_id": attachment_id,
                "stage": "uploaded_and_linked",
            }
            if inferred_contact_meta:
                for key in ("method", "value"):
                    value = inferred_contact_meta.get(key)
                    if value:
                        success_metadata[f"inferred_{key}"] = value

            self._audit_command(
                interaction=interaction,
                action="crm.upload_resume",
                result="success",
                metadata=success_metadata,
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await self._run_resume_extract_and_preview(
                interaction=interaction,
                contact_id=contact_id,
                contact_name=contact_name,
                attachment_id=attachment_id,
                filename=filename,
                link_member=link_user,
            )
        except EspoAPIError as e:
            logger.error(f"Failed to upload file to EspoCRM: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.upload_resume",
                result="error",
                metadata={
                    "search_term": search_term,
                    "filename": filename,
                    "error": str(e),
                },
                resource_type="crm_contact",
                resource_id=str(contact_id),
            )
            await interaction.followup.send(
                f"❌ Failed to upload file to CRM: {str(e)}"
            )

    @app_commands.command(
        name="upload-resume",
        description="Upload resume, extract profile fields, and preview CRM updates",
    )
    @app_commands.describe(
        file="Resume file to upload (PDF, DOC, DOCX, TXT)",
        search_term="Email, name, or contact ID (Steering Committee+ only). Omit to infer from resume.",
        overwrite="Replace existing resumes instead of appending",
        link_user="Discord user to link to this CRM contact (optional, Steering Committee+ for others)",
    )
    async def upload_resume(
        self,
        interaction: discord.Interaction,
        file: discord.Attachment,
        search_term: str | None = None,
        overwrite: bool = False,
        link_user: discord.Member | None = None,
    ) -> None:
        """Upload resume and run backend extraction to preview CRM updates."""
        try:
            await interaction.response.defer(ephemeral=True)

            if not settings.api_shared_secret:
                self._audit_command(
                    interaction=interaction,
                    action="crm.upload_resume",
                    result="error",
                    metadata={
                        "filename": file.filename,
                        "reason": "api_shared_secret_missing",
                    },
                )
                await interaction.followup.send(
                    "❌ API_SHARED_SECRET is not configured for backend API access."
                )
                return

            # Validate file type
            valid_extensions = {".pdf", ".doc", ".docx", ".txt"}
            file_extension = (
                "." + file.filename.split(".")[-1].lower()
                if "." in file.filename
                else ""
            )

            if file_extension not in valid_extensions:
                self._audit_command(
                    interaction=interaction,
                    action="crm.upload_resume",
                    result="error",
                    metadata={
                        "filename": file.filename,
                        "reason": "invalid_file_type",
                    },
                )
                await interaction.followup.send(
                    f"❌ Invalid file type. Please upload a PDF, DOC, DOCX, or TXT file.\nYou uploaded: `{file.filename}`"
                )
                return

            # Validate file size (10MB limit)
            max_size = 10 * 1024 * 1024  # 10MB in bytes
            if file.size > max_size:
                self._audit_command(
                    interaction=interaction,
                    action="crm.upload_resume",
                    result="error",
                    metadata={
                        "filename": file.filename,
                        "size_bytes": file.size,
                        "reason": "file_too_large",
                    },
                )
                await interaction.followup.send(
                    f"❌ File too large. Maximum size is 10MB.\nYour file: {file.size / (1024 * 1024):.1f}MB"
                )
                return

            is_steering = hasattr(
                interaction.user, "roles"
            ) and check_user_roles_with_hierarchy(
                interaction.user.roles, ["Steering Committee"]
            )
            if search_term and not is_steering:
                self._audit_command(
                    interaction=interaction,
                    action="crm.upload_resume",
                    result="denied",
                    metadata={
                        "search_term": search_term,
                        "filename": file.filename,
                        "target_scope": "other",
                        "reason": "missing_required_role_search",
                    },
                )
                await interaction.followup.send(
                    "❌ You must have Steering Committee role or higher to upload a resume for another contact."
                )
                return
            if (
                link_user is not None
                and link_user.id != interaction.user.id
                and not is_steering
            ):
                self._audit_command(
                    interaction=interaction,
                    action="crm.upload_resume",
                    result="denied",
                    metadata={
                        "filename": file.filename,
                        "target_scope": "other",
                        "reason": "missing_required_role_link_user",
                    },
                )
                await interaction.followup.send(
                    "❌ You must have Steering Committee role or higher to upload a resume for another Discord user."
                )
                return
            if not is_steering:
                link_user = None

            # Read attachment once for Steering Committee resume-based inference.
            file_content = await file.read()

            # Determine target contact
            target_contact = None
            target_scope = "self"
            inferred_contact_meta: dict[str, Any] | None = None

            if search_term:
                contacts = await self._search_contact_for_linking(search_term)
                if not contacts:
                    self._audit_command(
                        interaction=interaction,
                        action="crm.upload_resume",
                        result="error",
                        metadata={
                            "search_term": search_term,
                            "filename": file.filename,
                            "contact_found": False,
                            "target_scope": "other",
                        },
                    )
                    await interaction.followup.send(
                        f"❌ No contact found for: `{search_term}`"
                    )
                    return
                if len(contacts) > 1:
                    self._audit_command(
                        interaction=interaction,
                        action="crm.upload_resume",
                        result="error",
                        metadata={
                            "search_term": search_term,
                            "filename": file.filename,
                            "contact_found": False,
                            "target_scope": "other",
                            "reason": "multiple_contacts",
                        },
                    )
                    await interaction.followup.send(
                        f"⚠️ Multiple contacts found for `{search_term}`. "
                        "Please be more specific or use the contact ID."
                    )
                    return
                target_contact = contacts[0]
                target_scope = "other"
            elif is_steering and link_user is not None:
                # Uploading for linked user (Steering Committee+ path)
                target_scope = "other"
                target_contact = await self._find_contact_by_discord_id(
                    str(link_user.id)
                )
                if not target_contact:
                    create_payload = (
                        await self._build_contact_payload_for_link_user_async(
                            user=link_user,
                            file_content=file_content,
                            filename=file.filename,
                        )
                    )
                    self._audit_command(
                        interaction=interaction,
                        action="crm.upload_resume",
                        result="error",
                        metadata={
                            "filename": file.filename,
                            "target_scope": target_scope,
                            "reason": "discord_not_linked",
                            "stage": "create_contact_prompt_shown",
                            "link_user_id": str(link_user.id),
                        },
                    )
                    view = ResumeCreateContactView(
                        crm_cog=self,
                        interaction=interaction,
                        file_content=file_content,
                        filename=file.filename,
                        file_size=file.size,
                        search_term=search_term,
                        overwrite=overwrite,
                        link_user=link_user,
                        inferred_contact_meta={
                            "reason": "discord_not_linked",
                            "link_user_id": str(link_user.id),
                        },
                        target_scope=target_scope,
                        create_payload_override=create_payload,
                        created_target_scope="other_autocreated",
                    )
                    await interaction.followup.send(
                        "⚠️ The provided Discord user is not linked to a CRM contact. "
                        "Would you like to create a new contact for this Discord user "
                        "from the resume details?",
                        view=view,
                    )
                    return
            elif is_steering:
                # Uploading for contact inferred from resume content
                (
                    target_contact,
                    inferred_contact_meta,
                ) = await self._infer_contact_from_resume(
                    file_content,
                    filename=file.filename,
                )
                if not target_contact:
                    inferred_method = (inferred_contact_meta or {}).get("method")
                    inferred_value = (inferred_contact_meta or {}).get("value")
                    inferred_reason = (inferred_contact_meta or {}).get(
                        "reason", "resume_contact_not_found"
                    )
                    inferred_attempts = (inferred_contact_meta or {}).get("attempts")
                    inference_metadata = {
                        "filename": file.filename,
                        "target_scope": "resume_inferred",
                        "reason": inferred_reason,
                    }
                    if inferred_method:
                        inference_metadata["inferred_method"] = inferred_method
                    if inferred_value:
                        inference_metadata["inferred_value"] = inferred_value
                    if inferred_attempts is not None:
                        inference_metadata["inferred_attempts"] = inferred_attempts

                    inferred_attempts_text = (
                        await self._build_inference_lookup_summary_async(
                            file_content=file_content,
                            attempts=inferred_attempts
                            if isinstance(inferred_attempts, list)
                            else None,
                            filename=file.filename,
                        )
                    )

                    if inferred_reason == "multiple_matches" and inferred_value:
                        self._audit_command(
                            interaction=interaction,
                            action="crm.upload_resume",
                            result="error",
                            metadata=inference_metadata,
                        )
                        await interaction.followup.send(
                            f"⚠️ Multiple contacts match `{inferred_value}` from the resume. "
                            "Please provide `search_term` or `link_user`."
                            + inferred_attempts_text,
                            ephemeral=True,
                        )
                    elif inferred_reason == "no_matching_contact":
                        self._audit_command(
                            interaction=interaction,
                            action="crm.upload_resume",
                            result="error",
                            metadata=inference_metadata,
                        )
                        view = ResumeCreateContactView(
                            crm_cog=self,
                            interaction=interaction,
                            file_content=file_content,
                            filename=file.filename,
                            file_size=file.size,
                            search_term=search_term,
                            overwrite=overwrite,
                            link_user=link_user,
                            inferred_contact_meta=inferred_contact_meta,
                            target_scope="resume_inferred",
                        )
                        await interaction.followup.send(
                            "⚠️ Could not find a unique contact from this resume. "
                            "Would you like to create a new contact from the parsed details?"
                            + inferred_attempts_text
                            + await self._build_resume_parsed_identity_summary_async(
                                file_content,
                                filename=file.filename,
                            ),
                            view=view,
                            ephemeral=True,
                        )
                    else:
                        self._audit_command(
                            interaction=interaction,
                            action="crm.upload_resume",
                            result="error",
                            metadata=inference_metadata,
                        )
                        await interaction.followup.send(
                            "⚠️ Resume-based contact inference failed. "
                            "Please provide `search_term` or `link_user`."
                        )
                    return
                target_scope = "resume"
            else:
                # Uploading own resume - find contact by Discord user ID
                target_contact = await self._find_contact_by_discord_id(
                    str(interaction.user.id)
                )
                if not target_contact:
                    self._audit_command(
                        interaction=interaction,
                        action="crm.upload_resume",
                        result="error",
                        metadata={
                            "filename": file.filename,
                            "target_scope": "self",
                            "reason": "discord_not_linked",
                        },
                    )
                    await interaction.followup.send(
                        "❌ Your Discord account is not linked to a CRM contact. "
                        "Please ask a Steering Committee member to link your account first."
                    )
                    return

            assert target_contact is not None
            await self._upload_resume_attachment_to_contact(
                interaction=interaction,
                file_content=file_content,
                filename=file.filename,
                file_size=file.size,
                contact=target_contact,
                target_scope=target_scope,
                search_term=search_term,
                overwrite=overwrite,
                link_user=link_user,
                inferred_contact_meta=inferred_contact_meta,
            )

        except Exception as e:
            logger.error(f"Unexpected error in upload_resume: {e}")
            self._audit_command(
                interaction=interaction,
                action="crm.upload_resume",
                result="error",
                metadata={
                    "search_term": search_term,
                    "filename": file.filename,
                    "error": str(e),
                },
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while uploading the resume."
            )

    @app_commands.command(
        name="reprocess-resume",
        description="Reprocess a resume and re-run profile extraction",
    )
    @app_commands.describe(search_term="Email, 508 username, 508 email, or contact ID.")
    async def reprocess_resume(
        self, interaction: discord.Interaction, search_term: str
    ) -> None:
        """Re-run resume extraction for a contact's latest resume."""
        try:
            await interaction.response.defer(ephemeral=True)

            if not settings.api_shared_secret:
                self._audit_command(
                    interaction=interaction,
                    action="crm.reprocess_resume",
                    result="error",
                    metadata={
                        "search_term": search_term,
                        "reason": "api_shared_secret_missing",
                    },
                )
                await interaction.followup.send(
                    "❌ API_SHARED_SECRET is not configured for backend API access."
                )
                return

            is_steering = hasattr(
                interaction.user, "roles"
            ) and check_user_roles_with_hierarchy(
                interaction.user.roles, ["Steering Committee"]
            )
            if not is_steering:
                self._audit_command(
                    interaction=interaction,
                    action="crm.reprocess_resume",
                    result="denied",
                    metadata={
                        "search_term": search_term,
                        "target_scope": "other",
                        "reason": "missing_required_role",
                    },
                )
                await interaction.followup.send(
                    "❌ You must have Steering Committee role or higher to reprocess another resume."
                )
                return

            contacts = await self._search_contacts_for_reprocess_resume(search_term)
            if not contacts:
                self._audit_command(
                    interaction=interaction,
                    action="crm.reprocess_resume",
                    result="denied",
                    metadata={
                        "search_term": search_term,
                        "contact_found": False,
                        "target_scope": "other",
                    },
                )
                await interaction.followup.send(
                    f"❌ No contact found for: `{search_term}`"
                )
                return

            if len(contacts) > 1:
                self._audit_command(
                    interaction=interaction,
                    action="crm.reprocess_resume",
                    result="denied",
                    metadata={
                        "search_term": search_term,
                        "contact_found": False,
                        "target_scope": "other",
                        "reason": "multiple_contacts",
                    },
                )
                await interaction.followup.send(
                    f"⚠️ Multiple contacts found for `{search_term}`. "
                    "Please be more specific or use the contact ID."
                )
                return

            contact = contacts[0]
            contact_id = str(contact.get("id", ""))
            if not contact_id:
                self._audit_command(
                    interaction=interaction,
                    action="crm.reprocess_resume",
                    result="error",
                    metadata={
                        "search_term": search_term,
                        "reason": "contact_id_missing",
                    },
                )
                await interaction.followup.send("❌ Contact ID not found.")
                return

            contact_name = str(contact.get("name", "Unknown"))
            (
                attachment_id,
                filename,
            ) = await self._get_latest_resume_attachment_for_contact(contact_id)
            if not attachment_id:
                self._audit_command(
                    interaction=interaction,
                    action="crm.reprocess_resume",
                    result="denied",
                    metadata={
                        "search_term": search_term,
                        "contact_id": contact_id,
                        "contact_name": contact_name,
                        "stage": "no_resume_on_file",
                    },
                    resource_type="crm_contact",
                    resource_id=contact_id,
                )
                await interaction.followup.send(
                    f"❌ No resume found for `{contact_name}`. Upload a resume first."
                )
                return

            self._audit_command(
                interaction=interaction,
                action="crm.reprocess_resume",
                result="success",
                metadata={
                    "search_term": search_term,
                    "contact_id": contact_id,
                    "contact_name": contact_name,
                    "attachment_id": attachment_id,
                    "filename": filename,
                    "stage": "reprocess_confirmation_prompt",
                },
                resource_type="crm_contact",
                resource_id=contact_id,
            )

            display_filename = filename or "latest resume"
            view = ResumeReprocessConfirmationView(
                crm_cog=self,
                interaction=interaction,
                contact_id=contact_id,
                contact_name=contact_name,
                attachment_id=attachment_id,
                filename=display_filename,
            )
            await interaction.followup.send(
                f"⚠️ Reprocess resume `{display_filename}` for `{contact_name}`?",
                view=view,
                ephemeral=True,
            )
        except Exception as e:
            logger.error("Unexpected error in reprocess_resume: %s", e)
            self._audit_command(
                interaction=interaction,
                action="crm.reprocess_resume",
                result="error",
                metadata={
                    "search_term": search_term,
                    "error": str(e),
                },
            )
            await interaction.followup.send(
                "❌ An unexpected error occurred while reprocessing the resume."
            )

    @app_commands.command(
        name="match-candidates",
        description="Reads this thread's opening message as a job posting and returns ranked matching candidates.",
    )
    @require_role("Member")
    async def match_candidates(
        self,
        interaction: discord.Interaction,
    ) -> None:
        """Parse the thread's starter message and find matching candidates ranked by fit.

        Must be invoked inside a thread. The starter message is used as the job posting text.
        The response is posted publicly in the thread.
        """
        if not isinstance(interaction.channel, discord.Thread):
            await interaction.response.send_message(
                "⚠️ This command must be used inside a thread. "
                "Open a thread on the job posting message and run `/match-candidates` there.",
                ephemeral=True,
            )
            return

        thread: discord.Thread = interaction.channel
        starter = thread.starter_message
        if starter is None:
            try:
                starter = await thread.fetch_message(thread.id)
            except Exception:
                starter = None

        if starter is None or not starter.content.strip():
            await interaction.response.send_message(
                "⚠️ Could not read the thread's opening message. "
                "Make sure the thread was created from a job posting message.",
                ephemeral=True,
            )
            return

        posting = starter.content
        if thread.applied_tags:
            tag_names = ", ".join(t.name for t in thread.applied_tags)
            posting = f"Thread tags: {tag_names}\n\n{posting}"

        await interaction.response.defer(ephemeral=False)

        try:
            requirements = await asyncio.to_thread(
                extract_job_requirements,
                posting,
                api_key=settings.openai_api_key,
                base_url=settings.openai_base_url or None,
                model=settings.openai_model,
                webhook_url=settings.discord_logs_webhook_url,
            )
        except RuntimeError as exc:
            self._audit_command(
                interaction=interaction,
                action="crm.match_candidates",
                result="error",
                metadata={"stage": "extract_requirements", "error": str(exc)},
            )
            await interaction.followup.send(
                f"❌ Failed to analyze the job posting: {exc}",
                ephemeral=True,
            )
            return

        if not requirements.required_skills:
            self._audit_command(
                interaction=interaction,
                action="crm.match_candidates",
                result="denied",
                metadata={"stage": "no_required_skills_extracted"},
            )
            await interaction.followup.send(
                "⚠️ No required skills could be extracted from this posting. "
                "Please include explicit skill requirements and try again.",
                ephemeral=True,
            )
            return

        try:
            candidates = await asyncio.to_thread(
                search_candidates,
                settings,
                requirements,
                limit=20,
                min_match_score=8.0,
            )
        except Exception as exc:
            logger.error("Candidate search failed: %s", exc)
            self._audit_command(
                interaction=interaction,
                action="crm.match_candidates",
                result="error",
                metadata={"stage": "search_candidates", "error": str(exc)},
            )
            await interaction.followup.send(
                "❌ Candidate search failed. Please try again later.",
                ephemeral=True,
            )
            return

        lines: list[str] = []

        header_parts: list[str] = []
        if requirements.title:
            header_parts.append(f"**{requirements.title}**")
        if requirements.discord_role_types:
            header_parts.append(
                "Role: " + ", ".join(f"`{r}`" for r in requirements.discord_role_types)
            )
        if requirements.required_skills:
            header_parts.append(
                "Skills: "
                + ", ".join(f"`{s}`" for s in requirements.required_skills[:8])
            )
        if requirements.seniority:
            header_parts.append(f"Seniority: `{requirements.seniority}`")
        if requirements.location_type == "us_only":
            header_parts.append("📍 US only")
        elif requirements.raw_location_text:
            header_parts.append(f"📍 {requirements.raw_location_text}")

        lines.append("## Job Match Results")
        if header_parts:
            lines.append(" · ".join(header_parts))
        lines.append(f"Found **{len(candidates)}** candidate(s).\n")

        crm_base = settings.espo_base_url.rstrip("/")
        resume_options: list[tuple[str, str, str]] = []

        for i, c in enumerate(candidates, start=1):
            label = "**[Member]**" if c.is_member else "[Prospect]"
            name = c.name or "Unknown"
            email = c.email_508 or c.email or "—"
            crm_link = (
                f"{crm_base}/#Contact/view/{c.crm_contact_id}"
                if c.has_crm_link and c.crm_contact_id
                else None
            )
            if crm_link:
                parts = [f"{i}. {label} [{name}](<{crm_link}>) · {email}"]
            else:
                parts = [f"{i}. {label} {name} · {email}"]
                if c.discord_user_id:
                    parts.append(f"Discord: <@{c.discord_user_id}>")

            if c.linkedin:
                parts.append(f"[LinkedIn](<{c.linkedin}>)")
            if c.latest_resume_id and c.latest_resume_name:
                parts.append(f"Resume: `{c.latest_resume_name}`")
                resume_options.append((name, c.latest_resume_id, c.latest_resume_name))

            skill_info: list[str] = []
            skill_info.append(f"score: {c.match_score:.1f}")
            if c.matched_required_skills:
                skill_info.append(
                    "✅ " + ", ".join(f"`{s}`" for s in c.matched_required_skills[:5])
                )
            if c.matched_discord_roles:
                skill_info.append(
                    "🏷️ " + ", ".join(f"`{r}`" for r in c.matched_discord_roles)
                )
            if c.seniority:
                skill_info.append(f"seniority: `{c.seniority}`")
            if c.timezone:
                skill_info.append(f"tz: `{c.timezone}`")
            if skill_info:
                parts.append("   " + " · ".join(skill_info))

            lines.append("\n".join(parts))

        # Paginate: Discord followup allows multiple sends; split on 1900-char chunks.
        messages: list[str] = []
        current = ""
        for line in lines:
            candidate_block = line + "\n"
            if len(current) + len(candidate_block) > 1900:
                messages.append(current.rstrip())
                current = candidate_block
            else:
                current += candidate_block
        if current.strip():
            messages.append(current.rstrip())

        for msg in messages:
            await interaction.followup.send(msg)
        if resume_options:
            await interaction.followup.send(
                "Resume download:",
                view=MatchResumeSelectView(resume_options),
            )

        self._audit_command(
            interaction=interaction,
            action="crm.match_candidates",
            result="success",
            metadata={
                "title": requirements.title,
                "required_skills_count": len(requirements.required_skills),
                "preferred_skills_count": len(requirements.preferred_skills),
                "discord_role_types": requirements.discord_role_types,
                "candidates_returned": len(candidates),
            },
        )

    async def _bulk_sync_guild_roles(
        self, guild: discord.Guild
    ) -> tuple[int, int, int]:
        """Sync discord_roles for all non-bot guild members.

        Returns (updated, skipped, failed). Per-member failures are logged and
        skipped so one bad record never aborts the full run.
        Roles in DISCORD_ROLES_EXCLUDE_FROM_SYNC (Bots, FixTweet, @everyone)
        are excluded from the stored list.
        """
        updated = 0
        skipped = 0
        failed = 0
        for member in guild.members:
            if member.bot:
                continue
            role_names = [
                r.name
                for r in member.roles
                if r.name not in DISCORD_ROLES_EXCLUDE_FROM_SYNC
            ]
            try:
                await asyncio.to_thread(
                    upsert_discord_member,
                    settings,
                    discord_user_id=str(member.id),
                    guild_id=str(guild.id),
                    discord_username=member.name,
                    display_name=member.display_name,
                    roles=role_names,
                )
                did_update = await asyncio.to_thread(
                    update_person_discord_roles,
                    settings,
                    str(member.id),
                    role_names,
                )
            except Exception as exc:
                failed += 1
                logger.warning(
                    "bulk role sync: failed for user_id=%s: %s", member.id, exc
                )
                continue
            if did_update:
                updated += 1
            else:
                skipped += 1
        return updated, skipped, failed

    @commands.Cog.listener()
    async def on_ready(self) -> None:
        """Bulk-sync all guild member roles on startup."""
        for guild in self.bot.guilds:
            try:
                updated, skipped, failed = await self._bulk_sync_guild_roles(guild)
                logger.info(
                    "Startup discord role sync: guild=%s updated=%d skipped=%d failed=%d",
                    guild.name,
                    updated,
                    skipped,
                    failed,
                )
            except Exception as exc:
                logger.warning(
                    "Startup discord role sync failed for guild %s: %s", guild.name, exc
                )

    @app_commands.command(
        name="sync-discord-roles",
        description="Re-sync all server members' Discord roles into the candidate database.",
    )
    @require_role("Steering Committee")
    async def sync_discord_roles(
        self,
        interaction: discord.Interaction,
    ) -> None:
        """Manually trigger a full guild role sync (also runs automatically on startup)."""
        guild = interaction.guild
        if guild is None:
            await interaction.response.send_message(
                "⚠️ This command must be used inside a server.", ephemeral=True
            )
            return

        await interaction.response.defer(ephemeral=True)

        updated, skipped, failed = await self._bulk_sync_guild_roles(guild)

        self._audit_command(
            interaction=interaction,
            action="crm.sync_discord_roles",
            result="success",
            metadata={
                "updated": updated,
                "skipped_no_db_match": skipped,
                "failed": failed,
                "total_members_scanned": updated + skipped + failed,
            },
        )

        await interaction.followup.send(
            f"✅ Discord role sync complete.\n"
            f"Updated: **{updated}** · No DB match (skipped): **{skipped}** · Failed: **{failed}**",
            ephemeral=True,
        )

    @commands.Cog.listener()
    async def on_member_update(
        self,
        before: discord.Member,
        after: discord.Member,
    ) -> None:
        """Automatically sync discord_roles when a member's roles change."""
        if before.roles == after.roles:
            return

        if after.guild is None:
            return

        role_names = [
            r.name for r in after.roles if r.name not in DISCORD_ROLES_EXCLUDE_FROM_SYNC
        ]

        try:
            await asyncio.to_thread(
                upsert_discord_member,
                settings,
                discord_user_id=str(after.id),
                guild_id=str(after.guild.id),
                discord_username=after.name,
                display_name=after.display_name,
                roles=role_names,
            )
            await asyncio.to_thread(
                update_person_discord_roles,
                settings,
                str(after.id),
                role_names,
            )
        except Exception as exc:
            logger.warning(
                "on_member_update: failed to sync roles for user %s: %s",
                after.id,
                exc,
            )


async def setup(bot: commands.Bot) -> None:
    """Add the CRM cog to the bot."""
    cog = CRMCog(bot)
    await bot.add_cog(cog)
    # Slash commands will be synced automatically in bot.py
